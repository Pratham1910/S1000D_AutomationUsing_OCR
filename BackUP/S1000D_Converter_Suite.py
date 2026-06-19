#!/usr/bin/env python3
"""
S1000D Converter Suite
======================
Converts raw files (PDF · DOCX · TXT · MD · Images · Scanned PDFs) into a
complete S1000D output package in one click:

    01_raw_json/          Raw element list (glmocr / pymupdf / python-docx)
    02_semantic_json/     Same list + "semantic" classification per element
    03_s1000d_xml/        S1000D Issue 4.2 XML data modules (all 21 schemas)
    04_adoc/              AsciiDoc source files (ready for Asciidoctor backend)
    conversion_report.txt Summary of every file processed

Supported input types
    .pdf        Text-layer PDF  (pymupdf) or scanned/image PDF (glmocr)
    .docx       Word document   (python-docx → JSON;  pandoc → AsciiDoc)
    .txt        Plain text
    .md .markdown  Markdown
    .png .jpg .jpeg .tiff .bmp .webp  Raster images (glmocr)

Dependencies (install what you need; others degrade gracefully):
    pip install pymupdf python-docx tqdm requests
    pip install glmocr          # for scanned PDFs and images
    pip install opendataloader-pdf  # PDF extraction alternative when not using GLM OCR
    pandoc (CLI)                # for DOCX / MD → AsciiDoc conversion
    
"""

# ─── stdlib ─────────────────────────────────────────────────────────────────
import os, sys, json, re, subprocess, threading, traceback, time, shutil, html, hashlib
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from xml.etree import ElementTree as ET
from xml.dom import minidom
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, font as tkfont, messagebox
 
# ─── Optional heavy deps – detected at runtime ──────────────────────────────
try:
    import fitz as _fitz          # pymupdf
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import docx as _docx          # python-docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False  

try:
    from glmocr import GlmOcr as _GlmOcr
    HAS_GLMOCR = True
except ImportError:
    HAS_GLMOCR = False

try:
    from opendataloader_pdf import convert as _odl_pdf_convert
    HAS_OPENDATALOADER_PDF = True
except ImportError:
    HAS_OPENDATALOADER_PDF = False

try:
    from PIL import Image as _PILImage
    import io as _pil_io
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# ────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ────────────────────────────────────────────────────────────────────────────

IMAGE_EXTS  = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}
TEXT_EXTS   = {".txt", ".text"}
MD_EXTS     = {".md", ".markdown"}
MD_LINK_RE  = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", re.IGNORECASE)
URL_RE      = re.compile(r'\b(https?://[^\s<>"]+)', re.IGNORECASE)
WWW_RE      = re.compile(r'\b(www\.[^\s<>"]+)', re.IGNORECASE)
DOMAIN_RE   = re.compile(
    r"\b((?:[a-z0-9-]+\.)+(?:com|org|net|edu|gov|mil|io|co|in|eu|info|biz|ai|dev|tech)"
    r"(?:/[^\s<>\"]*)?)\b",
    re.IGNORECASE,
)

# All 21 S1000D DM schema types
S1000D_DM_TYPES = [
    "auto",
    "descript", "procedure", "fault", "proced", "sched",
    "container", "crew", "sb", "pim", "chkl", "learning",
    "frontmatter", "appliccrossreftable", "condcrossreftable",
    "functionalitem", "partrepository", "illustratedpartscatalog",
    "wrngdata", "comrepository", "brex", "techrep",
]

DM_TYPE_DESC = {
    "auto":                    "Auto detect from filename + content",
    "descript":                "Descriptive",
    "procedure":               "Procedural (generic)",
    "fault":                   "Fault isolation",
    "proced":                  "Procedural (structured)",
    "sched":                   "Scheduled maintenance",
    "container":               "Container",
    "crew":                    "Crew / operator",
    "sb":                      "Service bulletin",
    "pim":                     "Preliminary information",
    "chkl":                    "Checklist",
    "learning":                "Learning / CBT",
    "frontmatter":             "Front matter",
    "appliccrossreftable":     "Applicability XRef table",
    "condcrossreftable":       "Condition XRef table",
    "functionalitem":          "Functional item repository",
    "partrepository":          "Parts repository",
    "illustratedpartscatalog": "Illustrated parts catalog (IPC)",
    "wrngdata":                "Wiring data",
    "comrepository":           "Comments repository",
    "brex":                    "Business rules exchange",
    "techrep":                 "Technical report",
}

# XSD schema base URL
S1000D_42_XSD = "http://www.s1000d.org/S1000D_4-2/xml_schema_flat/{dm_type}.xsd"

# Schema types that have dedicated XML content builders (others use generic)
RICH_XML_TYPES = {"descript", "proced", "procedure", "fault", "sched",
                  "chkl", "container", "frontmatter"}

# Pandoc-aware types (prefer pandoc for AsciiDoc generation)
PANDOC_IN_EXTS = {".docx", ".md", ".markdown", ".txt", ".text"}

TEMPLATE_MAP = {
    "procedure": "DMC-ProceduralData.adoc",
    "proced": "DMC-ProceduralData.adoc",
    "fault": "Type_fault.adoc",
    "sched": "Type_sched.adoc",
    "illustratedpartscatalog": "IPD.adoc",
}

OLLAMA_URL = os.environ.get("S1000D_OLLAMA_URL", "http://127.0.0.1:11434/api/generate")
OLLAMA_MODEL = os.environ.get("S1000D_OLLAMA_MODEL", "llama3.1:8b")
USE_OLLAMA_TEMPLATE = os.environ.get("S1000D_USE_OLLAMA_TEMPLATE", "1") not in ("0", "false", "False")
USE_OLLAMA_TABLE_MERGE = os.environ.get("S1000D_USE_OLLAMA_TABLE_MERGE", "1") not in ("0", "false", "False")
S1000D_FORCE_DESCRIPT_ONLY = os.environ.get("S1000D_FORCE_DESCRIPT_ONLY", "1") not in ("0", "false", "False")

GLMOCR_BACKEND = os.environ.get("S1000D_GLMOCR_BACKEND", "ollama").strip().lower()
if GLMOCR_BACKEND not in {"default", "ollama"}:
    GLMOCR_BACKEND = "default"
GLMOCR_OLLAMA_URL = os.environ.get("S1000D_GLMOCR_OLLAMA_URL", "http://127.0.0.1:11434/api/generate")
GLMOCR_OLLAMA_MODEL = os.environ.get("S1000D_GLMOCR_OLLAMA_MODEL", "glm-ocr:latest")
GLMOCR_LAYOUT_MODEL_DIR = os.environ.get("S1000D_GLMOCR_LAYOUT_MODEL_DIR", "").strip()

# Number of pages OCR'd concurrently. The backend (GPU / Ollama server) is the
# real bottleneck — raising this past what the server can serve in parallel
# just queues requests. With Ollama, also raise OLLAMA_NUM_PARALLEL server-side.
try:
    OCR_WORKERS = max(1, int(os.environ.get("S1000D_OCR_WORKERS", "4")))
except ValueError:
    OCR_WORKERS = 4

# OpenDataLoader-PDF config (fallback high-accuracy PDF extractor)
ODL_USE_HYBRID  = os.environ.get("S1000D_ODL_HYBRID",      "0") not in ("0", "false", "False")
ODL_HYBRID_URL  = os.environ.get("S1000D_ODL_HYBRID_URL",  "http://127.0.0.1:5002")


def _autodetect_layout_model_dir() -> str:
    """Best-effort local PP-DocLayout path resolution from HF cache."""
    # Priority 1: exact snapshot under user cache.
    cache_root = Path.home() / ".cache" / "huggingface" / "hub"
    snapshots_dir = cache_root / "models--PaddlePaddle--PP-DocLayoutV3_safetensors" / "snapshots"
    if snapshots_dir.is_dir():
        candidates = sorted([p for p in snapshots_dir.iterdir() if p.is_dir()], key=lambda p: p.stat().st_mtime, reverse=True)
        for cand in candidates:
            if (cand / "preprocessor_config.json").is_file() and (cand / "config.json").is_file():
                return str(cand)

    # Priority 2: model root itself (works if files are in root, no snapshots).
    root_dir = cache_root / "models--PaddlePaddle--PP-DocLayoutV3_safetensors"
    if root_dir.is_dir() and (root_dir / "preprocessor_config.json").is_file() and (root_dir / "config.json").is_file():
        return str(root_dir)

    return ""

# ────────────────────────────────────────────────────────────────────────────
# ELEMENT NORMALIZATION
# ────────────────────────────────────────────────────────────────────────────

def _el(content: str, native_label: str = "paragraph",
        page_no: int = 0, meta: Optional[dict] = None) -> Dict:
    """Create a normalised element dict compatible with semantic_annotate /
    json_to_s1000d_xml conventions."""
    e = {"content": content, "native_label": native_label, "page_no": page_no}
    if meta:
        e.update(meta)
    return e


def _enforce_descript_only(dm_type: str, log=None) -> str:
    """Force all flows to `descript` when descript-only mode is enabled."""
    normalized = (dm_type or "").strip().lower() or "descript"
    if not S1000D_FORCE_DESCRIPT_ONLY:
        return normalized
    if normalized != "descript" and log:
        log(f"  Descript-only mode active: overriding DM type '{dm_type}' -> 'descript'.")
    return "descript"


def _chunk_pages(elements: List[Dict], page_size: int = 40) -> List[List[Dict]]:
    """Group flat element list into page-sized buckets (for JSON structure)."""
    if not elements:
        return []
    pages: Dict[int, List[Dict]] = {}
    for el in elements:
        pn = el.get("page_no", 0)
        pages.setdefault(pn, []).append(el)
    # Sort by page number
    return [pages[k] for k in sorted(pages)]


def _clean_link_text(text: str) -> str:
    """Normalize extracted link text for downstream rendering."""
    if not text:
        return ""
    return " ".join(str(text).split())


def _md_links_to_text_and_url(text: str) -> str:
    """Convert markdown links to plain text + URL for XML-safe storage."""
    if not text:
        return ""
    return MD_LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)


def _text_to_adoc_links(text: str) -> str:
    """Convert markdown and plain URLs into AsciiDoc link syntax."""
    if not text:
        return ""

    converted = MD_LINK_RE.sub(lambda m: f"link:{m.group(2)}[{m.group(1)}]", text)

    # Convert bare URLs that are not already inside link:...[]
    def _url_repl(match):
        url = match.group(1)
        prefix = converted[max(0, match.start() - 6):match.start()].lower()
        if prefix.endswith("link:"):
            return url
        return f"link:{url}[{url}]"

    return URL_RE.sub(_url_repl, converted)


def _looks_like_merged_table_text(text: str) -> bool:
    """Heuristic for AsciiDoc tables that likely require merge syntax handling."""
    if not text:
        return False
    t = text
    if "|===" in t and re.search(r"(?m)^\s*\d+\+\^?\|", t):
        return True
    if re.search(r"(?m)^\s*\d+\+\^?\|", t):
        return True
    return False


def _looks_like_flattened_pipe_table(text: str) -> bool:
    """Detect OCR-flattened tables represented as one long '|' stream."""
    t = (text or "").strip()
    if not t:
        return False
    if "|" not in t:
        return False
    # Typical failure mode: no line breaks, many pipe separators.
    if "\n" not in t and t.count("|") >= 8:
        return True
    return False


def _basic_table_tokens(text: str) -> List[str]:
    """Extract coarse lexical tokens used for safety-checking model table rewrites."""
    toks = re.findall(r"[A-Za-z0-9][A-Za-z0-9\-_/\.]*", text or "")
    # Ignore very small noisy tokens.
    return [t for t in toks if len(t) >= 2]


def _is_plausible_adoc_table(text: str) -> bool:
    if not text:
        return False
    t = text.strip()
    delim_lines = re.findall(r"(?m)^\s*\|===\s*$", t)
    if len(delim_lines) < 2:
        return False
    return True


def _is_well_formed_adoc_table(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    lines = [ln.rstrip() for ln in t.splitlines() if ln.strip()]
    if len(lines) < 3:
        return False
    if lines[0].strip() != "|===" or lines[-1].strip() != "|===":
        return False
    return True


def _ollama_repair_merged_table(table_text: str) -> Optional[str]:
    """Use Ollama to normalize merged-cell AsciiDoc table blocks.

    Returns repaired table text, or None if unavailable/invalid.
    """
    if not USE_OLLAMA_TABLE_MERGE:
        return None
    if not _looks_like_merged_table_text(table_text):
        return None

    try:
        import requests

        # Cap table input — large tables can't be meaningfully repaired by LLM anyway
        if len(table_text) > 6000:
            return None

        prompt = (
            "You are an AsciiDoc table expert.\n"
            "Repair the following table into valid AsciiDoc table markup.\n"
            "Use merge syntax like 2+|, 6+^| where needed.\n"
            "Hard rules:\n"
            "1) Return only the table block between |=== and |===.\n"
            "2) Preserve all source content words; do not drop data.\n"
            "3) Do not add explanations.\n\n"
            "TABLE INPUT:\n"
            f"{table_text[:6000]}\n"
        )

        endpoints: List[Tuple[str, Dict[str, Any], str]] = []
        base = OLLAMA_URL
        if "/api/" in base:
            root = base.split("/api/", 1)[0]
        else:
            root = base.rstrip("/")

        _opts = {"temperature": 0, "top_p": 0.9, "num_predict": 2048}
        endpoints.append((
            base,
            {
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": _opts,
            },
            "generate",
        ))
        endpoints.append((
            root + "/api/chat",
            {
                "model": OLLAMA_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "options": _opts,
            },
            "chat",
        ))

        out = ""
        for url, payload, mode in endpoints:
            try:
                resp = requests.post(url, json=payload, timeout=(10, 45))
                resp.raise_for_status()
                data = resp.json()
                if mode == "generate":
                    out = (data.get("response") or "").strip()
                else:
                    out = ((data.get("message") or {}).get("content") or "").strip()
                if out:
                    break
            except Exception:
                continue

        if not out:
            return None
        if not _is_plausible_adoc_table(out):
            return None

        # Lightweight data-loss guard: keep most original lexical tokens.
        src_toks = _basic_table_tokens(table_text)
        out_norm = out.lower()
        if src_toks:
            kept = sum(1 for tok in src_toks if tok.lower() in out_norm)
            if kept / max(len(src_toks), 1) < 0.80:
                return None

        return out
    except Exception:
        return None


def _sanitize_cell_for_table(cell_text: str) -> str:
    """Remove/normalize image macros and extra whitespace in table cell text."""
    # Table cells should use inline image syntax (image:...[]), not block image syntax.
    cell_text = re.sub(r"\bimage::", "image:", cell_text)
    # Replace inline image macro references with a compact placeholder.
    cell_text = re.sub(r"image:{1,2}[^\[]*\[.*?\]", "[img]", cell_text)
    # Normalize runs of whitespace
    cell_text = re.sub(r"\s+", " ", cell_text).strip()
    return cell_text


def _html_table_to_slots(raw_table_text: str) -> Optional[List[List[Any]]]:
    """Parse a simple HTML table into slot rows preserving colspan/rowspan."""
    raw = (raw_table_text or "").strip()
    if not raw or "<table" not in raw.lower():
        return None

    m_table = re.search(r"<table\b[^>]*>(.*?)</table>", raw, re.IGNORECASE | re.DOTALL)
    if not m_table:
        return None

    table_inner = m_table.group(1)
    tr_matches = list(re.finditer(r"<tr\b[^>]*>(.*?)</tr>", table_inner, re.IGNORECASE | re.DOTALL))
    if not tr_matches:
        return None

    pending_rowspan: Dict[int, int] = {}
    rows: List[List[Any]] = []

    for trm in tr_matches:
        tr_html = trm.group(1)
        row: List[Any] = []
        c = 0

        def _consume_pending_at_col(col: int) -> bool:
            rem = int(pending_rowspan.get(col, 0) or 0)
            if rem <= 0:
                return False
            row.append("__MERGE_ROW__")
            if rem == 1:
                pending_rowspan.pop(col, None)
            else:
                pending_rowspan[col] = rem - 1
            return True

        # Place leading row-span continuations before first explicit cell.
        while _consume_pending_at_col(c):
            c += 1

        td_matches = list(re.finditer(r"<t[dh]\b([^>]*)>(.*?)</t[dh]>", tr_html, re.IGNORECASE | re.DOTALL))
        for tdm in td_matches:
            while _consume_pending_at_col(c):
                c += 1

            attr_s = tdm.group(1) or ""
            inner = tdm.group(2) or ""
            inner = re.sub(r"<br\s*/?>", "\n", inner, flags=re.IGNORECASE)
            text = re.sub(r"<[^>]+>", " ", inner)
            text = html.unescape(_sanitize_cell_for_table(text))

            m_cs = re.search(r"\bcolspan\s*=\s*['\"]?(\d+)['\"]?", attr_s, re.IGNORECASE)
            m_rs = re.search(r"\browspan\s*=\s*['\"]?(\d+)['\"]?", attr_s, re.IGNORECASE)
            colspan = max(1, int(m_cs.group(1))) if m_cs else 1
            rowspan = max(1, int(m_rs.group(1))) if m_rs else 1

            row.append({"text": text, "colspan": colspan, "rowspan": rowspan})
            for _ in range(colspan - 1):
                row.append("__MERGE_COL__")

            if rowspan > 1:
                for cc in range(c, c + colspan):
                    cur = int(pending_rowspan.get(cc, 0) or 0)
                    pending_rowspan[cc] = max(cur, rowspan - 1)
            c += colspan

        # Place trailing row-span continuations to keep geometry aligned.
        if pending_rowspan:
            max_pending_col = max(pending_rowspan.keys())
            while c <= max_pending_col:
                if _consume_pending_at_col(c):
                    c += 1
                    continue
                row.append("")
                c += 1

        if row:
            rows.append(row)

    return rows if rows else None


def _heuristic_unflatten_pipe_table(raw_text: str) -> Optional[str]:
    """Best-effort reconstruction of flattened pipe table text into AsciiDoc rows."""
    t = (raw_text or "").strip()
    if not _looks_like_flattened_pipe_table(t):
        return None

    # First pass: split by pipes and clean each cell.
    tokens = [_sanitize_cell_for_table(c) for c in t.split("|")]
    while tokens and not tokens[0]:
        tokens.pop(0)
    while tokens and not tokens[-1]:
        tokens.pop()
    if len(tokens) < 4:
        return None

    # Find the best column count by detecting domain-specific patterns.
    # Maintenance tables commonly have activity name + 4-5 schedule columns.
    col_scores: Dict[int, float] = {}  # column_count -> quality_score
    
    # Count cells that look like schedule headers (contain month/hour/daily patterns).
    schedule_keywords = sum(1 for tok in tokens if re.search(
        r"\b(daily|weekly|month|year|hour|hrs?|schedule|interval)\b", tok, re.IGNORECASE))
    
    # Try column groupings and score them. Prefer more columns (fewer rows).
    for potential_cols in range(2, min(13, len(tokens) // 2 + 1)):
        test_rows = []
        test_cur = []
        for tok in tokens:
            test_cur.append(tok)
            if len(test_cur) == potential_cols:
                test_rows.append(test_cur)
                test_cur = []
        if test_rows:
            complete_rows = len(test_rows)
            
            # Base score: fewer rows is better (prefer more columns).
            score = 1.0 / complete_rows  # Inverted: fewer rows = higher score
            # Bonus: col count aligns with schedule keywords (e.g., 4 + 1 activity = 5 cols)
            if schedule_keywords > 0 and potential_cols in (schedule_keywords, schedule_keywords + 1):
                score *= 2.5
            # Bonus for typical maintenance range
            if 5 <= potential_cols <= 7:
                score *= 1.5
            # Penalty: too few columns (leading to very few rows) might indicate wrong structure
            if complete_rows <= 1:
                score *= 0.3
            
            col_scores[potential_cols] = score

    # Pick the column count with the best score.
    if col_scores:
        ncols = max(col_scores.keys(), key=lambda k: col_scores[k])
    else:
        # Fallback: infer from first contiguous non-empty span.
        lead_non_empty = 0
        for tok in tokens:
            if tok:
                lead_non_empty += 1
            else:
                break
        ncols = lead_non_empty if 2 <= lead_non_empty <= 12 else 0
        if not ncols:
            # Domain-specific fallback for maintenance schedule matrices.
            if re.search(r"\b(daily|weekly|month|months|hrs?|hours?)\b", t, re.IGNORECASE):
                ncols = 5
            elif re.search(r"\btask\b", t, re.IGNORECASE):
                ncols = 6  # Common for maintenance tables: Activity + 5 columns
            else:
                ncols = 4

    if ncols < 2 or ncols > 12:
        return None

    rows: List[List[str]] = []
    cur: List[str] = []
    for tok in tokens:
        cur.append(tok)
        if len(cur) == ncols:
            rows.append(cur)
            cur = []
    if cur:
        if rows:
            # Attach leftovers to the last cell to avoid data loss.
            tail = " | ".join(x for x in cur if x)
            if tail:
                rows[-1][-1] = (rows[-1][-1] + " " + tail).strip()
        else:
            return None

    if not rows:
        return None

    out: List[str] = ["|==="]
    for row in rows:
        # Ensure row has exactly ncols cells, padding with empty as needed.
        row = (row + [""] * ncols)[:ncols]
        out.append("| " + " | ".join(cell for cell in row))
    out.append("|===")
    return "\n".join(out)


def _table_text_to_adoc_block(raw_table_text: str) -> str:
    """Convert extracted table text to an AsciiDoc table block with merge support."""
    raw = (raw_table_text or "").strip()
    if not raw:
        return ""

    # If source is HTML table (common in Markdown outputs), preserve row/col spans.
    if "<table" in raw.lower():
        slots = _html_table_to_slots(raw)
        if slots:
            block = _table_slots_to_adoc_block(slots)
            if block:
                return block.strip()

    # If it's already an adoc table, keep as-is (optionally repaired for merges).
    if "|===" in raw:
        if _is_well_formed_adoc_table(raw):
            return raw.strip()
        repaired = _ollama_repair_merged_table(raw)
        return (repaired or raw).strip()

    # If merge markers are present but delimiters are missing, wrap and optionally repair.
    if _looks_like_merged_table_text(raw):
        wrapped = "|===\n" + raw + "\n|==="
        repaired = _ollama_repair_merged_table(wrapped)
        return (repaired or wrapped).strip()

    # OCR can flatten whole tables into one long pipe-separated line.
    # Try deterministic reconstruction first, then Ollama repair for badly mangled tables.
    if _looks_like_flattened_pipe_table(raw):
        heur = _heuristic_unflatten_pipe_table(raw)
        if heur and _is_plausible_adoc_table(heur):
            # Attempt Ollama repair to improve structure and handle merges
            repaired = _ollama_repair_merged_table(heur)
            if repaired:
                return repaired.strip()
            # If Ollama can't repair, check if heuristic result seems reasonable
            # (at least 3-4 columns); if too narrow, try Ollama on raw table
            col_count = heur.count("|") // max(1, heur.count("\n") - 1)
            if col_count >= 3:
                return heur.strip()
            # Column count is too low; try Ollama on the raw flattened format
            wrapped = "|===\n" + raw + "\n|==="
            repaired = _ollama_repair_merged_table(wrapped)
            return (repaired or heur).strip()
        # Heuristic failed; try Ollama on raw flattened table
        wrapped = "|===\n" + raw + "\n|==="
        repaired = _ollama_repair_merged_table(wrapped)
        return (repaired or wrapped).strip()

    # Fallback plain pipe-split table.
    rows = [r.strip() for r in raw.splitlines() if r.strip()]
    lines: List[str] = ["|==="]
    for r in rows:
        cells = r.split("|")
        lines.append("| " + " | ".join(c.strip() for c in cells))
    lines.append("|===")
    return "\n".join(lines)


def _table_rows_to_adoc_block(table_rows: List[List[str]]) -> str:
    """Render table rows with merge markers into AsciiDoc table syntax.

    Special markers supported:
    - __MERGE_COL__: continuation of previous cell colspan
    - __MERGE_ROW__: vertical-merge continuation (suppressed in row output)
    """
    if not table_rows:
        return ""

    max_cols = max((len(r) for r in table_rows), default=0)
    if max_cols <= 0:
        return ""

    rows = [r + [""] * (max_cols - len(r)) for r in table_rows]
    lines: List[str] = ["|==="]

    for row in rows:
        parts: List[str] = []
        c = 0
        while c < max_cols:
            cell = row[c]
            if cell == "__MERGE_COL__" or cell == "__MERGE_ROW__":
                c += 1
                continue

            span = 1
            j = c + 1
            while j < max_cols and row[j] == "__MERGE_COL__":
                span += 1
                j += 1

            txt = (cell or "").strip()
            if span > 1:
                parts.append(f"{span}+| {txt}")
            else:
                parts.append(f"| {txt}")
            c = j

        if parts:
            lines.append(" ".join(parts))

    lines.append("|===")
    return "\n".join(lines)


def _table_slots_to_adoc_block(table_slots: List[List[Any]]) -> str:
    """Render structured table slots with explicit row/col spans to AsciiDoc."""
    if not table_slots:
        return ""
    max_cols = max((len(r) for r in table_slots), default=0)
    if max_cols <= 0:
        return ""

    rows = [r + [""] * (max_cols - len(r)) for r in table_slots]
    lines: List[str] = ["|==="]

    for row in rows:
        parts: List[str] = []
        c = 0
        while c < max_cols:
            tok = row[c]
            if tok in ("__MERGE_COL__", "__MERGE_ROW__", ""):
                c += 1
                continue

            if isinstance(tok, dict):
                txt = (tok.get("text") or "").strip()
                colspan = max(1, int(tok.get("colspan", 1) or 1))
                rowspan = max(1, int(tok.get("rowspan", 1) or 1))
            else:
                txt = str(tok).strip()
                colspan = 1
                rowspan = 1

            if colspan > 1 and rowspan > 1:
                prefix = f"{colspan}.{rowspan}+| "
            elif colspan > 1:
                prefix = f"{colspan}+| "
            elif rowspan > 1:
                prefix = f".{rowspan}+| "
            else:
                prefix = "| "

            parts.append(prefix + txt)
            c += max(1, colspan)

        if parts:
            lines.append(" ".join(parts))

    lines.append("|===")
    return "\n".join(lines)


def _table_slots_to_html(table_slots: List[List[Any]]) -> str:
    """Render structured table slots to HTML for Markdown fallback (supports rowspan/colspan)."""
    if not table_slots:
        return ""
    max_cols = max((len(r) for r in table_slots), default=0)
    if max_cols <= 0:
        return ""

    rows = [r + [""] * (max_cols - len(r)) for r in table_slots]
    out: List[str] = ["<table>"]

    for row in rows:
        out.append("  <tr>")
        c = 0
        while c < max_cols:
            tok = row[c]
            if tok in ("__MERGE_COL__", "__MERGE_ROW__", ""):
                c += 1
                continue

            if isinstance(tok, dict):
                txt = (tok.get("text") or "").strip()
                colspan = max(1, int(tok.get("colspan", 1) or 1))
                rowspan = max(1, int(tok.get("rowspan", 1) or 1))
            else:
                txt = str(tok).strip()
                colspan = 1
                rowspan = 1

            attrs: List[str] = []
            if colspan > 1:
                attrs.append(f'colspan="{colspan}"')
            if rowspan > 1:
                attrs.append(f'rowspan="{rowspan}"')

            # Preserve inline adoc image macro references as plain text in markdown HTML.
            safe_text = html.escape(txt)
            attr_txt = (" " + " ".join(attrs)) if attrs else ""
            out.append(f"    <td{attr_txt}>{safe_text}</td>")
            c += max(1, colspan)

        out.append("  </tr>")

    out.append("</table>")
    return "\n".join(out)


def _table_element_to_adoc_block(el: Dict) -> str:
    """Build AsciiDoc table block from element, preferring structured table metadata."""
    table_slots = el.get("table_slots")
    if isinstance(table_slots, list) and table_slots:
        block = _table_slots_to_adoc_block(table_slots)
        if block:
            return block

    table_rows = el.get("table_rows")
    if isinstance(table_rows, list) and table_rows:
        block = _table_rows_to_adoc_block(table_rows)
        if block:
            return block
    return _table_text_to_adoc_block((el.get("content") or ""))


def _table_element_to_markdown_block(el: Dict) -> str:
    """Build Markdown table, with HTML fallback when merge spans are present."""
    table_slots = el.get("table_slots")
    if isinstance(table_slots, list) and table_slots:
        has_span = False
        for row in table_slots:
            for tok in row:
                if isinstance(tok, dict) and (
                    int(tok.get("colspan", 1) or 1) > 1 or int(tok.get("rowspan", 1) or 1) > 1
                ):
                    has_span = True
                    break
            if has_span:
                break
        if has_span:
            html_tbl = _table_slots_to_html(table_slots)
            if html_tbl:
                return html_tbl + "\n"

    raw_content = (el.get("content") or "").strip()
    rows = [r.strip() for r in raw_content.splitlines() if r.strip()]
    if not rows:
        return ""
    first = [c.strip() for c in rows[0].split("|")]
    lines: List[str] = []
    lines.append("| " + " | ".join(first) + " |")
    lines.append("| " + " | ".join(["---"] * len(first)) + " |")
    for r in rows[1:]:
        cells = [c.strip() for c in r.split("|")]
        if len(cells) < len(first):
            cells += [""] * (len(first) - len(cells))
        lines.append("| " + " | ".join(cells[:len(first)]) + " |")
    return "\n".join(lines) + "\n"


def _normalize_ocr_url_text(text: str) -> str:
    """Repair OCR-fractured URL patterns such as 'h t t p s : / / example.com'."""
    if not text:
        return ""

    repaired = text

    # Common OCR spacing artefacts around URL scheme and separators.
    repaired = re.sub(
        r"\b[hH]\s*[tT]\s*[tT]\s*[pP](\s*[sS])?\s*:\s*/\s*/\s*",
        lambda m: "https://" if m.group(1) else "http://",
        repaired,
    )
    repaired = re.sub(r"\b[wW]\s*[wW]\s*[wW]\s*\.\s*", "www.", repaired)
    repaired = re.sub(r"\s*/\s*", "/", repaired)
    repaired = re.sub(r"\s*\.\s*", ".", repaired)

    return repaired


def _extract_links_from_text(text: str) -> List[Tuple[str, str]]:
    """Extract (link_text, uri) candidates from plain/OCR/markdown text."""
    if not text:
        return []

    source = _normalize_ocr_url_text(text)
    hits: List[Tuple[str, str]] = []
    seen = set()

    def _add(label: str, uri: str):
        label = _clean_link_text(label or uri)
        uri = (uri or "").strip().rstrip(".,;:)")
        if not uri:
            return
        key = (label.lower(), uri.lower())
        if key in seen:
            return
        seen.add(key)
        hits.append((label, uri))

    for m in MD_LINK_RE.finditer(source):
        _add(m.group(1), m.group(2))

    for m in URL_RE.finditer(source):
        _add(m.group(1), m.group(1))

    for m in WWW_RE.finditer(source):
        host = m.group(1)
        _add(host, f"https://{host}")

    for m in DOMAIN_RE.finditer(source):
        host_or_path = m.group(1)
        # Skip emails and already-captured URI forms.
        if "@" in host_or_path:
            continue
        if host_or_path.lower().startswith(("http://", "https://", "www.")):
            continue
        _add(host_or_path, f"https://{host_or_path}")

    return hits


def _enrich_pages_with_recovered_links(pages: List[List[Dict]], log) -> Tuple[List[List[Dict]], int]:
    """Append synthetic `link` elements by scanning extracted text for URL patterns."""
    if not isinstance(pages, list):
        return pages, 0

    added = 0
    for pg_idx, page in enumerate(pages):
        if not isinstance(page, list):
            continue

        existing = set()
        for el in page:
            uri = (el.get("uri") or "").strip().lower()
            if uri:
                existing.add(uri)
            if (el.get("native_label") or "") == "link":
                raw = el.get("content", "") or ""
                for _, parsed_uri in _extract_links_from_text(raw):
                    existing.add(parsed_uri.lower())

        new_links: List[Dict] = []
        for el in page:
            raw = el.get("content", "") or ""
            found_links = _extract_links_from_text(raw)
            if found_links:
                # Preserve link signal on original OCR block for downstream diagnostics.
                el["has_link"] = True
                el["detected_links"] = [u for _, u in found_links]

            for label, uri in found_links:
                if uri.lower() in existing:
                    continue
                existing.add(uri.lower())

                meta = {
                    "uri": uri,
                    "link_text": label,
                    "source": "ocr_recovered",
                    "label": "link",
                    "source_label": el.get("native_label", "paragraph"),
                }

                # Carry geometry/confidence from the source OCR region so link tags are visualizable.
                bbox = el.get("bbox_2d")
                if bbox is not None:
                    meta["bbox_2d"] = bbox
                if isinstance(el.get("score"), (int, float)):
                    meta["score"] = float(el.get("score"))
                if isinstance(el.get("confidence"), (int, float)):
                    meta["confidence"] = float(el.get("confidence"))

                new_links.append(
                    _el(
                        f"[{label}]({uri})",
                        "link",
                        pg_idx,
                        meta,
                    )
                )

        if new_links:
            page.extend(new_links)
            added += len(new_links)

    if added:
        log(f"  Recovered {added} link(s) from OCR/plain text patterns.")
    return pages, added


def _extract_pdf_link_elements(page, pg_idx: int) -> List[Dict]:
    """Extract hyperlink annotations from a PDF page as dedicated elements."""
    link_elements: List[Dict] = []
    seen = set()
    try:
        links = page.get_links() or []
    except Exception:
        links = []

    for ln in links:
        uri = (ln or {}).get("uri")
        rect = (ln or {}).get("from")
        if not uri:
            continue

        label = uri
        if rect is not None:
            try:
                words = page.get_text("words", clip=rect) or []
                # words format: (x0, y0, x1, y1, word, block, line, word_no)
                if words:
                    words_sorted = sorted(words, key=lambda w: (w[5], w[6], w[7]))
                    joined = " ".join(str(w[4]) for w in words_sorted).strip()
                    if joined:
                        label = joined
            except Exception:
                pass

        key = (label, uri)
        if key in seen:
            continue
        seen.add(key)
        link_elements.append(
            _el(
                f"[{_clean_link_text(label)}]({uri})",
                "link",
                pg_idx,
                {"uri": uri, "link_text": _clean_link_text(label)},
            )
        )

    return link_elements


# ── ODL element type constants ────────────────────────────────────────────────
_ODL_FIGURE_TYPES = frozenset({"image", "picture", "figure", "photo"})
_ODL_IMG_EXTS     = frozenset({".png", ".jpg", ".jpeg", ".tiff", ".tif",
                                ".bmp", ".webp", ".gif"})


def _odl_bbox_to_list(bbox: Any) -> Optional[List[float]]:
    """Normalise an ODL bounding box to [left, bottom, right, top] (PDF points).
    ODL coordinate origin is bottom-left; values in 1/72-inch points."""
    if bbox is None:
        return None
    if isinstance(bbox, (list, tuple)) and len(bbox) == 4:
        try:
            return [float(v) for v in bbox[:4]]
        except Exception:
            return None
    if isinstance(bbox, dict):
        kmap = {k.lower().replace("_", " "): v for k, v in bbox.items()}
        left   = kmap.get("left",   kmap.get("x0", kmap.get("x")))
        bottom = kmap.get("bottom", kmap.get("y0", kmap.get("y")))
        right  = kmap.get("right",  kmap.get("x1"))
        top    = kmap.get("top",    kmap.get("y1"))
        if right is None and left is not None and kmap.get("width") is not None:
            right = float(left) + float(kmap["width"])
        if top is None and bottom is not None and kmap.get("height") is not None:
            top = float(bottom) + float(kmap["height"])
        try:
            if None not in (left, bottom, right, top):
                return [float(left), float(bottom), float(right), float(top)]
        except Exception:
            return None
    return None


def _odl_flatten_table(node: dict) -> str:
    """Convert an ODL table node into a pipe-delimited text representation."""
    for key in ("rows", "cells", "content"):
        rows_raw = node.get(key)
        if rows_raw is None:
            continue
        if isinstance(rows_raw, str):
            return rows_raw.strip()
        if isinstance(rows_raw, list):
            lines = []
            for row in rows_raw:
                if isinstance(row, str):
                    lines.append(row.strip())
                elif isinstance(row, list):
                    cells = []
                    for cell in row:
                        if isinstance(cell, str):
                            cells.append(cell.strip())
                        elif isinstance(cell, dict):
                            ct = (cell.get("content") or cell.get("text") or "").strip()
                            if ct:
                                cells.append(ct)
                    if cells:
                        lines.append(" | ".join(cells))
                elif isinstance(row, dict):
                    ct = (row.get("content") or row.get("text") or "").strip()
                    if ct:
                        lines.append(ct)
            return "\n".join(lines)
    return ""


def _odl_node_text(node: dict) -> str:
    """Return the best available text content from an ODL JSON node."""
    for k in ("content", "text", "value", "description"):
        v = node.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    typ = str(node.get("type", "")).lower()
    if typ == "table":
        return _odl_flatten_table(node)
    return ""


def _odl_run_convert(path: Path, tdir: str, *,
                     use_hybrid: bool, use_struct: bool) -> Optional[dict]:
    """Call opendataloader_pdf.convert and return the parsed JSON root dict."""
    kwargs: dict = {
        "input_path": str(path),
        "output_dir": tdir,
        "format": "json",
        "image_output": "external",
        "image_format": "png",
    }
    if use_struct:
        kwargs["use_struct_tree"] = True
    if use_hybrid:
        kwargs["hybrid"] = "docling-fast"
        kwargs["hybrid_mode"] = "full"   # enables LaTeX formulas + AI picture descriptions

    # Some older installs may not accept all kwargs; degrade gracefully.
    try:
        _odl_pdf_convert(**kwargs)
    except TypeError:
        safe = {k: v for k, v in kwargs.items()
                if k in ("input_path", "output_dir", "format")}
        _odl_pdf_convert(**safe)

    json_files = sorted(Path(tdir).glob("*.json"))
    if not json_files:
        return None
    try:
        return json.loads(json_files[0].read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return None


def _odl_parse_root(root: Any, img_files: List[Path],
                    ocr_images_dir: Optional[Path]) -> List[List[Dict]]:
    """Walk the ODL JSON tree and emit pages-of-element-dicts."""
    collected: Dict[int, List[Dict]] = {}
    img_queue = list(img_files)
    img_cursor = [0]   # mutable counter shared across nested calls

    def _link_image() -> Optional[str]:
        """Pop the next extracted image, copy to ocr_images_dir, return filename."""
        idx = img_cursor[0]
        if idx >= len(img_queue):
            return None
        img_cursor[0] += 1
        src = img_queue[idx]
        if ocr_images_dir:
            ocr_images_dir.mkdir(parents=True, exist_ok=True)
            dest = ocr_images_dir / src.name
            if not dest.exists():
                try:
                    shutil.copy2(str(src), str(dest))
                except Exception:
                    pass
            return src.name
        return str(src)

    def _walk(node: Any) -> None:
        if isinstance(node, list):
            for item in node:
                _walk(item)
            return
        if not isinstance(node, dict):
            return

        typ = str(node.get("type", "paragraph")).lower()

        page_no = node.get("page number", node.get("page", 1))
        try:
            pg = max(0, int(page_no) - 1)
        except Exception:
            pg = 0

        meta: Dict[str, Any] = {}

        # Bounding box
        bb = _odl_bbox_to_list(
            node.get("bounding box") or node.get("bounding_box") or node.get("bbox"))
        if bb:
            meta["bbox_2d"] = bb

        # Font metadata
        try:
            fs = float(node.get("font size") or node.get("font_size") or 0)
        except Exception:
            fs = 0.0
        if fs > 0:
            meta["font_size"] = fs
        font_name = str(node.get("font") or "").lower()

        # ── HEADING ──────────────────────────────────────────────────────
        if typ in ("heading", "title", "header"):
            text = _odl_node_text(node)
            if text:
                level_raw = node.get("heading level") or node.get("level")
                try:
                    level = int(level_raw)
                except Exception:
                    # "Title" → 1, numeric string → int, fallback → 1
                    level = 1 if str(level_raw or "").lower() in ("title", "1", "") else 2
                meta["heading_level"] = level
                collected.setdefault(pg, []).append(
                    _el(text, "paragraph_title", pg, meta))

        # ── FIGURE / PICTURE / IMAGE ─────────────────────────────────────
        elif typ in _ODL_FIGURE_TYPES:
            img_name = _link_image()
            if img_name:
                meta["image_path"] = img_name
            # Content: prefer description (hybrid AI) then any text, then placeholder
            desc = node.get("description", "")
            text = _odl_node_text(node)
            display = text or (f"[Figure] {desc}".strip() if desc else "[Figure]")
            collected.setdefault(pg, []).append(_el(display, "figure", pg, meta))

        # ── FORMULA ──────────────────────────────────────────────────────
        elif typ == "formula":
            text = _odl_node_text(node)
            if text:
                meta["is_formula"] = True
                collected.setdefault(pg, []).append(_el(text, "para", pg, meta))

        # ── TABLE ────────────────────────────────────────────────────────
        elif typ == "table":
            table_text = _odl_flatten_table(node) or _odl_node_text(node)
            if table_text:
                collected.setdefault(pg, []).append(_el(table_text, "table", pg, meta))
            # Still descend so nested caption/footnote nodes are captured
            _walk_kids(node)
            return

        # ── LIST / LIST ITEM ─────────────────────────────────────────────
        elif typ in ("list", "list_item"):
            text = _odl_node_text(node)
            if text:
                collected.setdefault(pg, []).append(_el(text, "list", pg, meta))
            _walk_kids(node)
            return

        # ── CAPTION ──────────────────────────────────────────────────────
        elif typ == "caption":
            text = _odl_node_text(node)
            if text:
                collected.setdefault(pg, []).append(_el(text, "caption", pg, meta))

        # ── FOOTNOTE / FOOTER ────────────────────────────────────────────
        elif typ in ("footnote", "footer", "footnote_text"):
            text = _odl_node_text(node)
            if text:
                collected.setdefault(pg, []).append(_el(text, "para", pg, meta))

        # ── CODE BLOCK ───────────────────────────────────────────────────
        elif typ in ("code", "code_block", "pre"):
            text = _odl_node_text(node)
            if text:
                collected.setdefault(pg, []).append(_el(text, "para", pg, meta))

        # ── DEFAULT: paragraph + heuristic heading detection ─────────────
        else:
            text = _odl_node_text(node)
            if text:
                is_heading = (
                    fs >= 14
                    or ("bold" in font_name and fs >= 11)
                    or typ in ("subheading", "subtitle")
                )
                native = "paragraph_title" if is_heading else "paragraph"
                if is_heading:
                    meta["heading_level"] = 2
                collected.setdefault(pg, []).append(_el(text, native, pg, meta))

        _walk_kids(node)

    def _walk_kids(node: dict) -> None:
        for key in ("kids", "children", "items", "elements", "content_list"):
            v = node.get(key)
            if isinstance(v, list):
                for child in v:
                    _walk(child)

    _walk(root)

    if not collected:
        return []
    return [collected[pg] for pg in sorted(collected.keys())]


def extract_pdf_opendataloader(
    path: Path,
    log,
    ocr_images_dir: Optional[Path] = None,
) -> List[List[Dict]]:
    """High-accuracy PDF extraction via opendataloader-pdf.

    Extraction strategy (highest → lowest accuracy):
      1. Hybrid mode (docling-fast + full) — complex tables, formulas, AI figure captions.
         Requires `opendataloader-pdf-hybrid` server running on ODL_HYBRID_URL.
      2. Local mode + use_struct_tree=True — uses native PDF accessibility tags.
      3. Local mode (plain) — deterministic XY-Cut++ layout, fastest, no GPU.

    All extracted images are copied to *ocr_images_dir* and linked into element dicts.
    """
    if not HAS_OPENDATALOADER_PDF:
        return []

    try:
        with tempfile.TemporaryDirectory(prefix="odl_pdf_") as tdir:
            root: Optional[dict] = None

            # ── Pass 1: hybrid mode (needs server) ───────────────────────
            if ODL_USE_HYBRID:
                log("  opendataloader-pdf [hybrid + struct_tree]…")
                try:
                    root = _odl_run_convert(
                        path, tdir, use_hybrid=True, use_struct=True)
                except Exception as exc:
                    log(f"    hybrid mode failed ({exc}); falling back to local.")
                    root = None

            # ── Pass 2: local + use_struct_tree ──────────────────────────
            if root is None:
                log("  opendataloader-pdf [local, use_struct_tree=True]…")
                try:
                    root = _odl_run_convert(
                        path, tdir, use_hybrid=False, use_struct=True)
                except Exception as exc:
                    log(f"    struct_tree mode failed ({exc}); trying plain.")
                    root = None

            # ── Pass 3: plain local ───────────────────────────────────────
            if root is None:
                log("  opendataloader-pdf [local, plain]…")
                try:
                    root = _odl_run_convert(
                        path, tdir, use_hybrid=False, use_struct=False)
                except Exception as exc:
                    log(f"    opendataloader-pdf failed: {exc}")
                    return []

            if root is None:
                log("  opendataloader-pdf: no JSON output produced.")
                return []

            # Collect all extracted image files (image_output="external")
            img_files: List[Path] = sorted(
                p for p in Path(tdir).rglob("*")
                if p.suffix.lower() in _ODL_IMG_EXTS
            )

            pages = _odl_parse_root(root, img_files, ocr_images_dir)

        el_count = sum(len(p) for p in pages)
        fig_count = sum(
            1 for p in pages for e in p
            if e.get("native_label") == "figure" and e.get("image_path")
        )
        log(f"  opendataloader-pdf: {len(pages)} page(s), "
            f"{el_count} element(s), {fig_count} figure(s) extracted.")
        return pages

    except Exception as exc:
        log(f"  opendataloader-pdf extraction failed: {exc}")
        traceback.print_exc()
        return []


# ────────────────────────────────────────────────────────────────────────────
# EXTRACTORS
# ────────────────────────────────────────────────────────────────────────────

def extract_pdf_text(path: Path, log) -> List[List[Dict]]:
    """Extract text-layer PDF with pymupdf. Returns list of pages."""
    if not HAS_PYMUPDF:
        log("  WARNING: pymupdf not installed – cannot extract text PDF.")
        return []
    doc = _fitz.open(str(path))
    pages = []
    for pg_idx, page in enumerate(doc):
        elements = []
        blocks = page.get_text("dict")["blocks"]
        for blk in blocks:
            if blk.get("type") != 0:   # 0 = text
                continue
            for line in blk.get("lines", []):
                text = " ".join(sp["text"] for sp in line.get("spans", []))
                text = text.strip()
                if not text:
                    continue
                # Heuristic label: large font / bold first span → title
                spans = line.get("spans", [])
                avg_size = sum(s["size"] for s in spans) / max(len(spans), 1)
                is_bold = any("bold" in s.get("font", "").lower() for s in spans)
                if avg_size >= 14 or (is_bold and avg_size >= 11):
                    label = "paragraph_title"
                else:
                    label = "paragraph"
                elements.append(_el(text, label, pg_idx))
        # Preserve interactive PDF hyperlinks for downstream outputs.
        elements.extend(_extract_pdf_link_elements(page, pg_idx))
        pages.append(elements)
    doc.close()
    return pages


def extract_pdf_ocr(
    path: Path,
    log,
    glm_native_out_dir: Optional[Path] = None,
    ocr_images_dir: Optional[Path] = None,
) -> List[List[Dict]]:
    """Use glmocr to OCR a scanned/image PDF. Returns list of pages."""
    if not HAS_GLMOCR:
        log("  WARNING: glmocr not installed – trying opendataloader-pdf, then pymupdf.")
        pages = (extract_pdf_opendataloader(path, log, ocr_images_dir)
                 if HAS_OPENDATALOADER_PDF else [])
        if pages:
            return pages
        return extract_pdf_text(path, log)
    log("  Loading glmocr pipeline (this may take a moment)...")
    try:
        # Render each PDF page to an image and OCR page-by-page so that
        # progress is visible and a single slow page doesn't block everything.
        if HAS_PYMUPDF:
            # Render all pages up-front, then feed them as ONE batch into the
            # glmocr pipeline with stream=True. The pipeline runs load → layout
            # → recognition as concurrent stages and fans region OCR requests
            # out over pipeline.max_workers threads (set from OCR_WORKERS in
            # _create_glmocr_parser), while the stream yields per-page results
            # in order as they finish — giving both parallelism and progress.
            doc = _fitz.open(str(path))
            n_pages = len(doc)
            log(f"  GLM OCR: rendering {n_pages} page(s)...")
            page_imgs: List[bytes] = []
            for pg_idx in range(n_pages):
                pix = doc[pg_idx].get_pixmap(matrix=_fitz.Matrix(2.0, 2.0), alpha=False)
                page_imgs.append(_to_rgb_bytes(pix.tobytes("png")))
            doc.close()

            log(f"  GLM OCR: {n_pages} page(s), {OCR_WORKERS} concurrent OCR request(s)...")

            all_pages: List[List[Dict]] = []
            mapping: Dict[str, str] = {}
            t_start = time.time()
            last_activity = [time.time()]
            hb_stop = threading.Event()

            with _create_glmocr_parser(log) as parser:
                # Heartbeat: pages yield in order, so the first page can take
                # minutes with nothing printed. Log queue activity every 30s
                # of silence so the user can see the pipeline is alive.
                def _heartbeat():
                    while not hb_stop.wait(30):
                        if time.time() - last_activity[0] < 30:
                            continue
                        msg = f"  GLM OCR working... ({time.time() - t_start:.0f}s elapsed"
                        try:
                            pl = getattr(parser, "_pipeline", None)
                            stats = pl.get_queue_stats() if pl else None
                            if stats:
                                msg += (f", {stats['page_queue_size']} page(s) queued, "
                                        f"{stats['region_queue_size']} region(s) awaiting OCR")
                        except Exception:
                            pass
                        log(msg + ")")

                threading.Thread(target=_heartbeat, daemon=True).start()
                try:
                    for pg_idx, results in enumerate(parser.parse(page_imgs, stream=True)):
                        try:
                            raw_obj: Any = None
                            if hasattr(results, "to_json"):
                                try:
                                    raw_obj = results.to_json()
                                except Exception:
                                    raw_obj = None
                            if raw_obj is None and hasattr(results, "pages"):
                                raw_obj = results.pages
                            if raw_obj is None and hasattr(results, "json_result"):
                                raw_obj = getattr(results, "json_result")

                            pg_pages = _normalize_glm_pages(raw_obj, log, f"{path.name}:p{pg_idx+1}")
                            pg_map = _collect_glm_figure_images(
                                results, glm_native_out_dir, ocr_images_dir,
                                f"{path.stem}_p{pg_idx+1}_fig_", log
                            )
                            mapping.update(pg_map)
                            all_pages.extend(pg_pages if pg_pages else [[]])
                        except Exception as pg_exc:
                            log(f"    page {pg_idx + 1} failed: {pg_exc}; skipping.")
                            all_pages.append([])

                        done = pg_idx + 1
                        last_activity[0] = time.time()
                        elapsed = time.time() - t_start
                        rate = done / elapsed if elapsed > 0 else 0
                        eta = (n_pages - done) / rate if rate > 0 else 0
                        log(f"  GLM OCR page {done}/{n_pages} done "
                            f"({elapsed:.0f}s elapsed, ~{eta:.0f}s left)")
                finally:
                    hb_stop.set()

            if any(pg for pg in all_pages):
                _apply_image_path_remap(all_pages, mapping)
                return all_pages
            log("  GLM OCR returned no parseable content; falling back to pymupdf.")
            return extract_pdf_text(path, log)

        # Fallback: pass entire PDF at once (no per-page progress)
        with _create_glmocr_parser(log) as parser:
            with open(str(path), "rb") as f:
                pdf_bytes = f.read()
            log(f"  GLM OCR: processing entire PDF (no page-by-page progress)...")
            results = parser.parse(pdf_bytes)

        _preserve_glm_default_output(results, glm_native_out_dir, log, path.name)
        mapping = _collect_glm_figure_images(
            results, glm_native_out_dir, ocr_images_dir, f"{path.stem}_fig_", log
        )

        raw_obj = None
        if hasattr(results, "to_json"):
            try:
                raw_obj = results.to_json()
            except Exception:
                raw_obj = None
        if raw_obj is None and hasattr(results, "pages"):
            raw_obj = results.pages
        if raw_obj is None and hasattr(results, "json_result"):
            raw_obj = getattr(results, "json_result")

        pages = _normalize_glm_pages(raw_obj, log, path.name)
        if pages:
            _apply_image_path_remap(pages, mapping)
            return pages

        log("  GLM OCR returned no parseable pages; falling back to pymupdf text extraction.")
        return extract_pdf_text(path, log)
    except Exception as e:
        log(f"  glmocr failed: {e}. Falling back to pymupdf.")
        return extract_pdf_text(path, log)


def _create_glmocr_parser(log):
    """Build a GLM OCR parser using either default settings or Ollama backend."""
    # Concurrency: the pipeline's recognition stage runs OCR requests through
    # an internal thread pool sized by pipeline.max_workers (config.yaml ships
    # with 1, which serializes everything). Raise it so multiple text regions
    # are OCR'd in flight at once. The connection pool must be >= max_workers.
    conc = {
        "pipeline.max_workers": OCR_WORKERS,
        "pipeline.ocr_api.connection_pool_size": max(8, OCR_WORKERS * 2),
    }
    if GLMOCR_BACKEND == "ollama":
        log(f"  GLM OCR backend: ollama ({GLMOCR_OLLAMA_MODEL}), {OCR_WORKERS} worker(s)")
        dotted = {
            "pipeline.maas.enabled": False,
            "pipeline.ocr_api.api_mode": "ollama_generate",
            "pipeline.ocr_api.api_url": GLMOCR_OLLAMA_URL,
            "pipeline.ocr_api.model": GLMOCR_OLLAMA_MODEL,
            **conc,
        }
        layout_dir = GLMOCR_LAYOUT_MODEL_DIR or _autodetect_layout_model_dir()
        if layout_dir:
            dotted["pipeline.layout.model_dir"] = layout_dir
            log(f"  GLM OCR layout model dir: {layout_dir}")
        else:
            log("  GLM OCR layout model dir not set; first run may access Hugging Face cache/downloads.")
        return _GlmOcr(
            mode="selfhosted",
            model=GLMOCR_OLLAMA_MODEL,
            _dotted=dotted,
        )
    return _GlmOcr(_dotted=conc)


def _is_scanned_pdf(path: Path) -> bool:
    """Heuristic: if first 3 pages have < 50 chars of text, treat as scanned."""
    if not HAS_PYMUPDF:
        return False
    try:
        doc = _fitz.open(str(path))
        chars = sum(len(doc[i].get_text()) for i in range(min(3, len(doc))))
        doc.close()
        return chars < 150
    except Exception:
        return False


def _is_image_heavy_docx(path: Path) -> bool:
    """Heuristic: DOCX with embedded images and < 200 chars of paragraph text → treat as image-based."""
    if not HAS_DOCX:
        return False
    try:
        document = _docx.Document(str(path))
        text_chars = sum(len(p.text) for p in document.paragraphs)
        image_count = sum(
            1 for rel in document.part.rels.values()
            if "image" in rel.reltype.lower()
        )
        return image_count > 0 and text_chars < 200
    except Exception:
        return False


def _docx_image_ext(content_type: str) -> str:
    """Map a DOCX image content-type string to a file extension."""
    ct = (content_type or "").lower()
    if "png" in ct:
        return ".png"
    if "jpeg" in ct or "jpg" in ct:
        return ".jpg"
    if "gif" in ct:
        return ".gif"
    if "bmp" in ct:
        return ".bmp"
    if "tiff" in ct:
        return ".tiff"
    if "webp" in ct:
        return ".webp"
    return ".png"


def extract_docx_ocr(
    path: Path,
    log,
    glm_native_out_dir: Optional[Path] = None,
    ocr_images_dir: Optional[Path] = None,
) -> List[List[Dict]]:
    """OCR all embedded images in a DOCX with glmocr, producing the same output as extract_pdf_ocr."""
    if not HAS_DOCX:
        log("  WARNING: python-docx not installed – cannot open DOCX for image OCR.")
        return []

    try:
        document = _docx.Document(str(path))
    except Exception as e:
        log(f"  ERROR opening DOCX: {e}")
        return []

    # Collect (bytes, ext) for every unique image relationship in the document.
    images: List[Tuple[bytes, str]] = []
    seen_rids: set = set()
    for rid, rel in document.part.rels.items():
        if "image" not in rel.reltype.lower():
            continue
        if rid in seen_rids:
            continue
        seen_rids.add(rid)
        try:
            part = rel.target_part
            blob = part.blob
            if not blob:
                continue
            ext = _docx_image_ext(getattr(part, "content_type", ""))
            images.append((blob, ext))
        except Exception:
            continue

    if not images:
        log("  No embedded images found in DOCX; falling back to text extraction.")
        return extract_docx(path, log)

    if not HAS_GLMOCR:
        log("  WARNING: glmocr not installed – falling back to DOCX text extraction.")
        return extract_docx(path, log)

    # Save source images so they are accessible in the output images folder.
    src_img_names: List[str] = []
    if ocr_images_dir is not None:
        ocr_images_dir.mkdir(parents=True, exist_ok=True)
        for idx, (blob, ext) in enumerate(images):
            src_name = f"{path.stem}_img_{idx + 1:03d}{ext}"
            try:
                (ocr_images_dir / src_name).write_bytes(blob)
                src_img_names.append(src_name)
            except Exception as e:
                log(f"  Warning: could not save source image {src_name}: {e}")
                src_img_names.append("")
    else:
        src_img_names = [""] * len(images)

    log(f"  Found {len(images)} embedded image(s) in DOCX – running GLM OCR on each...")
    all_pages: List[List[Dict]] = []

    try:
        with _create_glmocr_parser(log) as parser:
            for img_idx, (img_bytes, img_ext) in enumerate(images):
                log(f"  OCR image {img_idx + 1}/{len(images)}...")
                try:
                    results = parser.parse(_to_rgb_bytes(img_bytes))

                    # Persist full GLM artifacts for the first image only (avoid dir clutter).
                    per_img_glm_dir = None
                    if glm_native_out_dir is not None:
                        per_img_glm_dir = glm_native_out_dir / f"img{img_idx + 1:03d}"
                        _preserve_glm_default_output(results, per_img_glm_dir, log, f"{path.name}[img{img_idx+1}]")

                    # Collect figure crops into ocr_images_dir.
                    mapping = _collect_glm_figure_images(
                        results,
                        per_img_glm_dir,
                        ocr_images_dir,
                        f"{path.stem}_fig{img_idx + 1:03d}_",
                        log,
                    )

                    raw_obj: Any = None
                    if hasattr(results, "to_json"):
                        try:
                            raw_obj = results.to_json()
                        except Exception:
                            raw_obj = None
                    if raw_obj is None and hasattr(results, "pages"):
                        raw_obj = results.pages
                    if raw_obj is None and hasattr(results, "json_result"):
                        raw_obj = getattr(results, "json_result")

                    img_pages = _normalize_glm_pages(
                        raw_obj, log, f"{path.name}[img{img_idx + 1}]"
                    )
                    if img_pages:
                        _apply_image_path_remap(img_pages, mapping)
                        page_offset = len(all_pages)
                        src_name = src_img_names[img_idx] if img_idx < len(src_img_names) else ""
                        for pg_idx, pg in enumerate(img_pages):
                            for el in pg:
                                el["page_no"] = page_offset + pg_idx
                                # If GLM left no image_path on a figure element, point to the source image.
                                if el.get("native_label") == "figure" and not el.get("image_path") and src_name:
                                    el["image_path"] = src_name
                        # PIL-based crop: extract sub-regions within this source image.
                        _crop_and_save_figure_regions(
                            img_bytes,
                            img_pages,
                            ocr_images_dir,
                            f"{path.stem}_img{img_idx + 1:03d}_fig_",
                            log,
                        )
                        all_pages.extend(img_pages)
                except Exception as e:
                    log(f"  GLM OCR failed for image {img_idx + 1}: {e}")
    except Exception as e:
        log(f"  glmocr pipeline error during DOCX image OCR: {e}")

    if all_pages:
        log(f"  DOCX image OCR complete: {len(all_pages)} page(s) from {len(images)} image(s).")
        return all_pages

    log("  GLM OCR returned no results for DOCX images; falling back to text extraction.")
    return extract_docx(path, log)


def extract_docx(path: Path, log, docx_media_dir: Optional[Path] = None) -> List[List[Dict]]:
    """Extract Word DOCX using python-docx → element list."""
    if not HAS_DOCX:
        log("  WARNING: python-docx not installed. Trying pandoc text fallback.")
        return _extract_via_pandoc_plain(path, log)

    document = _docx.Document(str(path))
    from docx.text.paragraph import Paragraph as _DocxParagraph
    from docx.table import Table as _DocxTable

    elements: List[Dict] = []
    pg = 0
    HEADINGS = {"Heading 1", "Heading 2", "Heading 3", "Heading 4",
                "Heading 5", "Heading 6", "Title", "Subtitle"}
    LIST_STYLES = {"List Paragraph", "List Bullet", "List Number",
                   "List Bullet 2", "List Number 2"}

    # Optional media export for embedded DOCX images (used by ADOC/XML figure refs).
    media_dir = docx_media_dir
    if media_dir is not None:
        media_dir.mkdir(parents=True, exist_ok=True)

    rid_to_name: Dict[str, str] = {}
    img_count = 0
    _docx_img_ocr_cache: Dict[str, str] = {}

    def _content_type_ext(content_type: str) -> str:
        ct = (content_type or "").lower()
        if "png" in ct:
            return ".png"
        if "jpeg" in ct or "jpg" in ct:
            return ".jpg"
        if "gif" in ct:
            return ".gif"
        if "bmp" in ct:
            return ".bmp"
        if "tiff" in ct:
            return ".tiff"
        if "webp" in ct:
            return ".webp"
        return ".img"

    def _extract_images_from_node(node) -> List[str]:
        """Extract embedded images from any DOCX XML-backed node and return saved names."""
        nonlocal img_count
        names: List[str] = []
        try:
            # Find all DrawingML blips in this paragraph (supports inline/anchored images).
            base_el = getattr(node, "_element", node)
            blips = base_el.xpath('.//*[local-name()="blip"]')
        except Exception:
            blips = []

        rids: List[str] = []
        for blip in blips:
            for attr_k, attr_v in getattr(blip, "attrib", {}).items():
                if str(attr_k).endswith("}embed") and attr_v:
                    rids.append(attr_v)

        # Support legacy VML icons/images (<v:imagedata r:id="..."></v:imagedata>).
        try:
            vml_ids = base_el.xpath('.//*[local-name()="imagedata"]/@*[local-name()="id"]')
            for rid in vml_ids:
                if rid:
                    rids.append(rid)
        except Exception:
            pass

        for rid in rids:
            if not rid:
                continue

            if rid in rid_to_name:
                names.append(rid_to_name[rid])
                continue

            try:
                part = document.part.related_parts[rid]
                blob = part.blob
                ext = _content_type_ext(getattr(part, "content_type", ""))
            except Exception:
                continue

            img_count += 1
            img_name = f"{path.stem}_img_{img_count:03d}{ext}"
            if media_dir is not None:
                try:
                    (media_dir / img_name).write_bytes(blob)
                except Exception:
                    continue
            rid_to_name[rid] = img_name
            names.append(img_name)
        return names

    def _ocr_docx_image_text(img_name: str) -> str:
        """Best-effort OCR text recovery for embedded DOCX images."""
        if not img_name:
            return ""
        if img_name in _docx_img_ocr_cache:
            return _docx_img_ocr_cache[img_name]
        if media_dir is None or not HAS_GLMOCR:
            _docx_img_ocr_cache[img_name] = ""
            return ""

        try:
            img_path = media_dir / img_name
            if not img_path.exists():
                _docx_img_ocr_cache[img_name] = ""
                return ""

            # Reuse image extractor pipeline to decode OCR blocks from image-only callouts.
            pages = extract_image(img_path, lambda *_a, **_k: None, None)
            chunks: List[str] = []
            for page in pages:
                for el in page:
                    txt = _clean_text(el.get("content", "") or "")
                    if not txt:
                        continue
                    # Prefer cautionary callout text and concise content.
                    if re.search(r"\b(warning|caution|notice|danger)\b", txt, re.IGNORECASE):
                        chunks.append(txt)
                    elif len(txt) <= 120:
                        chunks.append(txt)

            # Keep order while removing duplicates.
            seen = set()
            ordered = []
            for c in chunks:
                key = c.lower()
                if key in seen:
                    continue
                seen.add(key)
                ordered.append(c)

            out = " | ".join(ordered[:6]).strip()
            _docx_img_ocr_cache[img_name] = out
            return out
        except Exception:
            _docx_img_ocr_cache[img_name] = ""
            return ""

    table_count = 0

    def _extract_textbox_text_from_node(node) -> List[str]:
        """Extract text from Word textboxes/shapes attached to a node."""
        try:
            base_el = getattr(node, "_element", node)
        except Exception:
            return []

        lines: List[str] = []
        seen = set()
        try:
            # Prefer textbox paragraph-level extraction to preserve logical lines.
            p_nodes = base_el.xpath('.//*[local-name()="txbxContent"]//*[local-name()="p"]')
            for p in p_nodes:
                texts = p.xpath('.//*[local-name()="t"]/text()')
                if not texts:
                    continue
                line = " ".join(t.strip() for t in texts if t and t.strip())
                line = " ".join(line.split())
                if line and line not in seen:
                    seen.add(line)
                    lines.append(line)
        except Exception:
            pass

        return lines

    def _emit_paragraph(para):
        for img_name in _extract_images_from_node(para):
            elements.append(_el(f"[Image: {img_name}]", "figure", pg,
                                {"image_path": img_name}))

        text = para.text.strip()
        textbox_lines = _extract_textbox_text_from_node(para)

        # Normal paragraph text plus any textbox/shape text anchored to this paragraph.
        merged_lines: List[str] = []
        if text:
            merged_lines.append(text)
        for line in textbox_lines:
            if line and line not in merged_lines:
                merged_lines.append(line)

        if not merged_lines:
            return

        text = "\n".join(merged_lines)

        style = para.style.name if para.style else ""
        if style in HEADINGS or "Heading" in style:
            depth = 1
            m = re.match(r"Heading (\d)", style)
            if m:
                depth = int(m.group(1))
            label = "doc_title" if depth == 1 else "paragraph_title"
            prefix = "#" * depth + " "
            elements.append(_el(prefix + text, label, pg))
        elif style in LIST_STYLES:
            bullet = ".. " if "2" in style else ". "
            elements.append(_el(bullet + text, "paragraph", pg))
        else:
            elements.append(_el(text, "paragraph", pg))

    def _emit_table(tbl):
        nonlocal table_count
        rows: List[List[str]] = []
        rows_meta: List[List[str]] = []
        rows_slots: List[List[Any]] = []

        def _tc_text(tc) -> str:
            try:
                texts = tc.xpath('.//*[local-name()="t"]/text()')
                sym_chars: List[str] = []
                # Preserve symbol icons used by Word (e.g., Wingdings/Symbol glyph refs).
                for sym in tc.xpath('.//*[local-name()="sym"]'):
                    ch = None
                    try:
                        for attr_k, attr_v in getattr(sym, "attrib", {}).items():
                            if str(attr_k).endswith("}char") and attr_v:
                                ch = str(attr_v)
                                break
                    except Exception:
                        ch = None
                    if ch:
                        ch = ch.strip()
                        try:
                            if ch.lower().startswith("0x"):
                                ch = ch[2:]
                            sym_chars.append(chr(int(ch, 16)))
                        except Exception:
                            # Keep a visible placeholder if conversion fails.
                            sym_chars.append("[icon]")

                img_names = _extract_images_from_node(tc)
                # In table cells, emit inline image macros.
                img_tokens = [f"image:{nm}[]" for nm in img_names]
                ocr_chunks = [_ocr_docx_image_text(nm) for nm in img_names]
                ocr_text = " | ".join(x for x in ocr_chunks if x)
                parts = [" ".join(texts), " ".join(sym_chars), " ".join(img_tokens), ocr_text]
                return " ".join(" ".join(p for p in parts if p).split())
            except Exception:
                return ""

        # Build logical rows from XML to preserve merged-cell geometry.
        tr_list = getattr(tbl._tbl, "tr_lst", None)
        if tr_list is None:
            # Fallback to python-docx high-level rows if XML access is unavailable.
            for row in tbl.rows:
                cells = [" ".join(c.text.split()) for c in row.cells]
                rows.append(cells)
        else:
            for tr in tr_list:
                logical_row: List[str] = []
                logical_row_meta: List[str] = []
                logical_row_slots: List[Any] = []
                tc_list = getattr(tr, "tc_lst", [])
                for tc in tc_list:
                    txt = _tc_text(tc)
                    # Horizontal merge width.
                    gs = 1
                    try:
                        gs_attr = tc.xpath('./*[local-name()="tcPr"]/*[local-name()="gridSpan"]/@*[local-name()="val"]')
                        if gs_attr:
                            gs = max(1, int(gs_attr[0]))
                    except Exception:
                        gs = 1

                    # Vertical merge continuation cells must be blank to avoid duplication.
                    is_vmerge_continue = False
                    vmerge_state = "none"
                    try:
                        vm_nodes = tc.xpath('./*[local-name()="tcPr"]/*[local-name()="vMerge"]')
                        if vm_nodes:
                            vm_val = vm_nodes[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                            if vm_val is None or str(vm_val).lower() == "continue":
                                is_vmerge_continue = True
                                vmerge_state = "continue"
                            elif str(vm_val).lower() == "restart":
                                vmerge_state = "restart"
                    except Exception:
                        is_vmerge_continue = False

                    logical_row.append("" if is_vmerge_continue else txt)
                    logical_row_meta.append("__MERGE_ROW__" if is_vmerge_continue else txt)
                    if is_vmerge_continue:
                        logical_row_slots.append("__MERGE_ROW__")
                    else:
                        logical_row_slots.append({
                            "text": txt,
                            "colspan": gs,
                            "rowspan": 1,
                            "vmerge": vmerge_state,
                        })
                    for _ in range(gs - 1):
                        logical_row.append("")
                        logical_row_meta.append("__MERGE_COL__")
                        logical_row_slots.append("__MERGE_COL__")

                rows.append(logical_row)
                rows_meta.append(logical_row_meta)
                rows_slots.append(logical_row_slots)

        max_cols = max((len(r) for r in rows), default=0)
        if max_cols > 0:
            rows = [r + [""] * (max_cols - len(r)) for r in rows]
            rows_meta = [r + [""] * (max_cols - len(r)) for r in rows_meta]
            rows_slots = [r + [""] * (max_cols - len(r)) for r in rows_slots]

        # Resolve row spans: restart cell consumes following MERGE_ROW placeholders.
        for rr, row in enumerate(rows_slots):
            for cc, tok in enumerate(row):
                if not isinstance(tok, dict):
                    continue
                if tok.get("vmerge") != "restart":
                    continue
                span = 1
                r2 = rr + 1
                while r2 < len(rows_slots):
                    nxt = rows_slots[r2][cc] if cc < len(rows_slots[r2]) else ""
                    if nxt == "__MERGE_ROW__":
                        span += 1
                        r2 += 1
                        continue
                    break
                tok["rowspan"] = span

        row_lines = [" | ".join(c.strip() for c in row) for row in rows]
        table_text = "\n".join(row_lines)
        if table_text.strip():
            meta = {
                "table_rows": rows_meta,
                "table_slots": rows_slots,
            } if (rows_meta or rows_slots) else None
            elements.append(_el(table_text, "table", pg, meta))
            table_count += 1

    # Iterate DOCX body blocks in source order so tables are not dropped/reordered.
    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1].lower()
        if tag == "p":
            _emit_paragraph(_DocxParagraph(child, document))
        elif tag == "tbl":
            _emit_table(_DocxTable(child, document))

    if table_count or img_count:
        log(f"  DOCX extracted: {table_count} table(s), {img_count} image(s).")

    return _chunk_pages(elements)


def extract_txt(path: Path, log) -> List[List[Dict]]:
    """Parse plain-text or Markdown file into elements."""
    text = path.read_text(encoding="utf-8", errors="replace")
    text = text.lstrip("\ufeff")
    lines = text.splitlines()
    elements: List[Dict] = []
    pg = 0
    para_buf: List[str] = []
    table_buf: List[str] = []
    in_html_table = False

    def flush_para():
        body = "\n".join(para_buf).strip()
        if body:
            elements.append(_el(body, "paragraph", pg))
        para_buf.clear()

    for line in lines:
        stripped = line.strip()

        # HTML table blocks in Markdown: preserve as native table elements.
        if in_html_table:
            table_buf.append(line)
            if re.search(r"</table>", stripped, re.IGNORECASE):
                flush_para()
                tbl = "\n".join(table_buf).strip()
                if tbl:
                    elements.append(_el(tbl, "table", pg))
                table_buf.clear()
                in_html_table = False
            continue

        if re.search(r"<table\b", stripped, re.IGNORECASE):
            flush_para()
            table_buf = [line]
            if re.search(r"</table>", stripped, re.IGNORECASE):
                tbl = "\n".join(table_buf).strip()
                if tbl:
                    elements.append(_el(tbl, "table", pg))
                table_buf.clear()
                in_html_table = False
            else:
                in_html_table = True
            continue

        # Markdown heading
        m = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if m:
            flush_para()
            depth = len(m.group(1))
            lbl = "doc_title" if depth == 1 else "paragraph_title"
            elements.append(_el(stripped, lbl, pg))
        elif stripped == "":
            flush_para()
        else:
            para_buf.append(line)

    if in_html_table and table_buf:
        flush_para()
        tbl = "\n".join(table_buf).strip()
        if tbl:
            elements.append(_el(tbl, "table", pg))

    flush_para()
    return _chunk_pages(elements)


def extract_image(
    path: Path,
    log,
    glm_native_out_dir: Optional[Path] = None,
    ocr_images_dir: Optional[Path] = None,
) -> List[List[Dict]]:
    """OCR a single raster image with glmocr."""
    if not HAS_GLMOCR:
        log("  WARNING: glmocr not installed – cannot OCR image.")
        return [[_el(f"[Image: {path.name}]", "figure", 0,
                     {"image_path": str(path)})]]
    log("  Loading glmocr pipeline for image OCR...")
    try:
        with _create_glmocr_parser(log) as parser:
            img_bytes = path.read_bytes()
            results = parser.parse(_to_rgb_bytes(img_bytes))

        _preserve_glm_default_output(results, glm_native_out_dir, log, path.name)
        mapping = _collect_glm_figure_images(
            results, glm_native_out_dir, ocr_images_dir, f"{path.stem}_fig_", log
        )

        raw_obj: Any = None
        if hasattr(results, "to_json"):
            try:
                raw_obj = results.to_json()
            except Exception:
                raw_obj = None
        if raw_obj is None and hasattr(results, "pages"):
            raw_obj = results.pages

        pages = _normalize_glm_pages(raw_obj, log, path.name)
        if pages:
            _apply_image_path_remap(pages, mapping)
            # PIL-based fallback: crop figure regions directly from the source image
            # for any figure element that still has no image_path after the GLM save scan.
            _crop_and_save_figure_regions(
                img_bytes, pages, ocr_images_dir, f"{path.stem}_fig_", log
            )
            return pages

        log("  GLM OCR image parse returned no parseable pages.")
    except Exception as e:
        log(f"  glmocr image error: {e}")
    return [[_el(f"[Image: {path.name}]", "figure", 0)]]


def _normalize_glm_pages(raw_obj: Any, log, source_name: str) -> List[List[Dict]]:
    """Normalize GLM OCR outputs (string/list/dict) to pages of element dicts."""
    if raw_obj is None:
        return []

    obj = raw_obj
    if isinstance(obj, str):
        try:
            obj = json.loads(obj)
        except Exception:
            log(f"  GLM parse warning: unable to decode JSON payload for {source_name}.")
            return []

    def _to_page_elements(page_obj: Any, page_idx: int) -> List[Dict]:
        elements: List[Dict] = []
        if isinstance(page_obj, dict):
            # Shape A: {layout: {blocks: [...]}}
            if isinstance(page_obj.get("layout"), dict):
                blocks = page_obj.get("layout", {}).get("blocks", [])
            # Shape B: {blocks: [...]} or {regions: [...]} or {layout_details:[...]}
            elif isinstance(page_obj.get("blocks"), list):
                blocks = page_obj.get("blocks", [])
            elif isinstance(page_obj.get("regions"), list):
                blocks = page_obj.get("regions", [])
            elif isinstance(page_obj.get("layout_details"), list):
                blocks = page_obj.get("layout_details", [])
            else:
                blocks = []
        elif isinstance(page_obj, list):
            blocks = page_obj
        else:
            blocks = []

        for bi, b in enumerate(blocks):
            if not isinstance(b, dict):
                continue
            content = (
                b.get("content")
                or b.get("block_content")
                or b.get("text")
                or ""
            )
            _raw_label = (
                b.get("native_label")
                or b.get("layout_type")
                or b.get("label")
                or "paragraph"
            )
            # Normalize GLM OCR image-block label variants to the canonical "figure"
            label = "figure" if _raw_label.lower() in {"image", "figure", "picture", "photo"} else _raw_label
            bbox = b.get("bbox_2d") or b.get("layout_box") or b.get("bbox")
            image_path = b.get("image_path")

            meta = {
                "index": b.get("index", bi),
                "label": b.get("label", label),
            }
            if bbox is not None:
                meta["bbox_2d"] = bbox
            if image_path:
                meta["image_path"] = image_path
            score = b.get("score")
            confidence = b.get("confidence")
            if isinstance(score, (int, float)):
                meta["score"] = float(score)
            if isinstance(confidence, (int, float)):
                meta["confidence"] = float(confidence)

            elements.append(_el(str(content), str(label), page_idx, meta))
        return elements

    # Shape 1: [[...], [...]]
    if isinstance(obj, list):
        if obj and isinstance(obj[0], list):
            pages = [_to_page_elements(p, i) for i, p in enumerate(obj)]
            return [p for p in pages if p]
        # Shape 2: [{...page...}, {...page...}] or flat blocks list
        if obj and isinstance(obj[0], dict):
            # If entries look like blocks, treat as single page
            block_like = sum(1 for x in obj if isinstance(x, dict) and ("content" in x or "block_content" in x))
            if block_like >= max(1, len(obj) // 2):
                p0 = _to_page_elements(obj, 0)
                return [p0] if p0 else []
            pages = [_to_page_elements(p, i) for i, p in enumerate(obj)]
            return [p for p in pages if p]
        return []

    # Shape 3: dict wrapper with pages/json_result/layout_details
    if isinstance(obj, dict):
        if isinstance(obj.get("pages"), list):
            pages = [_to_page_elements(p, i) for i, p in enumerate(obj["pages"])]
            return [p for p in pages if p]
        if isinstance(obj.get("json_result"), list):
            pages = [_to_page_elements(p, i) for i, p in enumerate(obj["json_result"])]
            return [p for p in pages if p]
        if isinstance(obj.get("layout_details"), list):
            pages = [_to_page_elements(p, i) for i, p in enumerate(obj["layout_details"])]
            return [p for p in pages if p]

    log(f"  GLM parse warning: unsupported output shape for {source_name}: {type(obj).__name__}")
    return []


def _preserve_glm_default_output(results: Any, out_dir: Optional[Path], log, source_name: str) -> None:
    """Persist native GLM output artifacts (json/md/imgs/layout_vis) if supported."""
    if out_dir is None:
        return
    try:
        out_dir.mkdir(parents=True, exist_ok=True)
        if hasattr(results, "save"):
            results.save(output_dir=str(out_dir))
            log(f"  Preserved native GLM output at: {out_dir}")
        else:
            log(f"  GLM result object has no save() method for {source_name}; skipping native output export.")
    except Exception as e:
        log(f"  Could not preserve native GLM output for {source_name}: {e}")


_IMG_COPY_EXTS = frozenset({".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp", ".gif"})


def _collect_glm_figure_images(
    results: Any,
    glm_save_dir: Optional[Path],
    images_dir: Optional[Path],
    prefix: str,
    log,
) -> Dict[str, str]:
    """Copy figure/image files produced by GLM OCR into images_dir.

    If glm_save_dir already exists (populated by _preserve_glm_default_output),
    images are sourced from there.  Otherwise results.save() is called into a
    temporary directory and image files are pulled from that.

    Returns {original_basename_or_abs_path → saved_filename} for image_path updates.
    """
    if images_dir is None:
        return {}

    src_dir: Optional[Path] = None
    tmp_dir_obj = None

    if glm_save_dir is not None and glm_save_dir.is_dir():
        src_dir = glm_save_dir
    elif hasattr(results, "save"):
        tmp_dir_obj = tempfile.TemporaryDirectory(prefix="glm_fig_")
        src_dir = Path(tmp_dir_obj.name)
        try:
            results.save(output_dir=str(src_dir))
        except Exception as e:
            log(f"  Warning: GLM save for figure extraction failed: {e}")
            src_dir = None

    mapping: Dict[str, str] = {}
    if src_dir and src_dir.is_dir():
        try:
            images_dir.mkdir(parents=True, exist_ok=True)
            count = 0
            for f in sorted(src_dir.rglob("*")):
                if not f.is_file() or f.suffix.lower() not in _IMG_COPY_EXTS:
                    continue
                dst_name = f"{prefix}{f.name}"
                dst_path = images_dir / dst_name
                n = 1
                while dst_path.exists():
                    dst_name = f"{prefix}{f.stem}_{n}{f.suffix}"
                    dst_path = images_dir / dst_name
                    n += 1
                shutil.copy2(str(f), str(dst_path))
                mapping[f.name] = dst_name
                mapping[str(f)] = str(dst_path)
                count += 1
            if count:
                log(f"  Saved {count} figure image(s) to: {images_dir}")
        except Exception as e:
            log(f"  Warning: could not copy GLM figure images: {e}")

    if tmp_dir_obj is not None:
        try:
            tmp_dir_obj.cleanup()
        except Exception:
            pass

    return mapping


def _apply_image_path_remap(pages: List[List[Dict]], mapping: Dict[str, str]) -> None:
    """Rewrite image_path in every element using the provided old→new mapping."""
    if not mapping:
        return
    for page in pages:
        for el in page:
            ip = el.get("image_path")
            if not ip:
                continue
            new = mapping.get(ip) or mapping.get(Path(ip).name)
            if new:
                el["image_path"] = new


_FIGURE_LABELS = {"figure", "image", "picture", "photo"}


def _crop_and_save_figure_regions(
    src_img_bytes: bytes,
    pages: List[List[Dict]],
    images_dir: Optional[Path],
    prefix: str,
    log,
) -> None:
    """Crop figure blocks from src_img_bytes using bbox_2d and save to images_dir.

    Updates image_path in each element in-place.  Only processes elements where
    native_label is a figure variant AND image_path is not already set (so it
    complements, not replaces, the GLM-save path collected earlier).
    Requires Pillow; silently skips if not installed.
    """
    if images_dir is None or not HAS_PIL or not src_img_bytes:
        return
    try:
        src = _PILImage.open(_pil_io.BytesIO(src_img_bytes)).convert("RGB")
        w, h = src.size
        images_dir.mkdir(parents=True, exist_ok=True)
        count = 0
        for pg_idx, page in enumerate(pages):
            for el_idx, el in enumerate(page):
                nl = (el.get("native_label") or "").lower()
                if nl not in _FIGURE_LABELS:
                    continue
                # Skip elements that already got an image_path from GLM save artifacts.
                if el.get("image_path"):
                    continue
                bbox = el.get("bbox_2d")
                if not bbox or len(bbox) < 4:
                    continue
                try:
                    x0, y0, x1, y1 = (float(v) for v in bbox[:4])
                    x0, y0 = max(0.0, x0), max(0.0, y0)
                    x1, y1 = min(float(w), x1), min(float(h), y1)
                    if x1 <= x0 or y1 <= y0:
                        continue
                    crop = src.crop((int(x0), int(y0), int(x1), int(y1)))
                    count += 1
                    fname = f"{prefix}{count:03d}.png"
                    crop.save(str(images_dir / fname), format="PNG")
                    el["image_path"] = fname
                except Exception:
                    continue
        if count:
            log(f"  Cropped and saved {count} figure region(s) to: {images_dir}")
    except Exception as e:
        log(f"  Warning: PIL figure crop failed (install Pillow for image extraction): {e}")


def _to_rgb_bytes(img_bytes: bytes) -> bytes:
    """Convert image bytes to RGB PNG so GLM OCR can save JPEG crops internally.

    JPEG does not support alpha channels; if the source image is RGBA, LA, or
    palette mode the glmocr library raises "cannot write mode RGBA as JPEG".
    Converting to RGB before passing to the parser avoids this.
    Returns the original bytes unchanged if PIL is unavailable or conversion fails.
    """
    if not HAS_PIL or not img_bytes:
        return img_bytes
    try:
        img = _PILImage.open(_pil_io.BytesIO(img_bytes))
        if img.mode not in ("RGBA", "LA", "P", "PA"):
            return img_bytes
        rgb = img.convert("RGB")
        buf = _pil_io.BytesIO()
        rgb.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return img_bytes


def _extract_via_pandoc_plain(path: Path, log) -> List[List[Dict]]:
    """Fallback: use pandoc to extract plain text, then parse."""
    try:
        result = subprocess.run(
            ["pandoc", str(path), "-t", "plain"],
            capture_output=True, text=True, encoding="utf-8", shell=True
        )
        if result.returncode == 0 and result.stdout.strip():
            tmp = path.with_suffix(".tmp_plain.txt")
            tmp.write_text(result.stdout, encoding="utf-8")
            data = extract_txt(tmp, log)
            tmp.unlink(missing_ok=True)
            return data
    except Exception as e:
        log(f"  pandoc plain fallback failed: {e}")
    return []


def extract_file(
    path: Path,
    log,
    force_ocr: bool = False,
    glm_native_out_dir: Optional[Path] = None,
    docx_media_dir: Optional[Path] = None,
    ocr_images_dir: Optional[Path] = None,
) -> List[List[Dict]]:
    """
    Route *path* to the right extractor.
    Returns list-of-pages, each page is list-of-element-dicts.
    """
    ext = path.suffix.lower()
    log(f"  Extracting: {path.name}  [{ext}]")

    if ext == ".pdf":
        if force_ocr or _is_scanned_pdf(path):
            log("  → scanned/image PDF → glmocr")
            return extract_pdf_ocr(path, log, glm_native_out_dir, ocr_images_dir)
        else:
            if HAS_OPENDATALOADER_PDF:
                log("  → text-layer PDF → opendataloader-pdf")
                pages = extract_pdf_opendataloader(path, log, ocr_images_dir)
                if pages:
                    return pages
                log("  opendataloader-pdf returned no parseable content; falling back to pymupdf.")
            log("  → text-layer PDF → pymupdf")
            return extract_pdf_text(path, log)
    elif ext == ".docx":
        if force_ocr or _is_image_heavy_docx(path):
            log("  → image-based/force-OCR DOCX → glmocr")
            return extract_docx_ocr(path, log, glm_native_out_dir, ocr_images_dir)
        return extract_docx(path, log, docx_media_dir)
    elif ext in TEXT_EXTS | MD_EXTS:
        return extract_txt(path, log)
    elif ext in IMAGE_EXTS:
        return extract_image(path, log, glm_native_out_dir, ocr_images_dir)
    else:
        log(f"  Unsupported extension '{ext}' – trying plain text.")
        return extract_txt(path, log)


# ────────────────────────────────────────────────────────────────────────────
# SEMANTIC ANNOTATION  (condensed from semantic_annotate.py)
# ────────────────────────────────────────────────────────────────────────────

_PROCED_VERBS = {
    "remove","install","connect","disconnect","check","verify","ensure","set",
    "select","perform","apply","close","open","turn","press","push","pull",
    "insert","tighten","loosen","adjust","align","attach","detach","activate",
    "deactivate","enable","disable","start","stop","examine","inspect",
    "measure","replace","assemble","disassemble","clean","drain","fill",
    "lubricate","torque","test","operate","position","lock","unlock","secure",
    "release","engage","disengage","enter","exit","switch","rotate","slide",
    "interlock","monitor","record","note","caution","warn",
}

_PROCED_PATS = [
    r"^\s*\d+[\.\)]\s+[A-Z]", r"^\s*[a-z][\.\)]\s+[A-Z]",
    r"^\s*step\s+\d+",
    r"\b(must|shall)\s+be\s+(installed|removed|checked|applied|replaced|performed|verified|connected|secured)",
    r"\b(warning|caution|note)\b.*:", r"torque\s+to\b", r"\bdo\s+not\b",
]
_DESCRIPT_PATS = [
    r"\bis\s+used\s+to\b", r"\bare\s+used\s+to\b", r"\bis\s+defined\s+as\b",
    r"\bprovide[sd]?\b", r"\bcontain[sd]?\b", r"\binclude[sd]?\b",
    r"\bthis\s+(section|chapter|module|document)\b",
    r"\bthe\s+purpose\s+of\b", r"\bdescribe[sd]?\b", r"\bexplain[sd]?\b",
]


def _score(text: str) -> Tuple[str, float]:
    if not text or len(text.strip()) < 3:
        return "descript", 0.5
    lower = text.lower().strip()
    first = lower.split()[0].rstrip(".,;:") if lower.split() else ""
    ps, ds = 0.0, 0.0
    if first in _PROCED_VERBS:
        ps += 0.6
    for pat in _PROCED_PATS:
        if re.search(pat, text, re.IGNORECASE | re.MULTILINE):
            ps += 0.4; break
    steps = re.findall(r"^\s*(\d+[\.\)]|[-•])\s+", text, re.MULTILINE)
    if len(steps) >= 2:
        ps += 0.3
    for pat in _DESCRIPT_PATS:
        if re.search(pat, text, re.IGNORECASE):
            ds += 0.35; break
    if len(text) > 200 and not steps:
        ds += 0.25
    if re.match(r"^\s*#{1,6}\s", text):
        ds += 0.5
    if ps == 0 and ds == 0:
        return "descript", 0.55
    if ps > ds:
        return "proced", round(min(0.95, 0.5 + ps - ds), 2)
    return "descript", round(min(0.95, 0.5 + ds - ps), 2)


_ROMAN_STEP_RE = re.compile(r"^\s*((?:x|ix|iv|v?i{1,3}))\s*[\.)]\s+(.+)$", re.IGNORECASE)
_ALPHA_STEP_RE = re.compile(r"^\s*([a-z])\s*[\.)]\s+(.+)$", re.IGNORECASE)
_MAIN_STEP_RE = re.compile(r"^\s*(\d+)\s*[\.)]\s+(.+)$")
_STEP_WORD_RE = re.compile(r"^\s*(?:step\s+)?(\d+)\s*[:\.)-]\s+(.+)$", re.IGNORECASE)
_BULLET_STEP_RE = re.compile(r"^\s*[-*•]\s+(.+)$")
_WARN_RE = re.compile(r"^\s*WARNING\s*:\s*(.+)$", re.IGNORECASE)
_CAUTION_RE = re.compile(r"^\s*CAUTION\s*:\s*(.+)$", re.IGNORECASE)
_NOTE_RE = re.compile(r"^\s*NOTE\s*:\s*(.+)$", re.IGNORECASE)
_COND_RE = re.compile(r"^\s*(if\b|when\b|ensure(?:\s+that)?\b)(.+)$", re.IGNORECASE)
_META_KV_RE = re.compile(r"^\s*([A-Za-z][A-Za-z0-9\s\./_-]{1,40})\s*:\s*(.+)$")
_PART_RE = re.compile(r"\b(?:P/?N|Part(?:\s*No)?|Manf\s*Part\s*No|Part\s*Number)\s*[:#-]?\s*([A-Z0-9-]{2,})\b", re.IGNORECASE)
_TOOL_HINT_RE = re.compile(r"\b(wrench|torque\s*wrench|spanner|screwdriver|socket|pliers|gauge|fixture|tool)\b", re.IGNORECASE)


def _extract_metadata_refs(text: str) -> Tuple[List[str], List[str]]:
    tools: List[str] = []
    parts: List[str] = []
    for m in _PART_RE.finditer(text or ""):
        pn = m.group(1).strip()
        if pn and pn not in parts:
            parts.append(pn)
    for m in _TOOL_HINT_RE.finditer(text or ""):
        t = m.group(0).strip()
        if t and t.lower() not in [x.lower() for x in tools]:
            tools.append(t)
    return tools, parts


def _new_struct_step(step_number: str, text: str) -> Dict[str, Any]:
    tools, parts = _extract_metadata_refs(text)
    return {
        "step_number": str(step_number),
        "text": text.strip(),
        "substeps": [],
        "warnings": [],
        "cautions": [],
        "notes": [],
        "conditions": [],
        "tools": tools,
        "parts": parts,
    }


def _detect_step_line(line: str) -> Optional[Tuple[int, str, str]]:
    m = _MAIN_STEP_RE.match(line)
    if m:
        return (1, m.group(1), m.group(2).strip())
    m = _STEP_WORD_RE.match(line)
    if m:
        return (1, m.group(1), m.group(2).strip())
    m = _ALPHA_STEP_RE.match(line)
    if m:
        return (2, m.group(1).lower(), m.group(2).strip())
    m = _ROMAN_STEP_RE.match(line)
    if m:
        return (3, m.group(1).lower(), m.group(2).strip())
    m = _BULLET_STEP_RE.match(line)
    if m:
        return (2, "-", m.group(1).strip())

    first = (line.strip().split()[:1] or [""])[0].lower().rstrip(".,;:")
    if first in _PROCED_VERBS:
        return (1, "", line.strip())
    return None


def _extract_structured_steps(text: str) -> List[Dict[str, Any]]:
    lines = [l.rstrip() for l in (text or "").splitlines() if l.strip()]
    steps: List[Dict[str, Any]] = []
    stack: List[Dict[str, Any]] = []
    auto_main = 1

    def _attach_to_current(field: str, value: str) -> bool:
        if not stack:
            return False
        cur = stack[-1]
        if value and value not in cur[field]:
            cur[field].append(value)
        return True

    for raw in lines:
        line = raw.strip()

        m = _WARN_RE.match(line)
        if m:
            _attach_to_current("warnings", m.group(1).strip())
            continue
        m = _CAUTION_RE.match(line)
        if m:
            _attach_to_current("cautions", m.group(1).strip())
            continue
        m = _NOTE_RE.match(line)
        if m:
            _attach_to_current("notes", m.group(1).strip())
            continue

        m = _COND_RE.match(line)
        if m:
            cond = (m.group(1) + m.group(2)).strip()
            if not _attach_to_current("conditions", cond):
                node = _new_struct_step(str(auto_main), line)
                auto_main += 1
                node["conditions"].append(cond)
                steps.append(node)
                stack = [node]
            continue

        detected = _detect_step_line(line)
        if detected:
            level, step_no, txt = detected
            if not step_no:
                step_no = str(auto_main)
                auto_main += 1
            node = _new_struct_step(step_no, txt)
            while len(stack) >= level:
                stack.pop()
            if stack:
                stack[-1]["substeps"].append(node)
            else:
                steps.append(node)
            stack.append(node)
            continue

        # Metadata or continuation text.
        kv = _META_KV_RE.match(line)
        if kv and stack:
            k = kv.group(1).strip().lower()
            v = kv.group(2).strip()
            if "tool" in k:
                if v not in stack[-1]["tools"]:
                    stack[-1]["tools"].append(v)
            elif "part" in k or "p/n" in k:
                if v not in stack[-1]["parts"]:
                    stack[-1]["parts"].append(v)
            elif "person" in k:
                if v not in stack[-1]["notes"]:
                    stack[-1]["notes"].append(f"personnel: {v}")
            else:
                stack[-1]["notes"].append(f"{kv.group(1).strip()}: {v}")
            continue

        if stack:
            stack[-1]["text"] = (stack[-1]["text"] + " " + line).strip()
            t_tools, t_parts = _extract_metadata_refs(line)
            for t in t_tools:
                if t not in stack[-1]["tools"]:
                    stack[-1]["tools"].append(t)
            for p in t_parts:
                if p not in stack[-1]["parts"]:
                    stack[-1]["parts"].append(p)
        else:
            node = _new_struct_step(str(auto_main), line)
            auto_main += 1
            steps.append(node)
            stack = [node]

    return steps


def _semantic_heading_level(el: Dict) -> Optional[int]:
    raw = (el.get("content") or "").strip()
    nl = (el.get("native_label") or "").strip()
    m = _HEADING_RE.match(raw)
    if m:
        return max(1, min(6, len(m.group(1))))
    if nl == "doc_title":
        return 1
    if nl == "paragraph_title":
        return 2
    return None


def build_structured_semantic(pages: List[List[Dict]]) -> Dict[str, Any]:
    """Build heading-grouped semantic view with steps/metadata/schema hints."""
    root: Dict[str, Any] = {
        "title": "Document",
        "level": 0,
        "section_type": "root",
        "schema_hint": "unknown",
        "metadata": {},
        "steps": [],
        "warnings": [],
        "cautions": [],
        "notes": [],
        "paragraphs": [],
        "line_items": [],
        "children": [],
        "_raw_lines": [],
    }
    stack: List[Dict[str, Any]] = [root]

    def _new_section(title: str, level: int) -> Dict[str, Any]:
        return {
            "title": title.strip() or "Untitled",
            "level": level,
            "section_type": f"H{level}",
            "schema_hint": "unknown",
            "metadata": {},
            "steps": [],
            "warnings": [],
            "cautions": [],
            "notes": [],
            "paragraphs": [],
            "line_items": [],
            "children": [],
            "_raw_lines": [],
        }

    flat_elements = [el for page in pages for el in (page if isinstance(page, list) else [page])]
    order_counter = 0
    for el in flat_elements:
        raw = (el.get("content") or "").strip()
        if not raw:
            continue
        level = _semantic_heading_level(el)
        if level is not None:
            title = _clean_text(_HEADING_RE.sub(r"\2", raw))
            while len(stack) > 1 and stack[-1]["level"] >= level:
                stack.pop()
            sec = _new_section(title, level)
            stack[-1]["children"].append(sec)
            stack.append(sec)
            continue

        cur = stack[-1]
        lines = [l.strip() for l in raw.splitlines() if l.strip()]
        cur["_raw_lines"].extend(lines)

        for line in lines:
            order_counter += 1
            line_class = "paragraph"
            wm = _WARN_RE.match(line)
            cm = _CAUTION_RE.match(line)
            nm = _NOTE_RE.match(line)
            if wm:
                line_class = "warning"
                cur["warnings"].append(wm.group(1).strip())
            elif cm:
                line_class = "caution"
                cur["cautions"].append(cm.group(1).strip())
            elif nm:
                line_class = "note"
                cur["notes"].append(nm.group(1).strip())
            else:
                kv = _META_KV_RE.match(line)
                if kv:
                    line_class = "metadata"
                elif _detect_step_line(line):
                    line_class = "step"

            cur["line_items"].append({
                "order": order_counter,
                "text": line,
                "class": line_class,
                "bbox": el.get("bbox_2d") or el.get("bbox"),
                "source_label": el.get("native_label", "paragraph"),
            })

            if line_class in ("warning", "caution", "note"):
                continue

            kv = _META_KV_RE.match(line)
            if kv:
                k = kv.group(1).strip().lower()
                v = kv.group(2).strip()
                if "no. of persons" in k or "number of persons" in k or "personnel" in k:
                    try:
                        cur["metadata"]["personnel"] = int(re.search(r"\d+", v).group(0))
                    except Exception:
                        cur["metadata"]["personnel"] = v
                elif "manf part" in k or "part no" in k or "part number" in k:
                    cur["metadata"]["part_number"] = v
                else:
                    cur["metadata"][k.replace(" ", "_")] = v
            else:
                cur["paragraphs"].append(line)

    def _finalize(sec: Dict[str, Any]) -> None:
        raw_text = "\n".join(sec.pop("_raw_lines", []))
        steps = _extract_structured_steps(raw_text)
        sec["steps"] = steps
        step_count = len(steps)
        para_count = len(sec.get("paragraphs", []))
        lowered = raw_text.lower()
        if step_count > 0:
            sec["schema_hint"] = "proced"
        elif any(k in lowered for k in ("fault", "troubleshoot", "isolation")):
            sec["schema_hint"] = "fault"
        elif para_count > 0:
            sec["schema_hint"] = "descript"
        else:
            sec["schema_hint"] = "unknown"
        for child in sec.get("children", []):
            _finalize(child)

    _finalize(root)
    return {"sections": root.get("children", []), "schema_hint": root.get("schema_hint", "unknown")}


def _sem_block(text: str, sem_type: str, conf: float) -> Dict:
    if sem_type == "proced":
        steps = _extract_structured_steps(text)
        if not steps:
            steps = [_new_struct_step("1", text.strip())]
        title = steps[0]["text"][:60]
        return {"type": "proced", "title": title,
                "structure": {"steps": steps}, "confidence": conf}
    else:
        paras = [p.strip() for p in re.split(r"\n{2,}", text.strip()) if p.strip()]
        if not paras:
            paras = [text.strip()]
        m = re.match(r"^#{1,6}\s*(.+)", text.strip())
        title = m.group(1).strip() if m else paras[0][:80]
        return {"type": "descript", "title": title,
                "structure": {"paragraphs": paras}, "confidence": conf}


def annotate_pages(pages: List[List[Dict]]) -> List[List[Dict]]:
    """Add 'semantic' field to every element in every page."""
    annotated = []
    for page in pages:
        a_page = []
        for el in (page if isinstance(page, list) else [page]):
            content = el.get("content", "") or ""
            nl = el.get("native_label", "")
            if nl in ("paragraph_title", "doc_title"):
                sem = _sem_block(content, "descript", 0.95)
            else:
                t, c = _score(content)
                sem = _sem_block(content, t, c)
            a_page.append({**el, "semantic": sem})
        annotated.append(a_page)
    return annotated


# ────────────────────────────────────────────────────────────────────────────
# XML GENERATION  (S1000D Issue 4.2)
# ────────────────────────────────────────────────────────────────────────────

_HEADING_RE  = re.compile(r"^(#{1,6})\s+(.+)")
_STEP_PFX_RE = re.compile(r"^(Step|step)\s+\d+\s*[:\.\)]\s*")
_BULLET_RE   = re.compile(r"^[-•]\s+(.+)")
_NUM_RE      = re.compile(r"^\d+[\.\)]\s+(.+)")


def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = _md_links_to_text_and_url(text)
    text = _HEADING_RE.sub(r"\2", text)
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"<[^>]+>", "", text)
    return " ".join(text.split())


def _dominant_type(elements: List[Dict]) -> str:
    c = {"proced": 0, "descript": 0}
    for el in elements:
        t = el.get("semantic", {}).get("type", "descript")
        c[t] = c.get(t, 0) + 1
    return "proced" if c["proced"] > c["descript"] else "descript"


def _make_ident(parent: ET.Element, dm_code: str, title: str,
                dm_type: str, issue_no: str = "001", dm_variant: Optional[str] = None):
    ias = ET.SubElement(parent, "identAndStatusSection")
    dmAddr = ET.SubElement(ias, "dmAddress")
    dmIdent = ET.SubElement(dmAddr, "dmIdent")
    parts = (dm_code or "").split("-")
    dc = ET.SubElement(dmIdent, "dmCode")
    dc.set("modelIdentCode",      parts[0] if len(parts) > 0 else "UNKWN")
    dc.set("systemDiffCode",      parts[1] if len(parts) > 1 else "A")
    dc.set("systemCode",          parts[2] if len(parts) > 2 else "00")
    dc.set("subSystemCode",       parts[3][:1] if len(parts) > 3 else "0")
    dc.set("subSubSystemCode",    parts[3][1:] if len(parts) > 3 and len(parts[3]) > 1 else "0")
    dc.set("assyCode",            parts[4] if len(parts) > 4 else "0000")
    dc.set("disassyCode",         parts[5][:2] if len(parts) > 5 else "00")
    dc.set("disassyCodeVariant",  parts[5][2:] if len(parts) > 5 and len(parts[5]) > 2 else "A")
    dc.set("infoCode",            parts[6][:3] if len(parts) > 6 else "040")
    if dm_variant:
        dc.set("infoCodeVariant", dm_variant)
    else:
        dc.set("infoCodeVariant", parts[6][3:] if len(parts) > 6 and len(parts[6]) > 3 else "A")
    dc.set("itemLocationCode",    parts[7] if len(parts) > 7 else "A")

    lang = ET.SubElement(dmIdent, "language")
    lang.set("languageIsoCode", "en")
    lang.set("countryIsoCode", "US")
    ii = ET.SubElement(dmIdent, "issueInfo")
    ii.set("issueNumber", issue_no)
    ii.set("inWork", "00")

    items = ET.SubElement(dmAddr, "dmAddressItems")
    idate = ET.SubElement(items, "issueDate")
    idate.set("year", "2026"); idate.set("month", "04"); idate.set("day", "04")
    dmTitle = ET.SubElement(items, "dmTitle")
    if dm_type == "sched" and " - " in title:
        _tp = title.split(" - ", 1)
        ET.SubElement(dmTitle, "techName").text = _tp[0].strip()
        ET.SubElement(dmTitle, "infoName").text  = _tp[1].strip()
    else:
        ET.SubElement(dmTitle, "techName").text = title

    dmStatus = ET.SubElement(ias, "dmStatus")
    dmStatus.set("issueType", "new")
    sec = ET.SubElement(dmStatus, "security")
    sec.set("securityClassification", "01")
    ET.SubElement(dmStatus, "responsiblePartnerCompany")
    ET.SubElement(dmStatus, "originator")
    applic = ET.SubElement(dmStatus, "applic")
    disp = ET.SubElement(applic, "displayText")
    ET.SubElement(disp, "simplePara").text = "All"
    brexRef = ET.SubElement(dmStatus, "brexDmRef")
    dmRef = ET.SubElement(brexRef, "dmRef")
    dmRefI = ET.SubElement(dmRef, "dmRefIdent")
    bc = ET.SubElement(dmRefI, "dmCode")
    bc.set("modelIdentCode", "S1000D"); bc.set("systemDiffCode", "H")
    bc.set("systemCode", "041"); bc.set("subSystemCode", "1")
    bc.set("subSubSystemCode", "0"); bc.set("assyCode", "0301")
    bc.set("disassyCode", "00"); bc.set("disassyCodeVariant", "A")
    bc.set("infoCode", "022"); bc.set("infoCodeVariant", "A")
    bc.set("itemLocationCode", "D")
    qa = ET.SubElement(dmStatus, "qualityAssurance")
    ET.SubElement(qa, "unverified")


def _build_description(parent: ET.Element, elements: List[Dict]):
    desc = ET.SubElement(parent, "description")
    stack: List[Tuple[int, ET.Element]] = [(0, desc)]

    def cur():
        return stack[-1][1]

    for el in elements:
        raw = el.get("content", "") or ""
        ct  = _clean_text(raw)
        nl  = el.get("native_label", "") or ""
        if not ct:
            continue
        m = _HEADING_RE.match(raw.strip())
        if m or nl in ("doc_title", "paragraph_title"):
            depth = len(m.group(1)) if m else 1
            while len(stack) > 1 and stack[-1][0] >= depth:
                stack.pop()
            lp = ET.SubElement(cur(), "levelledPara")
            ET.SubElement(lp, "title").text = _clean_text(m.group(2) if m else raw)
            stack.append((depth, lp))
        elif nl == "table":
            _simple_table(cur(), raw)
        elif nl == "figure":
            fig = ET.SubElement(cur(), "figure")
            g = ET.SubElement(fig, "graphic")
            g.set("infoEntityIdent", el.get("image_path") or "UNKNOWN")
        else:
            # Check for list
            bullet_items = [_BULLET_RE.match(l) for l in raw.splitlines() if l.strip()]
            num_items    = [_NUM_RE.match(l)    for l in raw.splitlines() if l.strip()]
            b_count = sum(1 for x in bullet_items if x)
            n_count = sum(1 for x in num_items    if x)
            if b_count >= 2:
                _emit_list(cur(), [x.group(1) for x in bullet_items if x], "randomList")
            elif n_count >= 2:
                _emit_list(cur(), [x.group(1) for x in num_items if x],    "sequentialList")
            else:
                ET.SubElement(cur(), "para").text = ct


def _emit_list(parent: ET.Element, items: List[str], tag: str):
    wrapper = ET.SubElement(parent, "para")
    lst = ET.SubElement(wrapper, tag)
    for item in items:
        li = ET.SubElement(lst, "listItem")
        ET.SubElement(li, "para").text = _clean_text(item)


def _simple_table(parent: ET.Element, text: str):
    rows = [r.strip() for r in text.splitlines() if r.strip()]
    ncols = max((len(r.split("|")) for r in rows), default=1)
    tbl = ET.SubElement(parent, "table")
    tbl.set("frame", "all"); tbl.set("pgwide", "1")
    tg = ET.SubElement(tbl, "tgroup"); tg.set("cols", str(ncols))
    for i in range(1, ncols + 1):
        cs = ET.SubElement(tg, "colspec"); cs.set("colname", f"c{i}")
    tbody = ET.SubElement(tg, "tbody")
    for row_text in rows:
        cells = [c.strip() for c in row_text.split("|")]
        row_el = ET.SubElement(tbody, "row")
        for i in range(ncols):
            ct_text = cells[i] if i < len(cells) else ""
            entry = ET.SubElement(row_el, "entry")
            entry.set("namest", f"c{i+1}"); entry.set("nameend", f"c{i+1}")
            ET.SubElement(entry, "para").text = ct_text or None


def _append_structured_procedural_step(parent: ET.Element, step: Dict[str, Any]) -> None:
    """Append one structured step (with nested substeps) as S1000D proceduralStep XML."""
    ps = ET.SubElement(parent, "proceduralStep")
    main_txt = _clean_text(step.get("text", "") or "")
    if main_txt:
        ET.SubElement(ps, "para").text = main_txt

    for cond in step.get("conditions", []) or []:
        ET.SubElement(ps, "para").text = _clean_text(f"Condition: {cond}")
    for w in step.get("warnings", []) or []:
        ET.SubElement(ps, "para").text = _clean_text(f"WARNING: {w}")
    for c in step.get("cautions", []) or []:
        ET.SubElement(ps, "para").text = _clean_text(f"CAUTION: {c}")
    for n in step.get("notes", []) or []:
        ET.SubElement(ps, "para").text = _clean_text(f"NOTE: {n}")

    tools = step.get("tools", []) or []
    parts = step.get("parts", []) or []
    if tools:
        ET.SubElement(ps, "para").text = _clean_text("Tools: " + ", ".join(str(t) for t in tools if t))
    if parts:
        ET.SubElement(ps, "para").text = _clean_text("Parts: " + ", ".join(str(p) for p in parts if p))

    for sub in step.get("substeps", []) or []:
        if isinstance(sub, dict):
            _append_structured_procedural_step(ps, sub)


def _build_procedure(parent: ET.Element, elements: List[Dict]):
    proc = ET.SubElement(parent, "procedure")
    pr = ET.SubElement(proc, "preliminaryRqmts")
    for tag in ("reqCondGroup", "reqPersons", "reqSupportEquips",
                "reqSupplies", "reqSpares", "reqSafety"):
        ET.SubElement(pr, tag)
    mp = ET.SubElement(proc, "mainProcedure")
    for el in elements:
        raw = el.get("content", "") or ""
        ct  = _clean_text(raw)
        nl  = el.get("native_label", "") or ""
        if not ct:
            if nl != "figure":
                continue
        if nl == "table":
            ps = ET.SubElement(mp, "proceduralStep")
            _simple_table(ps, raw)
            continue
        if nl == "figure":
            ps = ET.SubElement(mp, "proceduralStep")
            fig = ET.SubElement(ps, "figure")
            g = ET.SubElement(fig, "graphic")
            g.set("infoEntityIdent", el.get("image_path") or "UNKNOWN")
            continue
        steps = el.get("semantic", {}).get("structure", {}).get("steps", [])
        if steps:
            for sd in steps:
                if isinstance(sd, dict):
                    _append_structured_procedural_step(mp, sd)
        else:
            ps = ET.SubElement(mp, "proceduralStep")
            ET.SubElement(ps, "para").text = ct
    cr = ET.SubElement(proc, "closeRqmts")
    ET.SubElement(cr, "reqCondGroup")


def _build_fault(parent: ET.Element, elements: List[Dict]):
    fr = ET.SubElement(parent, "faultReporting")
    ET.SubElement(fr, "fault")
    fi = ET.SubElement(parent, "faultIsolation")
    for el in elements:
        raw = el.get("content", "") or ""
        nl = (el.get("native_label", "") or "").strip()
        if nl == "table":
            _simple_table(fi, raw)
            continue
        if nl == "figure":
            fig = ET.SubElement(fi, "figure")
            g = ET.SubElement(fig, "graphic")
            g.set("infoEntityIdent", el.get("image_path") or "UNKNOWN")
            continue
        ct = _clean_text(raw)
        if ct:
            ET.SubElement(fi, "para").text = ct


# ─── Scheduled Maintenance helpers ─────────────────────────────────────────

_SCHED_SUBSEC_PATS: List[Tuple[str, str]] = [
    # Each pattern uses \b at the START only so plurals / partial stems are matched too.
    ("reqCondGroup",     r"\b(condition|precondition|prerequisite|initial\s+cond)"),
    ("reqPersons",       r"\b(persons?|personnel|operator|technician|mechanic|skill\s+level|trade)"),
    ("reqSupportEquips", r"\b(support\s+equip|equipment|tools?\b|instrument|gauge|jig|fixture)"),
    ("reqSupplies",      r"\b(suppli|supply\b|material|lubricant|cleaning|fluid|compound|sealant|coolant)"),
    ("reqSpares",        r"\b(spare|consumable|filter|plug|pad|bearing|seal|gasket|component|part\s+number)"),
    ("reqSafety",        r"\b(safety|warning|caution|protect|precaution|hazard|ppe)"),
    ("limits",           r"\b(time\s+limit|limits?|interval|threshold|schedule|period|frequency|overhaul)"),
    ("refs",             r"\b(reference|related\s+doc|see\s+also|dmref)"),
]

_SCHED_UNIT_OF_MEASURE: List[Tuple[str, str]] = [
    (r"\bflight[\s-]hour", "th09"),
    (r"\bfh\b",             "th09"),
    (r"\bhour",             "th01"),
    (r"\bday",              "th02"),
    (r"\bmonth",            "th03"),
    (r"\bquarter",          "th04"),
    (r"\byear|annual",      "th06"),
    (r"\bcycle",            "th07"),
    (r"\bkm\b|kilometer",   "th10"),
    (r"\bmile",             "th13"),
]

_SCHED_INSP_CATEGORY: List[Tuple[str, str]] = [
    (r"\bdaily\b",                          "Daily"),
    (r"\bweekly\b",                         "Weekly"),
    (r"\bmonthly\b",                        "Monthly"),
    (r"\bquarterly\b|\b3[\s-]month",        "Quarterly"),
    (r"\bsemi[\s-]?annual\b|\b6[\s-]month", "SemiAnnually"),
    (r"\bannual\b|\byearly\b",              "Annually"),
    (r"\bbiennial\b|\b2[\s-]year",          "Biennially"),
    (r"\boverhaul\b",                       "Overhaul"),
    (r"\bon[\s-]condition\b|as\s+required", "OnCondition"),
    (r"\b500[\s-]h",                        "500hr"),
    (r"\b1000[\s-]h",                       "1000hr"),
]

_SCHED_LIMIT_TYPE: List[Tuple[str, str]] = [
    (r"\bhard[\s-]time\b|\bperiodic\b|\bscheduled\b", "po"),
    (r"\bon[\s-]condition\b|\boc\b",                  "oc"),
    (r"\bsafe[\s-]life\b|\bsl\b|\blife[\s-]limit",   "sl"),
    (r"\bfailure[\s-]find\b|\bffd?\b",                "ff"),
]


def _sched_detect_unit(text: str) -> str:
    t = text.lower()
    for pat, code in _SCHED_UNIT_OF_MEASURE:
        if re.search(pat, t):
            return code
    return "th06"


def _sched_detect_insp(text: str) -> str:
    t = text.lower()
    for pat, val in _SCHED_INSP_CATEGORY:
        if re.search(pat, t):
            return val
    return "Periodic"


def _sched_detect_limit_type(text: str) -> str:
    t = text.lower()
    for pat, val in _SCHED_LIMIT_TYPE:
        if re.search(pat, t):
            return val
    return "po"


def _sched_extract_number(text: str) -> str:
    m = re.search(r"(\d+(?:[.,]\d+)?)", text)
    return m.group(1).replace(",", ".") if m else "1"


def _sched_is_task_heading(el: Dict) -> bool:
    """Return True if this element begins a new scheduled-maintenance task block."""
    nl  = el.get("native_label", "") or ""
    raw = (el.get("content", "") or "").strip()
    if nl not in ("paragraph_title", "doc_title"):
        return False
    clean = _clean_text(raw)
    if re.search(
        r"\btask\b|\bmaintenance\s+task\b|\binspect\b|\bcheck\b|\bservice\b|\brepla",
        clean, re.IGNORECASE,
    ):
        return True
    if re.match(r"^(?:#{1,6}\s*)?\d+[\.)\]\s]\s*\S", clean):
        return True
    return False


def _sched_new_task(ident: str, heading_text: str) -> Dict:
    applic_m  = re.search(r"applic[_\s-]?ref\s*=\s*([\w-]+)", heading_text, re.IGNORECASE)
    applic    = applic_m.group(1).strip('[]"\'  ') if applic_m else None
    code_m    = re.search(r"\btask[\s_-]?code\s*[=:]\s*([\w-]+)", heading_text, re.IGNORECASE)
    task_code = code_m.group(1) if code_m else ""
    worth     = "mandatory" if re.search(r"\bmandatory\b", heading_text, re.IGNORECASE) else "recommended"
    reduced   = "Yes" if re.search(r"\breduced\b|\bsimplif", heading_text, re.IGNORECASE) else "No"
    skill_m   = re.search(
        r"\bst(\d{2})\b|\bskill[\s-]*(?:level[\s-]*)?(\d)\b", heading_text, re.IGNORECASE
    )
    skill = (
        f"st{(skill_m.group(1) or skill_m.group(2)).zfill(2)}"
        if skill_m else "st01"
    )
    return {
        "task_ident": ident, "task_code": task_code,
        "worthiness": worth, "reduced_maint": reduced,
        "skill_type": skill, "applic_ref": applic,
        "description": "", "conditions": [], "persons": [],
        "equipment": [], "supplies": [], "spares": [],
        "safety": [], "limits": [], "refs": [], "extra": [],
    }


def _parse_sched_tasks(elements: List[Dict]) -> List[Dict]:
    """Group flat element list into structured scheduled-maintenance task dicts."""
    tasks:   List[Dict]         = []
    cur:     Optional[Dict]     = None
    section: str                = "description"
    counter: int                = 1

    def _add(text: str) -> None:
        nonlocal section
        if cur is None:
            return
        key_map = {
            "description":    "description",
            "reqCondGroup":   "conditions",
            "reqPersons":     "persons",
            "reqSupportEquips": "equipment",
            "reqSupplies":    "supplies",
            "reqSpares":      "spares",
            "reqSafety":      "safety",
            "limits":         "limits",
            "refs":           "refs",
        }
        field = key_map.get(section, "extra")
        if field == "description":
            if cur["description"]:
                cur["extra"].append(text)
            else:
                cur["description"] = text
        else:
            cur[field].append(text)

    for el in elements:
        raw = (el.get("content", "") or "").strip()
        nl  = el.get("native_label", "") or ""

        # Preserve table payload as extra task detail; table layout can be reconstructed downstream.
        if nl == "table" and raw:
            if cur is None:
                cur = _sched_new_task(str(counter).zfill(3), "Task from schedule table")
                counter += 1
            cur["extra"].append(raw)
            continue

        ct  = _clean_text(raw)
        if not ct:
            continue

        if _sched_is_task_heading(el):
            if cur is not None:
                tasks.append(cur)
            cur     = _sched_new_task(str(counter).zfill(3), ct)
            counter += 1
            section = "description"
        elif nl == "paragraph_title" and cur is not None:
            detected = next(
                (sec for sec, pat in _SCHED_SUBSEC_PATS if re.search(pat, ct, re.IGNORECASE)),
                None,
            )
            if detected:
                section = detected
            else:
                _add(ct)
        else:
            if cur is None:
                cur     = _sched_new_task(str(counter).zfill(3), ct)
                counter += 1
                section = "description"
            else:
                _add(ct)

    if cur is not None:
        tasks.append(cur)

    if not tasks:
        all_text = " ".join(
            _clean_text(el.get("content", "") or "")
            for el in elements if el.get("content")
        ).strip()
        if all_text:
            t = _sched_new_task("001", all_text)
            t["description"] = all_text
            tasks.append(t)
    return tasks


def _sched_add_item_descr(parent: ET.Element, item_text: str, tag: str) -> ET.Element:
    """Append a descriptor element (supportEquipDescr / supplyDescr / spareDescr)."""
    clean  = _clean_text(item_text)
    pn_m   = re.search(
        r"[\(\[](P/?N|Part|No)\.?\s*:?\s*([A-Z0-9\-]+)[\)\]]", clean, re.IGNORECASE
    )
    pn     = pn_m.group(2) if pn_m else re.sub(r"\s+", "-", clean.upper()[:20])
    descr  = ET.SubElement(parent, tag)
    ET.SubElement(descr, "name").text = clean
    idn    = ET.SubElement(descr, "identNumber")
    ET.SubElement(idn, "manufacturerCode").text = "UNKN"
    psn    = ET.SubElement(idn, "partAndSerialNumber")
    ET.SubElement(psn, "partNumber").text = pn
    return descr


def _build_sched_task_xml(mp: ET.Element, task: Dict) -> None:
    """Append one <taskDefinition> to a <maintPlanning> element."""
    ident = task["task_ident"]
    td = ET.SubElement(mp, "taskDefinition")
    td.set("taskIdent",        ident)
    td.set("worthinessLimit",  task.get("worthiness",    "recommended"))
    td.set("reducedMaint",     task.get("reduced_maint", "No"))
    if task.get("task_code"):  td.set("taskCode",   task["task_code"])
    if task.get("skill_type"): td.set("skillType",  task["skill_type"])
    if task.get("applic_ref"): td.set("applicRefId", task["applic_ref"])

    # <task><taskDescr>
    t_el = ET.SubElement(td, "task")
    d_el = ET.SubElement(t_el, "taskDescr")
    ET.SubElement(d_el, "simplePara").text = (
        _clean_text(task.get("description", "") or f"Task {ident}")
    )

    # Keep any unmapped/task-tail content to prevent schema-path data loss.
    for extra in task.get("extra", []) or []:
        for ln in str(extra).splitlines():
            txt = " ".join((ln or "").split())
            if txt:
                ET.SubElement(d_el, "simplePara").text = txt

    # <preliminaryRqmts>
    pr = ET.SubElement(td, "preliminaryRqmts")

    # reqCondGroup
    cg = ET.SubElement(pr, "reqCondGroup")
    if task["conditions"]:
        for c in task["conditions"]:
            rc = ET.SubElement(cg, "reqCond")
            ET.SubElement(rc, "simplePara").text = _clean_text(c)
    else:
        ET.SubElement(cg, "noConds")

    # reqPersons
    prs = ET.SubElement(pr, "reqPersons")
    if task["persons"]:
        for p_txt in task["persons"]:
            person = ET.SubElement(prs, "person")
            person.set("man", "A")
            cat = ET.SubElement(person, "personCategory")
            if re.search(r"\btechnician\b|\bspecialist\b|\bengineer\b", p_txt, re.IGNORECASE):
                cat.set("personCategoryCode", "Technician")
            elif re.search(r"\boperator\b|\buser\b", p_txt, re.IGNORECASE):
                cat.set("personCategoryCode", "Basic user")
            else:
                cat.set("personCategoryCode", "Maintainer")
            trade_m = re.search(
                r"\b(operator|technician|specialist|mechanic|engineer|maintainer)\b",
                p_txt, re.IGNORECASE,
            )
            ET.SubElement(person, "trade").text = trade_m.group(1).title() if trade_m else "Maintainer"
            time_m = re.search(r"(\d+(?:[.,]\d+)?)\s*h(?:our)?s?\b", p_txt, re.IGNORECASE)
            if time_m:
                et_el = ET.SubElement(person, "estimatedTime")
                et_el.set("unitOfMeasure", "h")
                et_el.text = time_m.group(1).replace(",", ".")
    else:
        p_default = ET.SubElement(prs, "person")
        p_default.set("man", "A")
        ET.SubElement(p_default, "personCategory").set("personCategoryCode", "Maintainer")
        ET.SubElement(p_default, "trade").text = "Maintainer"

    # reqSupportEquips
    se     = ET.SubElement(pr, "reqSupportEquips")
    equips = [x.strip() for txt in task["equipment"]
              for x in re.split(r"[,;]", txt) if x.strip()]
    if equips:
        grp = ET.SubElement(se, "supportEquipDescrGroup")
        for item in equips:
            sed = _sched_add_item_descr(grp, item, "supportEquipDescr")
            qty = ET.SubElement(sed, "reqQuantity")
            qty.set("unitOfMeasure", "EA")
            qty.text = "1"
    else:
        ET.SubElement(se, "noSupportEquips")

    # reqSupplies
    sup_el   = ET.SubElement(pr, "reqSupplies")
    supplies = [x.strip() for txt in task["supplies"]
                for x in re.split(r"[,;]", txt) if x.strip()]
    if supplies:
        grp = ET.SubElement(sup_el, "supplyDescrGroup")
        for item in supplies:
            sd = _sched_add_item_descr(grp, item, "supplyDescr")
            ET.SubElement(sd, "reqQuantity").text = "As required"
    else:
        ET.SubElement(sup_el, "noSupplies")

    # reqSpares
    sp_el  = ET.SubElement(pr, "reqSpares")
    spares = [x.strip() for txt in task["spares"]
              for x in re.split(r"[,;]", txt) if x.strip()]
    if spares:
        grp = ET.SubElement(sp_el, "spareDescrGroup")
        for item in spares:
            sd = _sched_add_item_descr(grp, item, "spareDescr")
            qty = ET.SubElement(sd, "reqQuantity")
            qty.set("unitOfMeasure", "EA")
            qty.text = "1"
    else:
        ET.SubElement(sp_el, "noSpares")

    # reqSafety
    safe_el = ET.SubElement(pr, "reqSafety")
    if task["safety"]:
        for s_txt in task["safety"]:
            sr = ET.SubElement(safe_el, "safetyRqmt")
            ET.SubElement(sr, "simplePara").text = _clean_text(s_txt)
    else:
        ET.SubElement(safe_el, "noSafety")

    # <refs>
    if task["refs"]:
        refs_el = ET.SubElement(td, "refs")
        for r_txt in task["refs"]:
            ep  = ET.SubElement(refs_el, "externalPubRef")
            eri = ET.SubElement(ep, "externalPubRefIdent")
            ET.SubElement(eri, "externalPubCode").text = _clean_text(r_txt)

    # <limit> elements
    for lim_txt in task["limits"]:
        lim_el = ET.SubElement(td, "limit")
        lim_el.set("limitTypeValue", _sched_detect_limit_type(lim_txt))
        val  = _sched_extract_number(lim_txt)
        unit = _sched_detect_unit(lim_txt)
        insp = _sched_detect_insp(lim_txt)
        thr  = ET.SubElement(lim_el, "threshold")
        thr.set("thresholdUnitOfMeasure", unit)
        thr.set("thresholdType", "interval")
        ET.SubElement(thr, "thresholdValue").text = val
        tol_m = re.search(r"[±+\-]\s*(\d+(?:[.,]\d+)?)", lim_txt)
        if tol_m:
            tol = ET.SubElement(thr, "tolerance")
            tol.set("toleranceType", "plusorminus")
            tol.set("toleranceValue", tol_m.group(1))
        ET.SubElement(lim_el, "inspectionType").set("inspectionTypeCategory", insp)


def _build_sched(parent: ET.Element, elements: List[Dict]) -> None:
    """Build S1000D schedul.xsd <maintPlanning> content from flat elements."""
    tasks = _parse_sched_tasks(elements)
    mp    = ET.SubElement(parent, "maintPlanning")
    if not tasks:
        td = ET.SubElement(mp, "taskDefinition")
        td.set("taskIdent",       "001")
        td.set("worthinessLimit", "recommended")
        td.set("reducedMaint",    "No")
        t_el = ET.SubElement(td, "task")
        d_el = ET.SubElement(t_el, "taskDescr")
        ET.SubElement(d_el, "simplePara").text = "Scheduled maintenance task"
        return
    for task in tasks:
        _build_sched_task_xml(mp, task)


def _build_generic(parent: ET.Element, dm_type: str, elements: List[Dict]):
    """Generic content wrapper for less-common schema types."""
    container = ET.SubElement(parent, dm_type + "Content")
    for el in elements:
        raw = el.get("content", "") or ""
        ct = _clean_text(el.get("content", "") or "")
        nl = el.get("native_label", "")
        if not ct:
            continue
        if nl in ("doc_title", "paragraph_title"):
            lp = ET.SubElement(container, "levelledPara")
            ET.SubElement(lp, "title").text = ct
        elif nl == "table":
            _simple_table(container, raw)
        else:
            ET.SubElement(container, "para").text = ct


def _pretty_xml(root: ET.Element) -> str:
    raw   = ET.tostring(root, encoding="unicode")
    dom   = minidom.parseString(raw)
    lines = [l for l in dom.toprettyxml(indent="  ").splitlines() if l.strip()]
    lines[0] = '<?xml version="1.0" encoding="UTF-8"?>'
    lines.insert(1, "<!DOCTYPE dmodule>")
    return "\n".join(lines)


def _split_heading_level(el: Dict) -> Optional[int]:
    """Return the AsciiDoc heading level the element will render as, if any."""
    raw = (el.get("content") or "").strip()
    nl = (el.get("native_label") or "").strip()
    if not raw:
        return None
    if nl == "doc_title":
        return 1
    m = _HEADING_RE.match(raw)
    if m:
        return len(m.group(1)) + 1
    if nl == "paragraph_title":
        return 2
    return None


def _slugify_title(text: str) -> str:
    text = _clean_text(text).lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    return text or "module"


def _safe_output_filename(name: str, max_stem_len: int = 96) -> str:
    """Return a filesystem-safe filename, shortening long stems with a stable hash.

    Windows can report Errno 2 for very long paths; keep filenames bounded and ASCII-safe.
    """
    base, ext = os.path.splitext(name or "")
    safe_base = re.sub(r"[^A-Za-z0-9._-]+", "-", base).strip("-._") or "module"
    if len(safe_base) <= max_stem_len:
        return safe_base + ext

    digest = hashlib.sha1((name or "").encode("utf-8", errors="ignore")).hexdigest()[:10]
    keep = max(16, max_stem_len - len(digest) - 1)
    return f"{safe_base[:keep]}_{digest}{ext}"


def _split_elements_into_modules(elements: List[Dict], fallback_title: str) -> List[Dict]:
    """Split a flat element list into DMC-sized modules using heading structure.

    Rule of thumb:
    - If there are any `==`-level headings, each `==` starts a new DMC.
    - Otherwise, if there are any `===`-level headings, each `===` starts a new DMC.
    - If neither exists, keep the content as a single DMC.
    """
    flat = [el for el in elements if el.get("content")]
    if not flat:
        return []

    heading_levels = [lvl for lvl in (_split_heading_level(el) for el in flat) if lvl is not None]
    doc_title_count = sum(1 for lvl in heading_levels if lvl == 1)
    split_level = 2 if 2 in heading_levels else 1 if doc_title_count > 1 else 3 if 3 in heading_levels else None

    if split_level is None:
        title = fallback_title or _clean_text(flat[0].get("content", "") or "Untitled Data Module")
        return [{
            "index": 0,
            "title": title,
            "slug": _slugify_title(title),
            "elements": flat,
            "split_level": None,
        }]

    modules: List[List[Dict]] = []
    current: List[Dict] = []
    for el in flat:
        level = _split_heading_level(el)
        if level == split_level and current:
            modules.append(current)
            current = [el]
            continue
        if level == split_level and not current:
            current = [el]
            continue
        current.append(el)

    if current:
        modules.append(current)

    split_modules: List[Dict] = []
    for idx, module_elements in enumerate(modules):
        module_title = fallback_title
        for el in module_elements:
            level = _split_heading_level(el)
            if level == split_level:
                raw = (el.get("content") or "").strip()
                m = _HEADING_RE.match(raw)
                module_title = _clean_text(m.group(2) if m else raw)
                break
        if not module_title:
            module_title = _clean_text(module_elements[0].get("content", "") or fallback_title or "Untitled Data Module")
        split_modules.append({
            "index": idx,
            "title": module_title,
            "slug": _slugify_title(module_title),
            "elements": module_elements,
            "split_level": split_level,
        })

    return split_modules


def _split_elements_for_adoc(elements: List[Dict], fallback_title: str) -> List[Dict]:
    """Split content for ADOC so every heading starts a new DMC segment.

    Behavior requested by user:
    - First heading starts module 1
    - Next heading starts module 2
    - Each module runs until just before the next heading
    """
    # Keep figure elements even when content is empty (image blocks have no text content)
    flat = [el for el in elements if el.get("content") or (el.get("native_label") == "figure" and el.get("image_path"))]
    if not flat:
        return []

    modules: List[List[Dict]] = []
    current: List[Dict] = []

    for el in flat:
        is_heading = _split_heading_level(el) is not None
        if is_heading and current:
            modules.append(current)
            current = [el]
        else:
            current.append(el)

    if current:
        modules.append(current)

    out: List[Dict] = []
    for idx, module_elements in enumerate(modules):
        module_title = fallback_title
        for el in module_elements:
            if _split_heading_level(el) is not None:
                raw = (el.get("content") or "").strip()
                m = _HEADING_RE.match(raw)
                module_title = _clean_text(m.group(2) if m else raw)
                break
        if not module_title:
            module_title = _clean_text(module_elements[0].get("content", "") or fallback_title or "Untitled Data Module")
        out.append({
            "index": idx,
            "title": module_title,
            "slug": _slugify_title(module_title),
            "elements": module_elements,
            "split_level": "any-heading",
        })
    return out


def _module_dm_variant(module_index: int) -> str:
    """Return a stable one-letter variant code for split DMC modules."""
    return chr(ord("A") + (module_index % 26))


def page_to_xml(elements: List[Dict], dm_code: str, title: str,
                dm_type: str, page_idx: int, dm_variant: Optional[str] = None) -> str:
    """Build a single S1000D Issue 4.2 XML data module for one page."""
    # For auto mode: derive type from semantic majority
    effective_type = dm_type
    if dm_type in ("auto", ""):
        dominant = _dominant_type(elements)
        effective_type = dominant  # proced or descript

    # Map schema aliases to correct XSD filenames
    xsd_type = effective_type
    if xsd_type == "procedure":
        xsd_type = "proced"
    elif xsd_type == "sched":
        xsd_type = "schedul"

    root = ET.Element("dmodule")
    root.set("xmlns:dc",    "http://www.purl.org/dc/elements/1.1/")
    root.set("xmlns:rdf",   "http://www.w3.org/1999/02/22-rdf-syntax-ns#")
    root.set("xmlns:xlink", "http://www.w3.org/1999/xlink")
    root.set("xmlns:xsi",   "http://www.w3.org/2001/XMLSchema-instance")
    root.set("xsi:noNamespaceSchemaLocation",
             S1000D_42_XSD.format(dm_type=xsd_type))

    _make_ident(root, dm_code, title, effective_type, str(page_idx + 1).zfill(3), dm_variant)
    content_el = ET.SubElement(root, "content")

    if effective_type in ("descript",):
        _build_description(content_el, elements)
    elif effective_type in ("procedure", "proced"):
        _build_procedure(content_el, elements)
    elif effective_type == "fault":
        _build_fault(content_el, elements)
    elif effective_type == "sched":
        _build_sched(content_el, elements)
    else:
        _build_generic(content_el, effective_type, elements)

    return _pretty_xml(root)


# ────────────────────────────────────────────────────────────────────────────
# ASCIIDOC GENERATION
# ────────────────────────────────────────────────────────────────────────────

# Lua filter content written alongside adoc files
_LUA_CONTENT = r"""
function Emph(el)
  return pandoc.RawInline("asciidoc", "<<" .. pandoc.utils.stringify(el.content) .. ">>")
end
function Strong(el)
  return pandoc.RawInline("asciidoc", "[[" .. pandoc.utils.stringify(el.content) .. "]]")
end
"""

_ADOC_HEADER_TEMPLATE = """\
:dmc: {dmc}
:dm-type: {dm_type}
:dm-title: {title}
:tech-name: {tech_name}
:revdate: 2026-04-04
:issue-number: 001
:in-work: 00
:lang: en
:security-classification: 01
:responsible-partner-company: EaseYourWork
:enterprise-code-rpc: 8910X
:originator-enterprise: EaseYourWork
:enterprise-code-originator: 8910X
:applicability: All applicable units and serial numbers.
:brex-dmc: DMC-S1000D-H-041-1-0-0301-00-A-022-A-D
:reason-for-update: Initial draft.

"""


def elements_to_adoc(pages: List[List[Dict]], dm_code: str,
                     dm_type: str, title: str) -> str:
    """Convert extracted elements directly to AsciiDoc (no pandoc needed)."""
    lines: List[str] = [_ADOC_HEADER_TEMPLATE.format(
        dmc=dm_code, dm_type=dm_type, title=title, tech_name=title
    )]
    for page in pages:
        for el in page:
            raw_content = el.get("content", "") or ""
            nl          = el.get("native_label", "") or ""
            if not raw_content.strip() and not (nl == "figure" and el.get("image_path")):
                continue

            if nl == "doc_title":
                lines.append(f"= {_clean_text(raw_content)}\n")
            elif nl == "paragraph_title":
                m = _HEADING_RE.match(raw_content.strip())
                if m:
                    depth = len(m.group(1))
                    lines.append("=" * (depth + 1) + " " + m.group(2).strip() + "\n")
                else:
                    lines.append(f"== {_clean_text(raw_content)}\n")
            elif nl == "table":
                table_block = _table_element_to_adoc_block(el)
                if table_block:
                    lines.append(table_block + "\n")
            elif nl == "figure":
                img = _adoc_image_ref(el.get("image_path", ""))
                lines.append(f"image::{img}[]\n")
            elif nl == "link":
                uri = (el.get("uri") or "").strip()
                link_text = _clean_link_text(el.get("link_text") or raw_content)
                if not uri:
                    # Fall back to markdown pattern inside raw content.
                    m = MD_LINK_RE.search(raw_content)
                    if m:
                        uri = m.group(2)
                        link_text = _clean_link_text(m.group(1))
                if uri:
                    lines.append(f"link:{uri}[{link_text}]\n")
                else:
                    lines.append(_text_to_adoc_links(raw_content) + "\n")
            else:
                # Detect lists in content
                content_lines = raw_content.splitlines()
                in_list = False
                for cl in content_lines:
                    cs = cl.strip()
                    if not cs:
                        if in_list:
                            lines.append("")
                            in_list = False
                        continue
                    nm = _NUM_RE.match(cs)
                    bm = _BULLET_RE.match(cs)
                    if nm:
                        lines.append(f". {nm.group(1)}")
                        in_list = True
                    elif bm:
                        lines.append(f"* {bm.group(1)}")
                        in_list = True
                    else:
                        if in_list:
                            lines.append("")
                            in_list = False
                        lines.append(_text_to_adoc_links(cs))
                lines.append("")
    return "\n".join(lines)


def _adoc_image_ref(img: str) -> str:
    """Figure images are saved to 04_adoc/images/; reference bare filenames there."""
    if not img or "/" in img or "\\" in img or ":" in img:
        return img
    return f"images/{img}"


def _md_image_ref(img: str) -> str:
    """Markdown lives in 05_markdown/; images sit in ../04_adoc/images/."""
    if not img or "/" in img or "\\" in img or ":" in img:
        return img
    return f"../04_adoc/images/{img}"


def elements_to_adoc_body(pages: List[List[Dict]]) -> str:
    """Generate body-only AsciiDoc (without dynamic header attributes)."""
    lines: List[str] = []
    for page in pages:
        for el in page:
            raw_content = el.get("content", "") or ""
            nl = el.get("native_label", "") or ""
            if not raw_content.strip() and not (nl == "figure" and el.get("image_path")):
                continue

            if nl == "doc_title":
                lines.append(f"= {_clean_text(raw_content)}\n")
            elif nl == "paragraph_title":
                m = _HEADING_RE.match(raw_content.strip())
                if m:
                    depth = len(m.group(1))
                    lines.append("=" * (depth + 1) + " " + m.group(2).strip() + "\n")
                else:
                    lines.append(f"== {_clean_text(raw_content)}\n")
            elif nl == "table":
                table_block = _table_element_to_adoc_block(el)
                if table_block:
                    lines.append(table_block + "\n")
            elif nl == "figure":
                img = _adoc_image_ref(el.get("image_path", ""))
                lines.append(f"image::{img}[]\n")
            elif nl == "link":
                uri = (el.get("uri") or "").strip()
                link_text = _clean_link_text(el.get("link_text") or raw_content)
                if uri:
                    lines.append(f"link:{uri}[{link_text}]\n")
                else:
                    lines.append(_text_to_adoc_links(raw_content) + "\n")
            else:
                content_lines = raw_content.splitlines()
                in_list = False
                for cl in content_lines:
                    cs = cl.strip()
                    if not cs:
                        if in_list:
                            lines.append("")
                            in_list = False
                        continue
                    nm = _NUM_RE.match(cs)
                    bm = _BULLET_RE.match(cs)
                    if nm:
                        lines.append(f". {nm.group(1)}")
                        in_list = True
                    elif bm:
                        lines.append(f"* {bm.group(1)}")
                        in_list = True
                    else:
                        if in_list:
                            lines.append("")
                            in_list = False
                        lines.append(_text_to_adoc_links(cs))
                lines.append("")
    return "\n".join(lines).strip() + "\n"


def elements_to_markdown(pages: List[List[Dict]], title: str) -> str:
    """Generate Markdown output from normalized elements."""
    lines: List[str] = [f"# {title}", ""]

    for page in pages:
        for el in page:
            raw_content = (el.get("content") or "").strip()
            nl = (el.get("native_label") or "").strip()
            if not raw_content:
                continue

            if nl == "doc_title":
                lines.append(f"# {_clean_text(raw_content)}")
                lines.append("")
            elif nl == "paragraph_title":
                m = _HEADING_RE.match(raw_content)
                if m:
                    depth = min(6, len(m.group(1)) + 1)
                    lines.append("#" * depth + " " + _clean_text(m.group(2)))
                else:
                    lines.append(f"## {_clean_text(raw_content)}")
                lines.append("")
            elif nl == "table":
                md_block = _table_element_to_markdown_block(el)
                if md_block:
                    lines.extend(md_block.rstrip("\n").splitlines())
                lines.append("")
            elif nl == "figure":
                img = _md_image_ref(el.get("image_path", ""))
                lines.append(f"![image]({img})")
                lines.append("")
            elif nl == "link":
                uri = (el.get("uri") or "").strip()
                link_text = _clean_link_text(el.get("link_text") or raw_content)
                if uri:
                    lines.append(f"[{link_text}]({uri})")
                else:
                    lines.append(raw_content)
                lines.append("")
            else:
                for cl in raw_content.splitlines():
                    cs = cl.strip()
                    if not cs:
                        lines.append("")
                        continue
                    nm = _NUM_RE.match(cs)
                    bm = _BULLET_RE.match(cs)
                    if nm:
                        lines.append(f"1. {nm.group(1)}")
                    elif bm:
                        lines.append(f"- {bm.group(1)}")
                    else:
                        lines.append(cs)
                lines.append("")

    return "\n".join(lines).strip() + "\n"


def _scale_bbox_if_needed(bbox: List[float], page_rect) -> Optional[Tuple[float, float, float, float]]:
    """Normalize bbox coordinates to page units when OCR returns relative values."""
    if not isinstance(bbox, (list, tuple)) or len(bbox) != 4:
        return None
    try:
        x0, y0, x1, y1 = [float(v) for v in bbox]
    except Exception:
        return None

    max_v = max(abs(x0), abs(y0), abs(x1), abs(y1))
    w = float(page_rect.width)
    h = float(page_rect.height)

    # 0..1 normalized
    if max_v <= 1.0:
        x0, x1 = x0 * w, x1 * w
        y0, y1 = y0 * h, y1 * h
    # 0..100 percentage-style
    elif max_v <= 100.0 and w > 200 and h > 200:
        x0, x1 = (x0 / 100.0) * w, (x1 / 100.0) * w
        y0, y1 = (y0 / 100.0) * h, (y1 / 100.0) * h

    # Ensure valid rect orientation
    lx, rx = min(x0, x1), max(x0, x1)
    ty, by = min(y0, y1), max(y0, y1)
    return (lx, ty, rx, by)


def export_assets(src_path: Path, pages: List[List[Dict]], out_root: Path, log) -> None:
    """Export imgs and layout_vis assets into package output."""
    assets_dir = out_root / "06_assets"
    imgs_dir = assets_dir / "imgs"
    vis_dir = assets_dir / "layout_vis"
    raw_vis_dir = assets_dir / "raw_layout_vis"
    imgs_dir.mkdir(parents=True, exist_ok=True)
    vis_dir.mkdir(parents=True, exist_ok=True)
    raw_vis_dir.mkdir(parents=True, exist_ok=True)

    ext = src_path.suffix.lower()

    if ext == ".pdf" and HAS_PYMUPDF:
        try:
            doc = _fitz.open(str(src_path))

            def _label_color(name: str) -> Tuple[float, float, float]:
                n = (name or "").lower()
                if "title" in n or "header" in n:
                    return (0.1, 0.8, 0.4)
                if "image" in n or "figure" in n:
                    return (0.4, 0.9, 0.1)
                if "table" in n:
                    return (0.9, 0.6, 0.1)
                if "footer" in n:
                    return (0.8, 0.8, 0.1)
                if "link" in n:
                    return (0.1, 0.6, 0.9)
                return (0.1, 0.45, 0.9)

            for i in range(len(doc)):
                base = f"{src_path.stem}_page{i+1:04d}.png"
                page = doc[i]

                # Raw page image
                pix = page.get_pixmap(matrix=_fitz.Matrix(2, 2), alpha=False)
                pix.save(str(imgs_dir / base))

                # Build a clean layout visualization (normalized result boxes).
                clean_doc = _fitz.open()
                clean_doc.insert_pdf(doc, from_page=i, to_page=i)
                clean_page = clean_doc[0]
                page_elems = pages[i] if i < len(pages) and isinstance(pages[i], list) else []
                for el in page_elems:
                    bbox = el.get("bbox_2d")
                    scaled = _scale_bbox_if_needed(bbox, clean_page.rect)
                    if not scaled:
                        continue
                    try:
                        clean_page.draw_rect(_fitz.Rect(*scaled), color=(1, 0, 0), width=0.8)
                    except Exception:
                        continue

                vis = clean_page.get_pixmap(matrix=_fitz.Matrix(2, 2), alpha=False)
                vis.save(str(vis_dir / base))
                clean_doc.close()

                # Build GLM-like raw visualization with class color + label/confidence.
                raw_doc = _fitz.open()
                raw_doc.insert_pdf(doc, from_page=i, to_page=i)
                raw_page = raw_doc[0]
                for el in page_elems:
                    bbox = el.get("bbox_2d")
                    scaled = _scale_bbox_if_needed(bbox, raw_page.rect)
                    if not scaled:
                        continue
                    label = str(el.get("label") or el.get("native_label") or "text")
                    conf = el.get("score")
                    if conf is None:
                        conf = el.get("confidence")
                    color = _label_color(label)
                    try:
                        r = _fitz.Rect(*scaled)
                        raw_page.draw_rect(r, color=color, width=1.2)
                        caption = label
                        if isinstance(conf, (int, float)):
                            caption = f"{label} {float(conf):.2f}"
                        text_point = _fitz.Point(r.x0, max(6, r.y0 - 2))
                        raw_page.insert_text(text_point, caption, fontsize=7, color=color)
                    except Exception:
                        continue

                raw_vis = raw_page.get_pixmap(matrix=_fitz.Matrix(2, 2), alpha=False)
                raw_vis.save(str(raw_vis_dir / base))
                raw_doc.close()
            doc.close()
            log(f"  → {imgs_dir} ({len(doc)} image(s))")
            log(f"  → {vis_dir} ({len(doc)} visualization(s))")
            log(f"  → {raw_vis_dir} ({len(doc)} raw visualization(s))")
            return
        except Exception as e:
            log(f"  Asset export failed for PDF rendering: {e}")

    # For image inputs, copy source image as both raw and layout_vis placeholder.
    if ext in IMAGE_EXTS:
        try:
            dst_img = imgs_dir / src_path.name
            dst_vis = vis_dir / src_path.name
            dst_raw = raw_vis_dir / src_path.name
            shutil.copy2(src_path, dst_img)
            shutil.copy2(src_path, dst_vis)
            shutil.copy2(src_path, dst_raw)
            log(f"  → {imgs_dir} (1 image)")
            log(f"  → {vis_dir} (1 visualization placeholder)")
            log(f"  → {raw_vis_dir} (1 raw visualization placeholder)")
            return
        except Exception as e:
            log(f"  Asset export failed for image input: {e}")

    # Non-PDF/text sources have no natural page render or layout boxes.
    log("  Assets skipped (imgs/layout_vis/raw_layout_vis available for PDF/image sources).")


def _set_or_add_adoc_attr(template_text: str, attr_name: str, attr_value: str) -> str:
    """Set an AsciiDoc attribute in template text, adding it if missing."""
    pattern = re.compile(rf"(?m)^:{re.escape(attr_name)}:\s*.*$")
    line = f":{attr_name}: {attr_value}"
    if pattern.search(template_text):
        return pattern.sub(line, template_text)

    lines = template_text.splitlines()
    insert_at = 0
    if lines and lines[0].startswith("="):
        insert_at = 1
    lines.insert(insert_at, line)
    return "\n".join(lines) + ("\n" if template_text.endswith("\n") else "")


def _set_adoc_doc_title(template_text: str, title: str) -> str:
    """Set top-level AsciiDoc document title line (`= ...`)."""
    lines = template_text.splitlines()
    if not lines:
        return f"= {title}\n"
    if lines[0].startswith("="):
        lines[0] = f"= {title}"
    else:
        lines.insert(0, f"= {title}")
    return "\n".join(lines) + ("\n" if template_text.endswith("\n") else "")


def _sanitize_procedure_template(text: str) -> str:
    """Remove demo/sample payload from procedural template while keeping scaffold."""
    cleaned = text

    # Remove optional global applicability listing block.
    cleaned = re.sub(
        r"(?ms)^\[listing\.global_applicability_definition\]\s*\n----\s*\n.*?\n----\s*\n",
        "",
        cleaned,
    )

    # Remove sample content in required conditions.
    cleaned = re.sub(r"(?m)^\[pmref=.*$\n", "", cleaned)
    cleaned = re.sub(r"(?m)^\[dmref=.*$\n", "", cleaned)
    cleaned = re.sub(r"(?m)^\* Ensure the General Safety Publication.*$\n?", "", cleaned)
    cleaned = re.sub(r"(?m)^\* Verify that the system power is OFF.*$\n?", "", cleaned)
    cleaned = re.sub(r"(?m)^\* The work area must be clear.*$\n?", "", cleaned)

    # Remove sample technical-info list items.
    cleaned = re.sub(r"(?m)^\* DMC-S1KDTOOLS-.*$\n?", "", cleaned)

    # Remove sample safety warning.
    cleaned = re.sub(r"(?m)^\[WARNING\]\s*$\n?", "", cleaned)
    cleaned = re.sub(r"(?m)^Ensure all safety pins are installed.*$\n?", "", cleaned)

    # Remove sample procedural body between the section starter and closeout block.
    cleaned = re.sub(
        r"(?ms)^=== PURPOSE\n.*?(?=^// Closeout Requirements Section$)",
        "",
        cleaned,
    )

    # Remove sample closeout bullet/id payload.
    cleaned = re.sub(r"(?m)^\[id=\"cc_panel_close\"\]\s*$\n?", "", cleaned)
    cleaned = re.sub(r"(?m)^\[id=\"cc_tools_clear\"\]\s*$\n?", "", cleaned)
    cleaned = re.sub(r"(?m)^\[id=\"cc_logbook_update\"\]\s*$\n?", "", cleaned)
    cleaned = re.sub(r"(?m)^\* Close and securely fasten.*$\n?", "", cleaned)
    cleaned = re.sub(r"(?m)^\* Remove all tools, test equipment.*$\n?", "", cleaned)
    cleaned = re.sub(r"(?m)^\* Update the equipment maintenance logbook.*$\n?", "", cleaned)
    cleaned = re.sub(r"(?ms)^\[#logbook_entry_target\]\n=== Logbook Entry Details\n\. Log the Serial Number and Calibration Result\.\n?", "", cleaned)

    # Normalize extra blank lines.
    cleaned = re.sub(r"\n{4,}", "\n\n\n", cleaned)
    return cleaned


def _sanitize_sched_template(text: str) -> str:
    """Remove demo/sample payload from scheduled-maintenance template scaffold."""
    cleaned = text or ""

    # Drop example applicability blocks from the default template.
    cleaned = re.sub(
        r"(?ms)^\[#app-variant-std.*?(?=^// ========== PRELIMINARY REQUIREMENTS ==========)",
        "",
        cleaned,
    )

    # Remove canned preliminary requirement bullets to avoid irrelevant filler.
    prelim_samples = [
        r"\* Equipment is de-energised and isolated\.",
        r"\* Work area is clear, well-lit and safe\.",
        r"\* Maintainer \(Skill Level 1 or as specified per task\)",
        r"\* Safety Precautions Data Module",
        r"\* General-purpose toolset",
        r"\* Inspection lamp / torch",
        r"\* Cleaning cloth, lint-free",
        r"\* Approved cleaning solvent",
        r"\* Refer to individual task definitions below\.",
        r"\* Wear appropriate PPE \(gloves, eye protection\)\.",
        r"\* Follow local safety regulations\.",
    ]
    for pat in prelim_samples:
        cleaned = re.sub(rf"(?m)^\s*{pat}\s*$\n?", "", cleaned)

    # Remove canned closeout bullets.
    closeout_samples = [
        r"\* Confirm all tools and materials accounted for and removed\.",
        r"\* Confirm all access panels and covers secured\.",
        r"\* Confirm maintenance records updated\.",
        r"\* Complete authorised dealer maintenance certificate\.",
        r"\* Obtain customer acknowledgement signature\.",
        r"\* File original documentation per dealer records procedure\.",
        r"\* Verify that all performance parameters are within specification\.",
        r"\* Record diagnostic readings in the maintenance logbook\.",
        r"\* Provide equipment owner with performance summary report\.",
    ]
    for pat in closeout_samples:
        cleaned = re.sub(rf"(?m)^\s*{pat}\s*$\n?", "", cleaned)

    # Remove empty applicability placeholders left behind by sample removal.
    cleaned = re.sub(r"(?ms)^\[#app-[^\]]+\].*?(?=^\s*$)", "", cleaned)

    cleaned = re.sub(r"\n{4,}", "\n\n\n", cleaned)
    return cleaned


def _build_sched_adoc_body(tasks: List[Dict]) -> str:
    """Generate AsciiDoc task body for scheduled maintenance data modules."""
    _UNIT_NAMES: Dict[str, str] = {
        "th01": "hour(s)",        "th02": "day(s)",
        "th03": "month(s)",       "th04": "quarter(s)",
        "th06": "year(s)",        "th07": "cycle(s)",
        "th09": "flight hour(s)", "th10": "km",
        "th13": "mile(s)",
    }
    lines: List[str] = []

    for task in tasks:
        ident      = task.get("task_ident", "001")
        desc       = _clean_text(task.get("description", "") or f"Task {ident}")
        applic     = task.get("applic_ref")
        worthiness = task.get("worthiness", "recommended").title()
        reduced    = task.get("reduced_maint", "No")
        skill      = task.get("skill_type", "st01")

        if applic:
            lines.append(f"[applic_ref={applic}]")
        lines.append(f"=== Task {ident}: {desc}")
        lines.append("")

        show_note = (
            worthiness.lower() != "recommended" or
            str(reduced).strip().lower() not in ("no", "false", "0", "") or
            str(skill).strip().lower() not in ("st01", "")
        )
        if show_note:
            lines.append(
                f"NOTE: Worthiness: {worthiness} | "
                f"Reduced Maintenance: {reduced} | Skill Type: {skill}"
            )
            lines.append("")

        conds   = task.get("conditions", [])
        persons = task.get("persons",    [])
        equip   = task.get("equipment",  [])
        supps   = task.get("supplies",   [])
        spares  = task.get("spares",     [])
        safety  = task.get("safety",     [])
        limits  = task.get("limits",     [])
        refs    = task.get("refs",        [])
        extra   = task.get("extra",       [])

        if conds:
            lines.append("==== Required Conditions")
            for c in conds:
                lines.append(f"* {c}")
            lines.append("")

        if persons:
            lines.append("==== Required Persons")
            for p in persons:
                lines.append(f"* {p}")
            lines.append("")

        if equip:
            lines.append("==== Required Support Equipment")
            for e in equip:
                for item in re.split(r"[,;]", e):
                    item = item.strip()
                    if item:
                        lines.append(f"* {item}")
            lines.append("")

        if supps:
            lines.append("==== Required Supplies")
            for s in supps:
                for item in re.split(r"[,;]", s):
                    item = item.strip()
                    if item:
                        lines.append(f"* {item}")
            lines.append("")

        if spares:
            lines.append("==== Required Spares")
            for s in spares:
                for item in re.split(r"[,;]", s):
                    item = item.strip()
                    if item:
                        lines.append(f"* {item}")
            lines.append("")

        if safety:
            lines.append("==== Safety Requirements")
            for s in safety:
                lines.append(f"WARNING: {s}")
            lines.append("")

        if limits:
            lines.append("==== Time Limits / Intervals")
            for lim in limits:
                val   = _sched_extract_number(lim)
                unit  = _sched_detect_unit(lim)
                insp  = _sched_detect_insp(lim)
                lt    = _sched_detect_limit_type(lim)
                uname = _UNIT_NAMES.get(unit, unit)
                lines.append(
                    f"* *{insp}:* {val} {uname}"
                    f" (Limit type: `{lt}`, Unit code: `{unit}`)"
                )
            lines.append("")

        if refs:
            lines.append("==== References")
            for r in refs:
                lines.append(f"* {r}")
            lines.append("")

        if extra:
            lines.append("==== Additional Notes")
            for e in extra:
                lines.append(e)
            lines.append("")

        lines.append("'''")
        lines.append("")

    return "\n".join(lines)


def _collapse_tables_to_headers_only(adoc_text: str) -> str:
    """Keep only table headers; drop data rows for template-fill readability."""
    lines = adoc_text.splitlines()
    out: List[str] = []
    i = 0
    while i < len(lines):
        if lines[i].strip() == "|===":
            tbl_block = [lines[i]]
            i += 1
            while i < len(lines) and lines[i].strip() != "|===":
                tbl_block.append(lines[i])
                i += 1
            if i < len(lines):
                tbl_block.append(lines[i])

            header_row = ""
            for row in tbl_block[1:-1]:
                if row.strip().startswith("|"):
                    header_row = row
                    break

            out.append("|===")
            if header_row:
                out.append(header_row)
            out.append("|===")
            i += 1
            continue

        out.append(lines[i])
        i += 1
    return "\n".join(out)


def _parse_adoc_sections(adoc_text: str) -> Dict[str, List[str]]:
    """Parse heading-based sections from adoc body."""
    sections: Dict[str, List[str]] = {}
    current = ""
    for ln in adoc_text.splitlines():
        m = re.match(r"^(={1,6})\s+(.+)$", ln.strip())
        if m:
            current = m.group(2).strip().lower()
            sections.setdefault(current, [])
            continue
        if current:
            sections[current].append(ln)
    return sections


def _pick_section(sections: Dict[str, List[str]], keys: List[str]) -> str:
    chunks: List[str] = []
    for title, lines in sections.items():
        if any(k in title for k in keys):
            body = "\n".join(lines).strip()
            if body:
                chunks.append(body)
    return "\n\n".join(chunks).strip()


def _replace_heading_body(template_text: str, heading_title: str, replacement: str) -> str:
    """Replace content under a heading with replacement until next heading <= current level."""
    lines = template_text.splitlines()
    head_idx = -1
    head_lvl = 0
    for idx, ln in enumerate(lines):
        m = re.match(r"^(={1,6})\s+(.+)$", ln.strip())
        if not m:
            continue
        if m.group(2).strip().lower() == heading_title.strip().lower():
            head_idx = idx
            head_lvl = len(m.group(1))
            break
    if head_idx < 0:
        return template_text

    end_idx = len(lines)
    for j in range(head_idx + 1, len(lines)):
        m = re.match(r"^(={1,6})\s+(.+)$", lines[j].strip())
        if m and len(m.group(1)) <= head_lvl:
            end_idx = j
            break

    rep_lines = [""] + replacement.strip().splitlines() + [""] if replacement.strip() else [""]
    new_lines = lines[:head_idx + 1] + rep_lines + lines[end_idx:]
    return "\n".join(new_lines)


def _replace_heading_body_any(template_text: str, heading_titles: List[str], replacement: str) -> str:
    """Replace body for the first matching heading title from candidates."""
    updated = template_text
    for h in heading_titles:
        newer = _replace_heading_body(updated, h, replacement)
        if newer != updated:
            return newer
    return updated


def _fill_template_structured(dm_type: str, template_text: str, body_adoc: str) -> Optional[str]:
    """Fill known template sections from extracted content instead of pasting whole body."""
    if dm_type not in ("procedure", "proced", "fault"):
        return None

    prepared = _collapse_tables_to_headers_only(body_adoc)
    sections = _parse_adoc_sections(prepared)
    if not sections:
        # If extracted body has no headings, still use it as core content.
        sections = {"__core__": prepared.splitlines()}

    out = template_text
    # Preliminary requirements (best-effort keyword mapping)
    mapping = [
        (["Required Conditions"], ["required condition", "condition", "precondition"]),
        (["Required Persons"], ["required person", "person", "personnel", "operator", "technician"]),
        (["Required Technical Information"], ["required technical information", "technical information", "reference"]),
        (["Required Support Equipment", "Support Equipment"], ["required support equipment", "support equipment", "equipment", "tool"]),
        (["Required Supplies", "Supplies"], ["required supplies", "supplies", "consumable"]),
        (["Required Spares", "Spares"], ["required spares", "spares", "spare"]),
        (["Required Safety", "Safety Conditions"], ["required safety", "safety", "warning", "caution"]),
    ]
    for heading_candidates, keys in mapping:
        val = _pick_section(sections, keys)
        out = _replace_heading_body_any(out, heading_candidates, val)

    # Main procedure / fault steps: everything not mapped is treated as procedural core.
    mapped_keys = [k for _, ks in mapping for k in ks]
    core_chunks: List[str] = []
    for title, lines in sections.items():
        if any(k in title for k in mapped_keys):
            continue
        body = "\n".join(lines).strip()
        if body:
            if title == "__core__":
                core_chunks.append(body)
            else:
                core_chunks.append(f"=== {title.title()}\n{body}")
    core = "\n\n".join(core_chunks).strip()
    if dm_type in ("procedure", "proced"):
        out = _replace_heading_body(out, "Main Procedure", core)
    elif dm_type == "fault":
        out = _replace_heading_body(out, "Fault Isolation Procedure", core)

    # Keep closeout concise.
    out = _replace_heading_body(out, "Closeout Requirements", "")
    closeout = _pick_section(sections, ["closeout", "after job", "completion"])
    out = _replace_heading_body(out, "Required Conditions After Job Completion", closeout)

    return out


def _norm_text_line(line: str) -> str:
    return " ".join((line or "").strip().split())


def _validate_no_data_loss(source_body: str, candidate_text: str) -> bool:
    """Require every non-empty source body line to remain present in candidate text."""
    if not source_body.strip():
        return True
    cand_norm = "\n".join(_norm_text_line(l) for l in candidate_text.splitlines())
    source_lines = [_norm_text_line(l) for l in source_body.splitlines() if _norm_text_line(l)]
    # Keep this strict: if any line disappears, reject LLM output.
    for line in source_lines:
        if line not in cand_norm:
            return False
    return True


def _is_plausible_template_output(candidate_text: str) -> bool:
    """Reject common LLM wrapper/meta outputs that are not final ADOC docs."""
    lowered = (candidate_text or "").lower()
    bad_markers = [
        "=== template start ===",
        "=== template end ===",
        "=== body start ===",
        "=== body end ===",
        "here is the merged document",
    ]
    if any(m in lowered for m in bad_markers):
        return False
    if ":dmc:" not in lowered:
        return False
    return True


def _ollama_structure_template(
    *,
    template_text: str,
    body_adoc: str,
    dm_type: str,
    title: str,
    log,
) -> Optional[str]:
    """Ask Ollama to place body content into template structure without data loss."""
    if not USE_OLLAMA_TEMPLATE:
        return None
    if dm_type == "sched":
        # Scheduled maintenance uses deterministic template insertion; avoid
        # a slow/optional Ollama round-trip that can block the converter.
        return None
    if not template_text.strip() or not body_adoc.strip():
        return None

    # Skip LLM pass for large documents — sending 100K+ chars would take hours
    # and the deterministic fill already handled the structure above.
    _OLLAMA_TEMPLATE_CHAR_LIMIT = 8000
    if len(body_adoc) > _OLLAMA_TEMPLATE_CHAR_LIMIT:
        return None

    try:
        import requests

        tmpl_snippet = template_text[:4000]
        body_snippet = body_adoc[:_OLLAMA_TEMPLATE_CHAR_LIMIT]
        prompt = (
            "You are an S1000D AsciiDoc formatter.\n"
            "Task: merge BODY content into TEMPLATE structure.\n"
            "Hard rules:\n"
            "1) Output valid AsciiDoc only.\n"
            "2) Preserve ALL BODY lines exactly (no omission, no paraphrase).\n"
            "3) Keep TEMPLATE headings/anchors/attributes intact unless strictly required to insert BODY.\n"
            "4) Do not add explanations.\n"
            "5) Return only the final merged document.\n\n"
            f"DM Type: {dm_type}\n"
            f"Title: {title}\n\n"
            "=== TEMPLATE START ===\n"
            f"{tmpl_snippet}\n"
            "=== TEMPLATE END ===\n\n"
            "=== BODY START ===\n"
            f"{body_snippet}\n"
            "=== BODY END ===\n"
        )

        endpoints: List[Tuple[str, Dict[str, Any], str]] = []
        base = OLLAMA_URL
        if "/api/" in base:
            root = base.split("/api/", 1)[0]
        else:
            root = base.rstrip("/")

        _opts = {"temperature": 0, "top_p": 0.9, "num_predict": 4096}
        endpoints.append((
            base,
            {
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": _opts,
            },
            "generate",
        ))
        endpoints.append((
            root + "/api/chat",
            {
                "model": OLLAMA_MODEL,
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "stream": False,
                "options": _opts,
            },
            "chat",
        ))

        out = ""
        last_err = None
        for url, payload, mode in endpoints:
            try:
                resp = requests.post(url, json=payload, timeout=(10, 60))
                resp.raise_for_status()
                data = resp.json()
                if mode == "generate":
                    out = (data.get("response") or "").strip()
                else:
                    out = ((data.get("message") or {}).get("content") or "").strip()
                if out:
                    break
            except Exception as e:
                last_err = e
                continue

        if not out:
            if last_err:
                raise last_err
            return None
        if not _is_plausible_template_output(out):
            log("  Ollama template pass rejected (non-document wrapper/meta output).")
            return None
        if not _validate_no_data_loss(body_adoc, out):
            log("  Ollama template pass rejected (data-loss check failed); using deterministic template insertion.")
            return None
        log("  Applied Ollama template structuring pass.")
        return out
    except Exception as e:
        log(f"  Ollama template pass unavailable: {e}")
        return None


def apply_template_to_adoc(
    dm_type: str,
    dm_code: str,
    title: str,
    body_adoc: str,
    log,
) -> Optional[str]:
    """Render ADOC using user template for supported schema types."""
    template_name = TEMPLATE_MAP.get(dm_type)
    if not template_name:
        return None

    template_path = Path("Templates") / template_name
    if not template_path.exists():
        log(f"  Template not found: {template_path}. Falling back to default ADOC generation.")
        return None

    try:
        text = template_path.read_text(encoding="utf-8")
    except Exception as e:
        log(f"  Could not read template {template_path}: {e}")
        return None

    today_iso = time.strftime("%Y-%m-%d")

    def _normalize_common_meta(x: str) -> str:
        x = _set_or_add_adoc_attr(x, "issue-number", "001")
        x = _set_or_add_adoc_attr(x, "in-work", "00")
        x = _set_or_add_adoc_attr(x, "issue-date", today_iso)
        x = _set_or_add_adoc_attr(x, "revdate", today_iso)
        x = _set_or_add_adoc_attr(x, "responsible-partner-company", "Your Organization")
        x = _set_or_add_adoc_attr(x, "originator-enterprise", "Your Organization")
        return x

    # First pass: deterministic section-aware fill (fill-in-the-blank style)
    structured_text = _fill_template_structured(dm_type, text, body_adoc)
    if structured_text:
        text = _set_adoc_doc_title(structured_text, title)
        text = _set_or_add_adoc_attr(text, "dmc", dm_code)
        dm_type_attr = "IPD" if dm_type == "illustratedpartscatalog" else dm_type
        text = _set_or_add_adoc_attr(text, "dm-type", dm_type_attr)
        text = _set_or_add_adoc_attr(text, "dm-title", title)
        text = _set_or_add_adoc_attr(text, "tech-name", title)
        text = _set_or_add_adoc_attr(text, "infoName", title)
        text = _normalize_common_meta(text)
        llm_text = _ollama_structure_template(
            template_text=text,
            body_adoc=body_adoc,
            dm_type=dm_type,
            title=title,
            log=log,
        )
        if llm_text:
            text = llm_text
        log(f"  Applied template (structured fill): {template_path}")
        return text

    # ── Sched: direct marker replacement ─────────────────────────────────────
    if dm_type == "sched" and "// SCHED_TASKS_BODY_INSERT" in text:
        text = _sanitize_sched_template(text)
        text = _set_adoc_doc_title(text, title)
        text = _set_or_add_adoc_attr(text, "dmc",      dm_code)
        text = _set_or_add_adoc_attr(text, "dm-type",  dm_type)
        text = _set_or_add_adoc_attr(text, "dm-title", title)
        text = _set_or_add_adoc_attr(text, "tech-name", title.split(" - ")[0] if " - " in title else title)
        text = _set_or_add_adoc_attr(text, "infoName", title)
        text = _normalize_common_meta(text)
        text = text.replace("// SCHED_TASKS_BODY_INSERT", body_adoc.strip())
        log("  Skipping Ollama template pass for scheduled maintenance (deterministic path).")
        log(f"  Applied template: {template_path}")
        return text

    # Synchronize core attributes dynamically while keeping template structure.
    dm_type_attr = "IPD" if dm_type == "illustratedpartscatalog" else dm_type
    text = _set_adoc_doc_title(text, title)
    text = _set_or_add_adoc_attr(text, "dmc", dm_code)
    text = _set_or_add_adoc_attr(text, "dm-type", dm_type_attr)
    text = _set_or_add_adoc_attr(text, "dm-title", title)
    text = _set_or_add_adoc_attr(text, "tech-name", title)
    text = _set_or_add_adoc_attr(text, "infoName", title)
    text = _normalize_common_meta(text)

    if dm_type in ("procedure", "proced"):
        text = _sanitize_procedure_template(text)

    # For procedural templates, insert body directly after the main procedure heading.
    if dm_type in ("procedure", "proced"):
        main_heading = re.search(r"(?m)^==\s+Main\s+Procedure\s*$", text)
        if main_heading:
            insert_idx = main_heading.end()
            text = text[:insert_idx] + "\n\n" + body_adoc.strip() + "\n\n" + text[insert_idx:]
            log(f"  Applied template: {template_path}")
            return text

    insert_patterns = [
        r"(?m)^\[(?:#|\[#)?main_proc_steps\]?\]\s*$",
        r"(?m)^\[(?:#|\[#)?fault_iso_main\]?\]\s*$",
        r"(?m)^\[(?:#|\[#)?sched_tasks\]?\]\s*$",
        r"(?m)^==\s+Catalog\s+Sequence\s+Numbers\s*$",
    ]

    inserted = False
    for pat in insert_patterns:
        m = re.search(pat, text)
        if not m:
            continue
        next_heading = re.search(r"(?m)^==\s+.+$", text[m.end():])
        if next_heading:
            idx = m.end() + next_heading.start()
            text = text[:idx] + "\n\n" + body_adoc.strip() + "\n\n" + text[idx:]
        else:
            text = text.rstrip() + "\n\n" + body_adoc.strip() + "\n"
        inserted = True
        break

    if not inserted:
        text = text.rstrip() + "\n\n== Extracted Content\n\n" + body_adoc.strip() + "\n"

    llm_text = _ollama_structure_template(
        template_text=text,
        body_adoc=body_adoc,
        dm_type=dm_type,
        title=title,
        log=log,
    )
    if llm_text:
        text = llm_text

    log(f"  Applied template: {template_path}")
    return text


def pandoc_to_adoc(src_path: Path, dm_code: str, dm_type: str,
                   title: str, out_path: Path, log) -> bool:
    """Use pandoc to convert DOCX/MD → AsciiDoc and prepend S1000D header."""
    lua_path = out_path.parent / "s1000d_styles.lua"
    lua_path.write_text(_LUA_CONTENT, encoding="utf-8")

    try:
        cmd = ["pandoc", str(src_path),
               "--lua-filter", str(lua_path), "-t", "asciidoc"]
        r = subprocess.run(cmd, capture_output=True, text=True,
                           encoding="utf-8", shell=True)
        if r.returncode != 0:
            log(f"  pandoc stderr: {r.stderr[:300]}")
            return False
        body = r.stdout
        # Clean up Pandoc artefacts
        body = re.sub(r"\+\+(.*?)\+\+", r"\1", body)
        body = re.sub(r"^(\s*image):", r"\1::", body, flags=re.MULTILINE)
        body = re.sub(r"\[arabic(\s*,\s*start=\d+)?\]\s*", "", body, flags=re.MULTILINE)

        header = _ADOC_HEADER_TEMPLATE.format(
            dmc=dm_code, dm_type=dm_type, title=title, tech_name=title
        )
        out_path.write_text(header + body, encoding="utf-8")
        return True
    except Exception as e:
        log(f"  pandoc error: {e}")
        return False


# ────────────────────────────────────────────────────────────────────────────
# PACKAGE BUILDER
# ────────────────────────────────────────────────────────────────────────────

def detect_dm_code(pages: List[List[Dict]]) -> Optional[str]:
    """Try to find a DMC-like code inside the extracted content."""
    pat = re.compile(
        r"(DMC-[A-Z0-9]{2,14}-[A-Z0-9]{1,4}-[0-9A-Z]{2}-[0-9A-Z]{2,4}"
        r"(?:-[0-9A-Z]{1,5}){1,6})",
        re.IGNORECASE
    )
    for page in pages:
        for el in page:
            m = pat.search(el.get("content", "") or "")
            if m:
                return m.group(1).upper()
    return None


def infer_dm_type(src_path: Path, pages: List[List[Dict]], log) -> str:
    """Infer best DM type from filename/content. Falls back to procedure/descript."""
    file_text = src_path.stem.lower()
    content_chunks: List[str] = []
    for page in pages[:3]:
        if not isinstance(page, list):
            continue
        for el in page[:80]:
            content_chunks.append((el.get("content") or "")[:300])
    content_text = "\n".join(content_chunks).lower()
    probe = f"{file_text}\n{content_text}"

    # Scored detection to avoid false positives from single keyword matches.
    scores: Dict[str, int] = {k: 0 for k in S1000D_DM_TYPES if k != "auto"}
    traces: Dict[str, List[str]] = {k: [] for k in scores}

    def bump(dm: str, pts: int, reason: str):
        if dm in scores:
            scores[dm] += pts
            traces[dm].append(f"{reason} (+{pts})")

    # Strong schema-specific phrases
    strong_rules = [
        ("illustratedpartscatalog", ["illustrated parts catalog", "parts catalog (ipc)"]),
        ("appliccrossreftable", ["applicability cross-reference", "applicability cross reference table"]),
        ("condcrossreftable", ["condition cross-reference", "condition cross reference table"]),
        ("partrepository", ["parts repository", "part repository"]),
        ("functionalitem", ["functional item repository", "functional items"]),
        ("comrepository", ["comments repository", "comment repository"]),
        ("frontmatter", ["front matter", "table of contents", "title page"]),
        ("wrngdata", ["wiring data", "wiring diagram", "wire list"]),
        ("techrep", ["technical report", "analysis report", "test report"]),
        # BREX should only match on explicit BREX DM intent, not generic "BREX rules" mentions.
        ("brex", ["business rules exchange", "brex data module", "brex dm"]),
        ("sb", ["service bulletin"]),
        ("pim", ["preliminary information", "preliminary info"]),
        ("chkl", ["checklist", "inspection checklist"]),
        ("crew", ["crew", "operator information"]),
        ("fault", ["fault isolation", "troubleshooting", "fault code"]),
        ("sched", ["scheduled maintenance", "maintenance schedule", "maintenance interval"]),
    ]

    for dm, hints in strong_rules:
        for h in hints:
            if h in probe:
                bump(dm, 4, f"match '{h}'")

    # File-name hints
    name_rules = [
        ("illustratedpartscatalog", ["ipc", "illustrated", "partscatalog"]),
        ("brex", ["brex"]),
        ("fault", ["fault", "fim"]),
        ("sched", ["sched", "schedule"]),
        ("procedure", ["procedure", "proc", "steps", "guide"]),
        ("descript", ["description", "descript", "overview"]),
    ]
    for dm, hints in name_rules:
        for h in hints:
            if h in file_text:
                bump(dm, 2, f"filename contains '{h}'")

    # Procedural bias from step-heavy text
    step_count = len(re.findall(r"\bstep\s+\d+\b", probe, re.IGNORECASE))
    numbered_lines = len(re.findall(r"(?m)^\s*\d+[\.)]\s+", probe))
    if step_count >= 2:
        bump("procedure", 4, f"step patterns x{step_count}")
        bump("proced", 3, f"step patterns x{step_count}")
    if numbered_lines >= 2:
        bump("procedure", 3, f"numbered lines x{numbered_lines}")
        bump("proced", 2, f"numbered lines x{numbered_lines}")

    # Semantic majority still contributes, but with lower weight.
    flat_elements = [el for p in pages if isinstance(p, list) for el in p]
    dominant = _dominant_type(flat_elements)
    if dominant == "proced":
        bump("procedure", 2, "semantic majority proced")
        bump("proced", 2, "semantic majority proced")
    else:
        bump("descript", 2, "semantic majority descript")

    best_dm = max(scores, key=lambda k: scores[k])
    best_score = scores[best_dm]

    # If weak evidence, use semantic fallback.
    if best_score <= 2:
        inferred = "procedure" if dominant == "proced" else "descript"
        log(f"  Auto DM type fallback: {inferred} (weak signal, semantic majority: {dominant})")
        return inferred

    reasons = ", ".join(traces.get(best_dm, [])[:3])
    if reasons:
        log(f"  Auto DM type detected: {best_dm} ({reasons})")
    else:
        log(f"  Auto DM type detected: {best_dm}")
    return best_dm


def build_package(
    src_path: Path,
    out_root: Path,
    dm_type: str,
    force_ocr: bool,
    out_formats: Dict[str, bool],  # raw_json, sem_json, xml, adoc, md, assets, glm_default
    log,
    stop_event=None,       # threading.Event – checked between major steps
    progress_cb=None,      # callable(step:int, total:int, name:str)
):
    """
    Full conversion pipeline for one file.
    Creates:
        <out_root>/01_raw_json/<stem>_raw.json
        <out_root>/02_semantic_json/<stem>_semantic.json
        <out_root>/03_s1000d_xml/<stem>_page00.xml  (one per page)
        <out_root>/04_adoc/<stem>.adoc
    """
    stem = src_path.stem
    ext  = src_path.suffix.lower()

    def _prog(step: int, name: str):
        if progress_cb:
            try:
                progress_cb(step, 4, name)
            except Exception:
                pass

    def _stopped() -> bool:
        return stop_event is not None and stop_event.is_set()

    # ── Step 1 : Extract ──────────────────────────────────────────────────
    _prog(1, "Extracting content")
    log(f"\n{'─'*60}")
    log(f"[1/4] Extracting: {src_path.name}")
    glm_native_out_dir = None
    docx_media_dir = None
    ocr_images_dir = None
    if out_formats.get("glm_default"):
        glm_native_out_dir = out_root / "00_glm_default" / src_path.stem
    if ext == ".docx":
        # Extracted DOCX images live in 04_adoc/images/, referenced as image::images/…
        docx_media_dir = out_root / "04_adoc" / "images"
    # For any OCR path (scanned PDF, standalone image, or image-heavy/force-OCR DOCX),
    # save figure crops under 04_adoc/images/ so image::images/… macros resolve.
    if ext in IMAGE_EXTS or (ext == ".pdf" and (force_ocr or _is_scanned_pdf(src_path))) or (
        ext == ".docx" and (force_ocr or _is_image_heavy_docx(src_path))
    ):
        ocr_images_dir = out_root / "04_adoc" / "images"

    pages = extract_file(src_path, log, force_ocr, glm_native_out_dir, docx_media_dir, ocr_images_dir)
    if not pages:
        log("  ERROR: No content extracted. Skipping file.")
        return False

    # Recover visible URLs/domains from OCR/plain text when source metadata is absent.
    pages, _ = _enrich_pages_with_recovered_links(pages, log)

    dm_code = detect_dm_code(pages) or f"S1000D-A-00-00-0000-00A-040A-A"
    effective_dm_type = infer_dm_type(src_path, pages, log) if dm_type == "auto" else dm_type
    effective_dm_type = _enforce_descript_only(effective_dm_type, log)
    # Use stem as fallback title (first meaningful heading preferred)
    title = stem.replace("_", " ").replace("-", " ").title()
    # Try to pick a better title from first page
    for el in (pages[0] if pages else []):
        nl = el.get("native_label", "")
        if nl in ("doc_title", "paragraph_title"):
            title = _clean_text(el.get("content", "")) or title
            break

    log(
        f"  Pages extracted: {len(pages)} | DM code: {dm_code} | "
        f"DM type: {effective_dm_type} | Title: {title[:60]}"
    )

    if _stopped():
        log("  ⏹ Conversion stopped.")
        return False

    # ── Step 2 : Raw JSON ─────────────────────────────────────────────────
    _prog(2, "Writing JSON")
    if out_formats.get("raw_json"):
        log(f"[2/4] Writing raw JSON...")
        raw_dir = out_root / "01_raw_json"
        raw_dir.mkdir(parents=True, exist_ok=True)
        raw_path = raw_dir / f"{stem}_raw.json"
        raw_path.write_text(json.dumps(pages, indent=2, ensure_ascii=False),
                            encoding="utf-8")
        log(f"  → {raw_path}")

    if _stopped():
        log("  ⏹ Conversion stopped.")
        return False

    # ── Step 3 : Semantic JSON ────────────────────────────────────────────
    _prog(3, "Semantic annotation")
    log(f"[3/4] Semantic annotation...")
    sem_pages = annotate_pages(pages)
    structured_sem = build_structured_semantic(sem_pages)

    if out_formats.get("sem_json"):
        sem_dir = out_root / "02_semantic_json"
        sem_dir.mkdir(parents=True, exist_ok=True)
        sem_path = sem_dir / f"{stem}_semantic.json"
        sem_payload = {
            "pages": sem_pages,
            "structured": structured_sem,
        }
        sem_path.write_text(json.dumps(sem_payload, indent=2, ensure_ascii=False),
                            encoding="utf-8")
        log(f"  → {sem_path}")

    if _stopped():
        log("  ⏹ Conversion stopped.")
        return False

    # Flatten all elements for stats
    total_els = sum(len(p) for p in sem_pages)
    n_proced  = sum(1 for p in sem_pages for el in p
                   if el.get("semantic", {}).get("type") == "proced")
    log(f"  Elements: {total_els}  |  proced: {n_proced}  |  descript: {total_els - n_proced}")

    flat_elements = [el for p in sem_pages for el in p]
    modules = _split_elements_into_modules(flat_elements, title)
    if len(modules) > 1:
        log(f"  Split into {len(modules)} DMC modules using heading structure.")

    # ── Step 4a : S1000D XML ──────────────────────────────────────────────
    _prog(4, "Generating outputs")
    if out_formats.get("xml"):
        log(f"[4a/4] Generating S1000D XML ({effective_dm_type})...")
        xml_dir = out_root / "03_s1000d_xml"
        xml_dir.mkdir(parents=True, exist_ok=True)
        xml_written = 0
        for module in modules:
            module_index = module["index"]
            module_title = module["title"] or title
            module_elements = module["elements"]
            module_variant = _module_dm_variant(module_index) if len(modules) > 1 else None
            module_dm_type = infer_dm_type(src_path, [module_elements], log) if dm_type == "auto" else effective_dm_type
            module_dm_type = _enforce_descript_only(module_dm_type, log)
            xml_str = page_to_xml(
                module_elements,
                dm_code,
                module_title,
                module_dm_type,
                module_index,
                module_variant,
            )
            if len(modules) == 1:
                xml_name = f"{stem}_sched.xml" if module_dm_type == "sched" else f"{stem}.xml"
            else:
                xml_name = f"{stem}_dm{module_index + 1:02d}_{module['slug']}.xml"
            xml_name = _safe_output_filename(xml_name)
            xml_path = xml_dir / xml_name
            xml_path.write_text(xml_str, encoding="utf-8")
            xml_written += 1
        log(f"  → {xml_dir}  ({xml_written} file(s))")

    # ── Step 4b : AsciiDoc ────────────────────────────────────────────────
    if out_formats.get("adoc"):
        log(f"[4b/4] Generating AsciiDoc...")
        adoc_dir = out_root / "04_adoc"
        adoc_dir.mkdir(parents=True, exist_ok=True)
        adoc_modules = _split_elements_for_adoc(flat_elements, title)
        if len(adoc_modules) > 1:
            log(f"  ADOC split into {len(adoc_modules)} DMC modules (new heading starts new module).")
        adoc_written = 0
        for module in adoc_modules:
            module_index = module["index"]
            module_title = module["title"] or title
            module_elements = module["elements"]
            module_dm_code = dm_code
            module_dm_type = infer_dm_type(src_path, [module_elements], log) if dm_type == "auto" else effective_dm_type
            module_dm_type = _enforce_descript_only(module_dm_type, log)

            if module_dm_type == "sched":
                has_table = any((el.get("native_label") or "") == "table" for el in module_elements)
                # If a schedule matrix/table exists, preserve it verbatim in ADOC.
                # Task synthesis is only used for text-only schedule sources.
                if has_table:
                    body_adoc = elements_to_adoc_body([module_elements])
                else:
                    body_adoc = _build_sched_adoc_body(_parse_sched_tasks(module_elements))
            else:
                body_adoc = elements_to_adoc_body([module_elements])

            templated_adoc = apply_template_to_adoc(
                module_dm_type, module_dm_code, module_title, body_adoc, log
            )

            if templated_adoc is not None:
                adoc_text = templated_adoc
            else:
                adoc_text = elements_to_adoc([module_elements], module_dm_code, module_dm_type, module_title)

            if len(adoc_modules) == 1:
                adoc_name = f"{stem}.adoc"
            else:
                adoc_name = f"{stem}_dm{module_index + 1:02d}_{module['slug']}.adoc"
            adoc_name = _safe_output_filename(adoc_name)

            adoc_path = adoc_dir / adoc_name
            adoc_path.write_text(adoc_text, encoding="utf-8")
            adoc_written += 1

        log(f"  → {adoc_dir}  ({adoc_written} file(s))")

    # ── Step 4c : Markdown ────────────────────────────────────────────────
    if out_formats.get("md"):
        log(f"[4c/4] Generating Markdown...")
        md_dir = out_root / "05_markdown"
        md_dir.mkdir(parents=True, exist_ok=True)
        md_path = md_dir / f"{stem}.md"
        md_str = elements_to_markdown(sem_pages, title)
        md_path.write_text(md_str, encoding="utf-8")
        log(f"  → {md_path}")

    # ── Step 4d : Assets (imgs + layout_vis) ─────────────────────────────
    if out_formats.get("assets"):
        log(f"[4d/4] Exporting assets (imgs/layout_vis)...")
        export_assets(src_path, sem_pages, out_root, log)

    return True


# ────────────────────────────────────────────────────────────────────────────
# TKINTER  GUI
# ────────────────────────────────────────────────────────────────────────────

ACCEPT_EXTS = (
    ("All supported",  "*.pdf *.docx *.txt *.md *.markdown "
                        "*.png *.jpg *.jpeg *.tiff *.bmp *.webp"),
    ("PDF",            "*.pdf"),
    ("Word",           "*.docx"),
    ("Text / Markdown","*.txt *.md *.markdown"),
    ("Images",         "*.png *.jpg *.jpeg *.tiff *.bmp *.webp"),
    ("All files",      "*.*"),
)

_SPINNER_FRAMES = ["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"]


class ConverterSuiteApp(tk.Tk):
    # ── Palette ──────────────────────────────────────────────────────────────
    BG      = "#EEF2FF"   # lavender-50
    CARD    = "#FFFFFF"
    IND     = "#4F46E5"   # indigo-600  (primary)
    IND_H   = "#4338CA"   # hover
    GRN     = "#16A34A"   # green-600   (Go)
    GRN_H   = "#15803D"
    RED     = "#DC2626"   # red-600     (Stop / error)
    RED_H   = "#B91C1C"
    AMB     = "#D97706"   # amber-600   (warning)
    TEXT    = "#1E1B4B"   # deep indigo
    MUTE    = "#6B7280"   # grey-500
    BDR     = "#C7D2FE"   # indigo-200
    LOG_BG  = "#0F172A"   # slate-900
    LOG_DEF = "#94A3B8"   # slate-400
    LOG_STEP= "#38BDF8"   # sky-400
    LOG_OK  = "#4ADE80"   # green-400
    LOG_ERR = "#F87171"   # red-400
    LOG_WARN= "#FBBF24"   # amber-400
    LOG_ARR = "#A78BFA"   # violet-400
    LOG_SEP = "#334155"   # slate-700
    LOG_FILE= "#7DD3FC"   # sky-300

    def __init__(self):
        super().__init__()
        self.title("S1000D Converter Suite")
        self.geometry("1020x860")
        self.configure(bg=self.BG)
        self.resizable(True, True)
        self.minsize(780, 620)

        # ── Runtime state ────────────────────────────────────────────────
        self.file_list: List[Path] = []
        self._stop_event  = threading.Event()
        self._spinner_idx = 0
        self._spinner_job = None
        self._running     = False
        self._t0          = 0.0
        self._elapsed_job = None

        # ── Tkinter variables ─────────────────────────────────────────────
        self.output_dir            = tk.StringVar()
        self.dm_type               = tk.StringVar(value="auto")
        self.force_ocr             = tk.BooleanVar(value=False)
        self.out_raw               = tk.BooleanVar(value=True)
        self.out_sem               = tk.BooleanVar(value=True)
        self.out_xml               = tk.BooleanVar(value=True)
        self.out_adoc              = tk.BooleanVar(value=True)
        self.out_md                = tk.BooleanVar(value=True)
        self.out_assets            = tk.BooleanVar(value=True)
        self.out_glm_default       = tk.BooleanVar(value=True)
        self.use_ollama_template   = tk.BooleanVar(value=USE_OLLAMA_TEMPLATE)
        self.ollama_model_var      = tk.StringVar(value=OLLAMA_MODEL)
        self.glmocr_backend_var    = tk.StringVar(value=GLMOCR_BACKEND)
        self.glmocr_ollama_url_var = tk.StringVar(value=GLMOCR_OLLAMA_URL)
        self.glmocr_ollama_model_var = tk.StringVar(value=GLMOCR_OLLAMA_MODEL)
        self.odl_use_hybrid_var    = tk.BooleanVar(value=ODL_USE_HYBRID)
        self.odl_hybrid_url_var    = tk.StringVar(value=ODL_HYBRID_URL)
        self.dm_desc_var           = tk.StringVar(value=DM_TYPE_DESC["auto"])
        self._progress_var         = tk.DoubleVar(value=0.0)
        self._status_var           = tk.StringVar(value="")
        self._stats_var            = tk.StringVar(value="")

        self._build_styles()
        self._build_ui()

    # ── Styles ───────────────────────────────────────────────────────────────

    def _build_styles(self):
        st = ttk.Style()
        st.theme_use("clam")
        B, C, I, G, R, T, M = (self.BG, self.CARD, self.IND, self.GRN,
                                self.RED, self.TEXT, self.MUTE)
        UI = ("Segoe UI", 9)
        UIB = ("Segoe UI", 9, "bold")

        st.configure("TFrame",          background=B)
        st.configure("Card.TFrame",     background=C, relief="flat")
        st.configure("TLabel",          background=B,  foreground=T, font=UI)
        st.configure("Card.TLabel",     background=C,  foreground=T, font=UI)
        st.configure("Mute.TLabel",     background=C,  foreground=M, font=("Segoe UI", 8, "italic"))
        st.configure("TCheckbutton",    background=C,  foreground=T, font=UI)
        st.configure("TMenubutton",     background=C,  foreground=T, font=UI, padding=4)
        st.configure("TEntry",          fieldbackground=C, foreground=T, font=UI)
        st.configure("TLabelframe",     background=C,  bordercolor=self.BDR,
                     relief="solid", borderwidth=1)
        st.configure("TLabelframe.Label", background=C, foreground=I, font=UIB)

        st.configure("TButton", font=UIB, background=I, foreground="#FFF",
                     padding=(10, 6), relief="flat", borderwidth=0)
        st.map("TButton",
               background=[("active", self.IND_H), ("disabled", "#9CA3AF")],
               foreground=[("disabled", "#D1D5DB")])

        st.configure("Go.TButton",  font=("Segoe UI", 13, "bold"),
                     background=G, foreground="#FFF", padding=(18, 10), relief="flat")
        st.map("Go.TButton",
               background=[("active", self.GRN_H), ("disabled", "#9CA3AF")])

        st.configure("Stop.TButton", font=("Segoe UI", 13, "bold"),
                     background=R, foreground="#FFF", padding=(18, 10), relief="flat")
        st.map("Stop.TButton",
               background=[("active", self.RED_H), ("disabled", "#9CA3AF")])

        st.configure("Accent.Horizontal.TProgressbar",
                     troughcolor="#E0E7FF", background=I,
                     borderwidth=0, thickness=10)

    # ── UI Construction ───────────────────────────────────────────────────────

    def _build_ui(self):
        main = ttk.Frame(self, padding="14 10 14 10", style="TFrame")
        main.pack(fill=tk.BOTH, expand=True)
        main.columnconfigure(0, weight=1)
        main.rowconfigure(5, weight=1)   # log row

        self._build_header(main)      # row 0
        self._build_input(main)       # row 1
        self._build_output(main)      # row 2
        self._build_options(main)     # row 3
        self._build_action(main)      # row 4
        self._build_log(main)         # row 5

    def _build_header(self, parent):
        hdr = ttk.Frame(parent, style="TFrame")
        hdr.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        tk.Label(hdr, text="S1000D Converter Suite",
                 font=("Segoe UI", 17, "bold"),
                 bg=self.BG, fg=self.IND).pack(side=tk.LEFT)
        tk.Label(hdr, text="  PDF · DOCX · Images → XML · AsciiDoc · Markdown",
                 font=("Segoe UI", 9), bg=self.BG, fg=self.MUTE
                 ).pack(side=tk.LEFT, pady=(6, 0))

    def _build_input(self, parent):
        frm = ttk.LabelFrame(parent, text="Input Files", padding="8 6")
        frm.grid(row=1, column=0, sticky="ew", pady=(0, 8))
        frm.columnconfigure(0, weight=1)

        self.file_listbox = tk.Listbox(
            frm, selectmode=tk.EXTENDED, height=5,
            font=("Consolas", 9), bg=self.CARD, fg=self.TEXT,
            selectbackground=self.IND, selectforeground="#FFF",
            relief="flat", borderwidth=0, activestyle="none",
            highlightthickness=1, highlightcolor=self.BDR,
            highlightbackground=self.BDR,
        )
        self.file_listbox.grid(row=0, column=0, columnspan=5,
                               sticky="ew", pady=(0, 6))

        for col, (txt, cmd) in enumerate([
            ("＋ Add Files",   self._add_files),
            ("📁 Add Folder",  self._add_folder),
            ("✕ Remove",       self._remove_selected),
            ("⊘ Clear All",    self._clear_files),
        ]):
            ttk.Button(frm, text=txt, command=cmd).grid(
                row=1, column=col, sticky="w", padx=(0, 6))

    def _build_output(self, parent):
        frm = ttk.LabelFrame(parent, text="Output", padding="8 6")
        frm.grid(row=2, column=0, sticky="ew", pady=(0, 8))
        frm.columnconfigure(1, weight=1)

        ttk.Label(frm, text="Output Folder:", style="Card.TLabel"
                  ).grid(row=0, column=0, sticky="w", padx=(0, 8))
        tk.Label(frm, textvariable=self.output_dir, anchor="w",
                 bg="#F5F3FF", fg=self.TEXT, relief="flat",
                 font=("Segoe UI", 9), padx=6, pady=4,
                 highlightthickness=1, highlightbackground=self.BDR,
                 ).grid(row=0, column=1, sticky="ew")
        ttk.Button(frm, text="Browse", command=self._pick_output
                   ).grid(row=0, column=2, padx=(8, 0))

    def _build_options(self, parent):
        frm = ttk.LabelFrame(parent, text="Options", padding="8 6")
        frm.grid(row=3, column=0, sticky="ew", pady=(0, 8))
        frm.columnconfigure(1, weight=1)

        sf = ("Segoe UI", 8, "italic")

        # DM type
        ttk.Label(frm, text="DM Type:", style="Card.TLabel"
                  ).grid(row=0, column=0, sticky="w", padx=(0, 8))
        dm_menu = ttk.OptionMenu(
            frm, self.dm_type, "auto", *S1000D_DM_TYPES,
            command=lambda v: self.dm_desc_var.set(DM_TYPE_DESC.get(v, "")))
        dm_menu.config(width=26)
        dm_menu.grid(row=0, column=1, sticky="w", pady=2)
        ttk.Label(frm, textvariable=self.dm_desc_var,
                  background=self.CARD, foreground=self.MUTE, font=sf
                  ).grid(row=0, column=2, sticky="w", padx=(8, 0))

        ttk.Checkbutton(frm, text="Force OCR  (glmocr)",
                        variable=self.force_ocr
                        ).grid(row=1, column=0, columnspan=2, sticky="w", pady=2)
        ttk.Checkbutton(frm, text="Use Ollama to structure AsciiDoc templates",
                        variable=self.use_ollama_template
                        ).grid(row=2, column=0, columnspan=2, sticky="w", pady=2)

        # GLM backend
        ttk.Label(frm, text="GLM OCR Backend:", style="Card.TLabel"
                  ).grid(row=3, column=0, sticky="w", padx=(0, 8), pady=(4, 2))
        be_menu = ttk.OptionMenu(frm, self.glmocr_backend_var,
                                 self.glmocr_backend_var.get(), "default", "ollama")
        be_menu.config(width=14)
        be_menu.grid(row=3, column=1, sticky="w", pady=(4, 2))

        for r, (lbl, var, w) in enumerate([
            ("Ollama Model:",           self.ollama_model_var,       28),
            ("GLM OCR Ollama URL:",     self.glmocr_ollama_url_var,  42),
            ("GLM OCR Ollama Model:",   self.glmocr_ollama_model_var, 42),
        ], start=4):
            ttk.Label(frm, text=lbl, style="Card.TLabel"
                      ).grid(row=r, column=0, sticky="w", padx=(0, 8), pady=(2, 0))
            ttk.Entry(frm, textvariable=var, width=w
                      ).grid(row=r, column=1, columnspan=2, sticky="w", pady=(2, 0))

        # ── ODL Hybrid ────────────────────────────────────────────────────────
        sep = ttk.Label(frm, text="── OpenDataLoader Fallback ──",
                        background=self.CARD, foreground=self.MUTE,
                        font=("Segoe UI", 8, "italic"))
        sep.grid(row=7, column=0, columnspan=3, sticky="w", pady=(8, 2))

        ttk.Checkbutton(frm, text="ODL Hybrid mode  (requires opendataloader-pdf-hybrid server)",
                        variable=self.odl_use_hybrid_var
                        ).grid(row=8, column=0, columnspan=3, sticky="w", pady=(0, 2))

        ttk.Label(frm, text="ODL Hybrid URL:", style="Card.TLabel"
                  ).grid(row=9, column=0, sticky="w", padx=(0, 8), pady=(2, 0))
        ttk.Entry(frm, textvariable=self.odl_hybrid_url_var, width=42
                  ).grid(row=9, column=1, columnspan=2, sticky="w", pady=(2, 0))

        # Output format checkboxes
        fmt = ttk.Frame(frm, style="Card.TFrame")
        fmt.grid(row=10, column=0, columnspan=3, sticky="w", pady=(6, 0))
        ttk.Label(fmt, text="Outputs:", style="Card.TLabel").pack(side=tk.LEFT, padx=(0, 8))
        for txt, var in [
            ("Raw JSON",      self.out_raw),
            ("Semantic JSON", self.out_sem),
            ("S1000D XML",    self.out_xml),
            ("AsciiDoc",      self.out_adoc),
            ("Markdown",      self.out_md),
            ("Assets",        self.out_assets),
            ("GLM Default",   self.out_glm_default),
        ]:
            ttk.Checkbutton(fmt, text=txt, variable=var).pack(side=tk.LEFT, padx=(0, 6))

        # Layout label settings button
        ttk.Button(frm, text="Layout Labels…", command=self._open_layout_settings
                   ).grid(row=11, column=0, sticky="w", pady=(8, 0))

    # ── Layout label settings ─────────────────────────────────────────────────

    _ALL_LABELS = [
        "abstract", "algorithm", "aside_text", "chart", "content",
        "display_formula", "doc_title", "figure_title", "footer", "footer_image",
        "footnote", "formula_number", "header", "header_image", "image",
        "inline_formula", "number", "paragraph_title", "reference",
        "reference_content", "seal", "table", "text", "vertical_text",
        "vision_footnote",
    ]
    _CATEGORIES = ["text", "table", "formula", "skip", "abandon"]

    _DEFAULT_MAPPING = {
        "text":    ["abstract", "algorithm", "content", "doc_title", "figure_title",
                    "paragraph_title", "reference_content", "text", "vertical_text",
                    "vision_footnote", "seal", "formula_number", "header", "footer", "footnote"],
        "table":   ["table"],
        "formula": ["display_formula", "inline_formula"],
        "skip":    ["chart", "image"],
        "abandon": ["number", "aside_text", "reference", "footer_image", "header_image"],
    }

    def _config_yaml_path(self) -> Optional[Path]:
        """Locate glmocr/config.yaml relative to this script."""
        here = Path(__file__).parent
        for candidate in [
            here / "glmocr" / "config.yaml",
            here.parent / "glmocr" / "config.yaml",
        ]:
            if candidate.exists():
                return candidate
        return None

    def _load_layout_mapping(self) -> Dict[str, List[str]]:
        """Read label_task_mapping from config.yaml; fall back to defaults."""
        p = self._config_yaml_path()
        if p is None:
            return {k: list(v) for k, v in self._DEFAULT_MAPPING.items()}
        try:
            import yaml
            with open(p, encoding="utf-8") as f:
                cfg = yaml.safe_load(f)
            raw = (cfg.get("pipeline", {})
                      .get("layout", {})
                      .get("label_task_mapping", {}))
            if raw:
                return {cat: list(lbls or []) for cat, lbls in raw.items()}
        except Exception:
            pass
        return {k: list(v) for k, v in self._DEFAULT_MAPPING.items()}

    def _save_layout_mapping(self, mapping: Dict[str, List[str]]) -> bool:
        """Write label_task_mapping back into config.yaml."""
        p = self._config_yaml_path()
        if p is None:
            messagebox.showerror("Save failed",
                                 "Could not find glmocr/config.yaml next to this script.")
            return False
        try:
            import yaml
            with open(p, encoding="utf-8") as f:
                cfg = yaml.safe_load(f) or {}
            cfg.setdefault("pipeline", {}).setdefault("layout", {})["label_task_mapping"] = mapping
            with open(p, "w", encoding="utf-8") as f:
                yaml.dump(cfg, f, default_flow_style=False, allow_unicode=True, sort_keys=False)
            return True
        except ImportError:
            messagebox.showerror("Missing library",
                                 "PyYAML is required to save settings.\n"
                                 "Run:  pip install pyyaml")
            return False
        except Exception as e:
            messagebox.showerror("Save failed", str(e))
            return False

    def _open_layout_settings(self):
        """Open the Layout Label Settings dialog."""
        dlg = tk.Toplevel(self)
        dlg.title("Layout Label Settings")
        dlg.resizable(False, False)
        dlg.grab_set()

        B = self.BG; C = self.CARD; I = self.IND; T = self.TEXT; M = self.MUTE

        dlg.configure(bg=B)

        current = self._load_layout_mapping()

        # Build label→category map
        label_cat: Dict[str, tk.StringVar] = {}
        for cat, labels in current.items():
            for lbl in labels:
                if lbl in self._ALL_LABELS:
                    label_cat[lbl] = tk.StringVar(value=cat)
        # Any label not in config gets 'abandon' as default
        for lbl in self._ALL_LABELS:
            if lbl not in label_cat:
                label_cat[lbl] = tk.StringVar(value="abandon")

        # ── Header ───────────────────────────────────────────────────────────
        hdr = tk.Frame(dlg, bg=I, pady=8)
        hdr.pack(fill=tk.X)
        tk.Label(hdr, text="Layout Label Settings",
                 bg=I, fg="#FFF", font=("Segoe UI", 11, "bold")).pack()
        tk.Label(hdr, text="Assign each detected region label to a processing category.",
                 bg=I, fg="#C7D2FE", font=("Segoe UI", 8)).pack()

        # ── Category legend ───────────────────────────────────────────────────
        leg = tk.Frame(dlg, bg=B, pady=4)
        leg.pack(fill=tk.X, padx=16)
        cat_colors = {
            "text":    "#16A34A",
            "table":   "#2563EB",
            "formula": "#7C3AED",
            "skip":    "#D97706",
            "abandon": "#DC2626",
        }
        cat_tips = {
            "text":    "OCR as plain text",
            "table":   "OCR with table prompt",
            "formula": "OCR with formula prompt",
            "skip":    "Keep region, no OCR",
            "abandon": "Discard entirely",
        }
        tk.Label(leg, text="Categories:", bg=B, fg=T,
                 font=("Segoe UI", 8, "bold")).pack(side=tk.LEFT, padx=(0, 8))
        for cat in self._CATEGORIES:
            tk.Label(leg, text=f"● {cat}  ({cat_tips[cat]})",
                     bg=B, fg=cat_colors[cat],
                     font=("Segoe UI", 8)).pack(side=tk.LEFT, padx=(0, 12))

        # ── Scrollable label grid ─────────────────────────────────────────────
        canvas = tk.Canvas(dlg, bg=C, highlightthickness=0, width=520, height=460)
        sb = ttk.Scrollbar(dlg, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=sb.set)

        sb.pack(side=tk.RIGHT, fill=tk.Y, padx=(0, 8), pady=8)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(16, 0), pady=8)

        inner = tk.Frame(canvas, bg=C)
        win_id = canvas.create_window((0, 0), window=inner, anchor="nw")

        def _on_resize(e):
            canvas.itemconfig(win_id, width=e.width)
        canvas.bind("<Configure>", _on_resize)

        def _on_frame_configure(e):
            canvas.configure(scrollregion=canvas.bbox("all"))
        inner.bind("<Configure>", _on_frame_configure)

        def _on_mousewheel(e):
            canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        # Column headers
        for c_idx, (txt, fg) in enumerate([
            ("Label",    T),
            ("Category", T),
            ("Preview",  M),
        ]):
            tk.Label(inner, text=txt, bg=C, fg=fg,
                     font=("Segoe UI", 8, "bold"),
                     width=18 if c_idx == 0 else (14 if c_idx == 1 else 24),
                     anchor="w"
                     ).grid(row=0, column=c_idx, padx=(8, 4), pady=(4, 2), sticky="w")

        tk.Frame(inner, bg=self.BDR, height=1).grid(
            row=1, column=0, columnspan=3, sticky="ew", padx=8, pady=(0, 4))

        for r_idx, lbl in enumerate(sorted(self._ALL_LABELS), start=2):
            var = label_cat[lbl]

            def _make_preview(v=var, lbl_frame=None):
                def _update(*_):
                    cat = v.get()
                    if lbl_frame:
                        lbl_frame.config(fg=cat_colors.get(cat, M),
                                         text=f"→ {cat_tips.get(cat, cat)}")
                return _update

            # Label name
            tk.Label(inner, text=lbl, bg=C, fg=T,
                     font=("Segoe UI", 9), width=20, anchor="w"
                     ).grid(row=r_idx, column=0, padx=(8, 4), pady=1, sticky="w")

            # Category dropdown
            om = ttk.OptionMenu(inner, var, var.get(), *self._CATEGORIES)
            om.config(width=10)
            om.grid(row=r_idx, column=1, padx=4, pady=1, sticky="w")

            # Preview label (updates on change)
            preview = tk.Label(inner, text=f"→ {cat_tips.get(var.get(), var.get())}",
                               bg=C, fg=cat_colors.get(var.get(), M),
                               font=("Segoe UI", 8, "italic"), width=26, anchor="w")
            preview.grid(row=r_idx, column=2, padx=4, pady=1, sticky="w")

            updater = _make_preview(var, preview)
            var.trace_add("write", updater)

        # ── Footer buttons ────────────────────────────────────────────────────
        foot = tk.Frame(dlg, bg=B, pady=8)
        foot.pack(fill=tk.X, padx=16)

        def _reset():
            if messagebox.askyesno("Reset defaults",
                                   "Reset all labels to default categories?",
                                   parent=dlg):
                for lbl in self._ALL_LABELS:
                    for cat, lbls in self._DEFAULT_MAPPING.items():
                        if lbl in lbls:
                            label_cat[lbl].set(cat)
                            break
                    else:
                        label_cat[lbl].set("abandon")

        def _save():
            new_mapping: Dict[str, List[str]] = {cat: [] for cat in self._CATEGORIES}
            for lbl, var in label_cat.items():
                new_mapping[var.get()].append(lbl)
            if self._save_layout_mapping(new_mapping):
                messagebox.showinfo("Saved",
                                    "Layout label settings saved to config.yaml.\n"
                                    "Changes take effect on the next OCR run.",
                                    parent=dlg)
                dlg.destroy()

        ttk.Button(foot, text="Reset Defaults", command=_reset
                   ).pack(side=tk.LEFT)
        ttk.Button(foot, text="Cancel", command=dlg.destroy
                   ).pack(side=tk.RIGHT, padx=(8, 0))
        ttk.Button(foot, text="Save", command=_save, style="TButton"
                   ).pack(side=tk.RIGHT)

    def _build_action(self, parent):
        frm = ttk.LabelFrame(parent, text="Conversion", padding="10 8")
        frm.grid(row=4, column=0, sticky="ew", pady=(0, 8))
        frm.columnconfigure(1, weight=1)

        # ── Button row ──────────────────────────────────────────────────
        btn_row = ttk.Frame(frm, style="Card.TFrame")
        btn_row.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        btn_row.columnconfigure(2, weight=1)

        self.start_btn = ttk.Button(
            btn_row, text="🚀  Convert All Files",
            command=self._start_thread, style="Go.TButton")
        self.start_btn.grid(row=0, column=0, padx=(0, 8))

        self.stop_btn = ttk.Button(
            btn_row, text="⏹  Stop",
            command=self._stop_conversion, style="Stop.TButton",
            state=tk.DISABLED)
        self.stop_btn.grid(row=0, column=1)

        self._stats_label = tk.Label(
            btn_row, textvariable=self._stats_var,
            bg=self.CARD, fg=self.MUTE, font=("Segoe UI", 9), anchor="e")
        self._stats_label.grid(row=0, column=2, sticky="e", padx=(8, 0))

        # ── Progress bar ─────────────────────────────────────────────────
        self.progress_bar = ttk.Progressbar(
            frm, variable=self._progress_var,
            maximum=100.0, mode="determinate",
            style="Accent.Horizontal.TProgressbar")
        self.progress_bar.grid(row=1, column=0, columnspan=2,
                               sticky="ew", pady=(0, 4))

        # ── Status line ──────────────────────────────────────────────────
        self._status_label = tk.Label(
            frm, textvariable=self._status_var,
            bg=self.CARD, fg=self.IND, font=("Segoe UI", 9), anchor="w")
        self._status_label.grid(row=2, column=0, columnspan=2, sticky="ew")

    def _build_log(self, parent):
        frm = ttk.LabelFrame(parent, text="Conversion Log", padding="4 4")
        frm.grid(row=5, column=0, sticky="nsew")
        frm.rowconfigure(0, weight=1)
        frm.columnconfigure(0, weight=1)

        self.log_widget = scrolledtext.ScrolledText(
            frm, wrap=tk.WORD, font=("Consolas", 9),
            state=tk.DISABLED,
            bg=self.LOG_BG, fg=self.LOG_DEF,
            insertbackground=self.LOG_DEF,
            relief="flat", padx=8, pady=6,
        )
        self.log_widget.grid(row=0, column=0, sticky="nsew")

        # Color tags for log entries
        self.log_widget.tag_configure("step",  foreground=self.LOG_STEP)
        self.log_widget.tag_configure("ok",    foreground=self.LOG_OK)
        self.log_widget.tag_configure("err",   foreground=self.LOG_ERR)
        self.log_widget.tag_configure("warn",  foreground=self.LOG_WARN)
        self.log_widget.tag_configure("arrow", foreground=self.LOG_ARR)
        self.log_widget.tag_configure("sep",   foreground=self.LOG_SEP)
        self.log_widget.tag_configure("file",  foreground=self.LOG_FILE)

    # ── File management ───────────────────────────────────────────────────────

    def _add_files(self):
        for p in filedialog.askopenfilenames(filetypes=ACCEPT_EXTS):
            fp = Path(p)
            if fp not in self.file_list:
                self.file_list.append(fp)
                self.file_listbox.insert(tk.END, f"  {fp.name}  ({fp.parent})")

    def _add_folder(self):
        folder = filedialog.askdirectory(title="Add all supported files from folder")
        if not folder:
            return
        for fp in Path(folder).rglob("*"):
            if fp.suffix.lower() in ({".pdf", ".docx"} | TEXT_EXTS | MD_EXTS | IMAGE_EXTS
                                     ) and fp not in self.file_list:
                self.file_list.append(fp)
                self.file_listbox.insert(tk.END, f"  {fp.name}  ({fp.parent})")

    def _remove_selected(self):
        for idx in reversed(list(self.file_listbox.curselection())):
            self.file_listbox.delete(idx)
            self.file_list.pop(idx)

    def _clear_files(self):
        self.file_list.clear()
        self.file_listbox.delete(0, tk.END)

    def _pick_output(self):
        d = filedialog.askdirectory(title="Select base output folder")
        if d:
            self.output_dir.set(d)

    def _mark_file(self, idx: int, state: str):
        """Color-code a listbox row: processing/ok/skip/error."""
        colors = {
            "processing": ("#FEF3C7", self.AMB),  # amber
            "ok":         ("#DCFCE7", self.GRN),  # green
            "skip":       ("#F3F4F6", self.MUTE), # grey
            "error":      ("#FEE2E2", self.RED),  # red
        }
        bg, fg = colors.get(state, (self.CARD, self.TEXT))
        try:
            self.file_listbox.itemconfig(idx, bg=bg, fg=fg)
        except Exception:
            pass

    # ── Log helpers ───────────────────────────────────────────────────────────

    def _log(self, msg: str):
        self.after(0, self._write_log, msg)

    def _write_log(self, msg: str):
        w = self.log_widget
        w.config(state=tk.NORMAL)
        line = msg + "\n"

        # Pick a color tag based on message content
        m = msg.strip()
        if m.startswith("─") or m.startswith("="):
            tag = "sep"
        elif re.match(r"^\[([1-9]|4[abc])/4\]", m):
            tag = "step"
        elif re.search(r"\b(ok|done|finished|complete|saved|written)\b", m, re.I):
            tag = "ok"
        elif re.search(r"\b(error|exception|failed|cannot|critical)\b", m, re.I):
            tag = "err"
        elif re.search(r"\b(warning|warn|skipping|fallback|deprecated)\b", m, re.I):
            tag = "warn"
        elif m.startswith("  →") or m.startswith("→"):
            tag = "arrow"
        elif re.search(r"Extracting:|Output folder:|Files to process:|[1/4]", m):
            tag = "file"
        else:
            tag = None

        if tag:
            w.insert(tk.END, line, tag)
        else:
            w.insert(tk.END, line)

        w.config(state=tk.DISABLED)
        w.see(tk.END)

    # ── Progress / spinner animation ──────────────────────────────────────────

    def _start_spinner(self):
        self._spinner_idx = 0
        self._tick_spinner()

    def _tick_spinner(self):
        if not self._running:
            return
        self._spinner_idx = (self._spinner_idx + 1) % len(_SPINNER_FRAMES)
        # Refresh the status line prefix; the rest is set by _set_status
        cur = self._status_var.get()
        # Replace leading spinner char
        if cur and cur[0] in "⠋⠙⠹⠸⠼⠴⠦⠧⠇⠏":
            self._status_var.set(_SPINNER_FRAMES[self._spinner_idx] + cur[1:])
        self._spinner_job = self.after(90, self._tick_spinner)

    def _stop_spinner(self):
        if self._spinner_job:
            self.after_cancel(self._spinner_job)
            self._spinner_job = None

    def _set_status(self, text: str):
        spin = _SPINNER_FRAMES[self._spinner_idx]
        self.after(0, self._status_var.set, f"{spin}  {text}")

    def _tick_elapsed(self):
        if not self._running:
            return
        elapsed = time.perf_counter() - self._t0
        cur = self._stats_var.get()
        # Keep file count prefix, update time
        parts = cur.split("·")
        prefix = parts[0].strip() if parts else ""
        if prefix:
            self.after(0, self._stats_var.set, f"{prefix}  ·  {elapsed:.0f}s elapsed")
        self._elapsed_job = self.after(500, self._tick_elapsed)

    def _set_progress(self, value: float, file_idx: int, total: int,
                      step: int, total_steps: int, step_name: str):
        self._progress_var.set(max(0.0, min(100.0, value)))
        self._stats_var.set(
            f"File {file_idx + 1} / {total}  ·  Step {step}/{total_steps}"
        )
        self._set_status(f"{step_name}  —  {self.file_list[file_idx].name}")

    # ── Conversion thread ─────────────────────────────────────────────────────

    def _start_thread(self):
        if not self.file_list:
            messagebox.showwarning("No input", "Add at least one input file first.")
            return
        if not self.output_dir.get():
            messagebox.showwarning("No output", "Select an output folder first.")
            return
        self._stop_event.clear()
        self._running = True
        self._t0 = time.perf_counter()
        self._progress_var.set(0.0)
        self._stats_var.set("")
        self._status_var.set("")
        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self._start_spinner()
        self._tick_elapsed()
        threading.Thread(target=self._run, daemon=True).start()

    def _stop_conversion(self):
        self._stop_event.set()
        self.stop_btn.config(state=tk.DISABLED)
        self._set_status("Stop requested — finishing current step…")

    def _run(self):
        out_root = Path(self.output_dir.get()) / "S1000D_Package"
        out_root.mkdir(parents=True, exist_ok=True)

        dm_type   = self.dm_type.get()
        force_ocr = self.force_ocr.get()
        global OLLAMA_MODEL, USE_OLLAMA_TEMPLATE
        global GLMOCR_BACKEND, GLMOCR_OLLAMA_URL, GLMOCR_OLLAMA_MODEL
        global ODL_USE_HYBRID, ODL_HYBRID_URL
        OLLAMA_MODEL = self.ollama_model_var.get().strip() or OLLAMA_MODEL
        USE_OLLAMA_TEMPLATE = bool(self.use_ollama_template.get())
        GLMOCR_BACKEND = (self.glmocr_backend_var.get().strip().lower() or "default")
        if GLMOCR_BACKEND not in {"default", "ollama"}:
            GLMOCR_BACKEND = "default"
        GLMOCR_OLLAMA_URL   = self.glmocr_ollama_url_var.get().strip()   or GLMOCR_OLLAMA_URL
        GLMOCR_OLLAMA_MODEL = self.glmocr_ollama_model_var.get().strip() or GLMOCR_OLLAMA_MODEL
        ODL_USE_HYBRID = bool(self.odl_use_hybrid_var.get())
        ODL_HYBRID_URL = self.odl_hybrid_url_var.get().strip() or ODL_HYBRID_URL
        out_formats = {
            "raw_json":    self.out_raw.get(),
            "sem_json":    self.out_sem.get(),
            "xml":         self.out_xml.get(),
            "adoc":        self.out_adoc.get(),
            "md":          self.out_md.get(),
            "assets":      self.out_assets.get(),
            "glm_default": self.out_glm_default.get(),
        }

        report_lines: List[str] = [
            "S1000D Converter Suite – Conversion Report",
            "=" * 50,
            f"Output root : {out_root}",
            f"DM type     : {dm_type}",
            f"Force OCR   : {force_ocr}",
            f"OCR backend : {GLMOCR_BACKEND}",
            f"Ollama pass : {USE_OLLAMA_TEMPLATE}",
            f"Ollama model: {OLLAMA_MODEL}",
        ] + ([
            f"OCR Ollama URL  : {GLMOCR_OLLAMA_URL}",
            f"OCR Ollama model: {GLMOCR_OLLAMA_MODEL}",
        ] if GLMOCR_BACKEND == "ollama" else []) + [
            f"Outputs     : {', '.join(k for k, v in out_formats.items() if v)}",
            "",
        ]

        total = len(self.file_list)
        self._log(f"Output folder  : {out_root}")
        self._log(f"Files to process: {total}\n")

        ok = fail = 0
        t0 = time.perf_counter()

        for file_idx, src in enumerate(self.file_list):
            if self._stop_event.is_set():
                self._log("⏹  Conversion stopped by user.")
                break

            self.after(0, self._mark_file, file_idx, "processing")

            def _prog(step, total_steps, step_name, _idx=file_idx, _tot=total):
                pct = (_idx / _tot + step / (total_steps * _tot)) * 100
                self.after(0, self._set_progress,
                           pct, _idx, _tot, step, total_steps, step_name)

            try:
                success = build_package(
                    src_path=src,
                    out_root=out_root,
                    dm_type=dm_type,
                    force_ocr=force_ocr,
                    out_formats=out_formats,
                    log=self._log,
                    stop_event=self._stop_event,
                    progress_cb=_prog,
                )
                status = "OK" if success else "SKIPPED"
                state  = "ok" if success else "skip"
                if success:
                    ok += 1
                else:
                    fail += 1
            except Exception as e:
                self._log(f"  EXCEPTION for {src.name}:\n  {e}")
                traceback.print_exc()
                status, state = "ERROR", "error"
                fail += 1

            self.after(0, self._mark_file, file_idx, state)
            report_lines.append(f"  [{status:7s}]  {src.name}")

        elapsed = time.perf_counter() - t0
        summary = (
            f"\n✓ {ok} OK  ✗ {fail} failed  ({elapsed:.1f}s)\n"
            f"Output: {out_root}"
        )
        self._log(summary)
        report_lines += ["", summary]

        rpt = out_root / "conversion_report.txt"
        rpt.write_text("\n".join(report_lines), encoding="utf-8")
        self._log(f"\nReport written to: {rpt}")

        self._running = False
        self._stop_spinner()
        if self._elapsed_job:
            self.after_cancel(self._elapsed_job)
            self._elapsed_job = None

        stopped = self._stop_event.is_set()
        final_txt = "⏹  Stopped" if stopped else f"✓  Done — {ok}/{total} files converted"
        self.after(0, self._status_var.set, final_txt)
        self.after(0, self._progress_var.set, 0.0 if stopped else 100.0)
        self.after(0, self.start_btn.config, {"state": tk.NORMAL})
        self.after(0, self.stop_btn.config,  {"state": tk.DISABLED})


# ────────────────────────────────────────────────────────────────────────────

def main():
    app = ConverterSuiteApp()
    app.mainloop()


if __name__ == "__main__":
    main()